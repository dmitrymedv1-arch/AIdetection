import streamlit as st
import re
import numpy as np
import pandas as pd
import plotly.graph_objects as go
import plotly.express as px
from collections import Counter
from typing import List, Dict, Tuple, Optional, Any
import hashlib
import tempfile
import os
import subprocess
import sys
import warnings
import time
import io
from datetime import datetime
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import hashlib
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, Image

# Проверка наличия reportlab
REPORTLAB_AVAILABLE = False
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False
    # Создаем заглушку для функции
    def generate_pdf_report(results_data, topic_name="CT(A)I-detector Analysis"):
        return None

# Если reportlab не установлен, показываем предупреждение в интерфейсе
if not REPORTLAB_AVAILABLE:
    st.sidebar.warning(
        "⚠️ PDF export requires reportlab. Install with: pip install reportlab"
    )

# Force CSS injection
st.markdown("""
<style>
    .metric-grid {
        display: grid;
        grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
        gap: 1rem;
        margin: 1rem 0;
    }
    .metric-card {
        background: white;
        border-radius: 16px;
        padding: 1.25rem;
        box-shadow: 0 5px 15px rgba(0,0,0,0.05);
    }
    .metric-title {
        font-size: 0.9rem;
        color: #666;
        margin-bottom: 0.5rem;
        text-transform: uppercase;
    }
    .metric-value {
        font-size: 2rem;
        font-weight: 600;
        color: #333;
    }
    .metric-unit {
        font-size: 0.9rem;
        color: #999;
        margin-left: 0.25rem;
    }
</style>
""", unsafe_allow_html=True)

# Отключаем все предупреждения spaCy и pydantic до импорта
warnings.filterwarnings('ignore', category=UserWarning)
warnings.filterwarnings('ignore', category=DeprecationWarning)
warnings.filterwarnings('ignore', category=FutureWarning)
warnings.filterwarnings('ignore', module='spacy')
warnings.filterwarnings('ignore', module='pydantic')
warnings.filterwarnings('ignore', module='thinc')
warnings.filterwarnings('ignore', module='blis')

# Исправление для pydantic v1 и v2 совместимости
import pydantic
if hasattr(pydantic, 'v1'):
    from pydantic.v1 import BaseModel, ConfigDict
else:
    from pydantic import BaseModel

# NLP библиотеки с улучшенной обработкой ошибок
SPACY_AVAILABLE = False
try:
    # Принудительно устанавливаем переменные окружения для spaCy
    os.environ['SPACY_WARNING_IGNORE'] = '1'
    os.environ['TOKENIZERS_PARALLELISM'] = 'false'
    
    # Подавляем все возможные предупреждения при импорте
    with warnings.catch_warnings(record=True) as w:
        warnings.simplefilter("ignore")
        
        # Пытаемся импортировать spacy с правильной конфигурацией
        import spacy
        from spacy import displacy
        
        # Проверяем, загружена ли модель
        try:
            # Пробуем загрузить маленькую модель, если она есть
            nlp = spacy.load("en_core_web_sm", disable=["parser", "ner"])
            SPACY_AVAILABLE = True
            print("spaCy успешно загружен с моделью en_core_web_sm")
        except OSError:
            # Если модели нет, пытаемся скачать (но не блокируем приложение)
            try:
                print("Модель en_core_web_sm не найдена, но spaCy доступен")
                SPACY_AVAILABLE = True  # Библиотека доступна, даже без модели
            except:
                SPACY_AVAILABLE = True
        except Exception as e:
            print(f"Ошибка при загрузке модели spaCy: {e}")
            SPACY_AVAILABLE = True  # Библиотека все еще может быть доступна
            
except ImportError as e:
    SPACY_AVAILABLE = False
    # Тихий импорт - не показываем ошибку пользователю
except Exception as e:
    # Ловим любые другие ошибки
    SPACY_AVAILABLE = False

# Если spaCy не доступен, создаем заглушку для функций, которые его используют
if not SPACY_AVAILABLE:
    class SpacyStub:
        def __getattr__(self, name):
            return None
    spacy = SpacyStub()

# Для transformers - загружаем с обработкой ошибок
TRANSFORMERS_AVAILABLE = False
try:
    with warnings.catch_warnings():
        warnings.simplefilter("ignore")
        from transformers import AutoTokenizer, AutoModelForCausalLM, pipeline, AutoModelForSequenceClassification
        import torch
        TRANSFORMERS_AVAILABLE = True
except ImportError:
    TRANSFORMERS_AVAILABLE = False
except Exception:
    TRANSFORMERS_AVAILABLE = False

# Для sentence-transformers
SENTENCE_TRANSFORMERS_AVAILABLE = False
try:
    with warnings.catch_warnings():
        warnings.simplefilter("ignore")
        from sentence_transformers import SentenceTransformer
        from sklearn.metrics.pairwise import cosine_similarity
        SENTENCE_TRANSFORMERS_AVAILABLE = True
except ImportError:
    SENTENCE_TRANSFORMERS_AVAILABLE = False
except Exception:
    SENTENCE_TRANSFORMERS_AVAILABLE = False

# Для lexical diversity
LEXICAL_DIVERSITY_AVAILABLE = False
try:
    with warnings.catch_warnings():
        warnings.simplefilter("ignore")
        import lexical_diversity as lexdiv
        LEXICAL_DIVERSITY_AVAILABLE = True
except ImportError:
    LEXICAL_DIVERSITY_AVAILABLE = False
except Exception:
    LEXICAL_DIVERSITY_AVAILABLE = False

# Подавляем все оставшиеся предупреждения после импорта
warnings.filterwarnings('ignore')

# Page configuration
st.set_page_config(
    page_title="CT(A)I-detector",
    page_icon="🔬",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Modern lab-style design
st.markdown("""
<style>
    /* Modern lab design - clean, scientific, professional */
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');
    
    * {
        font-family: 'Inter', sans-serif;
    }
    
    .stApp {
        background: linear-gradient(135deg, #f5f7fa 0%, #e9ecf2 100%);
    }
    
    .main > div {
        padding: 0;
        background: transparent;
        box-shadow: none;
    }
    
    /* Step container */
    .step-container {
        background: white;
        border-radius: 24px;
        padding: 2rem;
        margin: 1rem 0;
        box-shadow: 0 10px 30px rgba(0,0,0,0.05);
        border: 1px solid rgba(255,255,255,0.5);
        backdrop-filter: blur(10px);
    }
    
    /* Step indicators */
    .step-indicator {
        display: flex;
        justify-content: space-between;
        margin-bottom: 3rem;
        position: relative;
        padding: 0 1rem;
    }
    
    .step-indicator::before {
        content: '';
        position: absolute;
        top: 24px;
        left: 80px;
        right: 80px;
        height: 2px;
        background: #e0e0e0;
        z-index: 0;
    }
    
    .step-item {
        text-align: center;
        position: relative;
        z-index: 1;
        flex: 1;
    }
    
    .step-circle {
        width: 48px;
        height: 48px;
        background: white;
        border: 2px solid #e0e0e0;
        border-radius: 50%;
        display: flex;
        align-items: center;
        justify-content: center;
        margin: 0 auto 0.5rem;
        font-weight: 600;
        color: #666;
        background: white;
        transition: all 0.3s ease;
    }
    
    .step-circle.active {
        background: #667eea;
        border-color: #667eea;
        color: white;
        box-shadow: 0 5px 15px rgba(102,126,234,0.3);
    }
    
    .step-circle.completed {
        background: #10b981;
        border-color: #10b981;
        color: white;
    }
    
    .step-label {
        font-size: 0.9rem;
        font-weight: 500;
        color: #666;
    }
    
    .step-label.active {
        color: #667eea;
        font-weight: 600;
    }
    
    /* Upload area */
    .upload-area {
        background: linear-gradient(135deg, #f8faff 0%, #f0f3fd 100%);
        border: 2px dashed #667eea;
        border-radius: 20px;
        padding: 3rem;
        text-align: center;
        transition: all 0.3s ease;
        cursor: pointer;
    }
    
    .upload-area:hover {
        border-color: #5a67d8;
        background: linear-gradient(135deg, #f0f4ff 0%, #e8ecfe 100%);
        transform: translateY(-2px);
    }
    
    .upload-icon {
        font-size: 3rem;
        color: #667eea;
        margin-bottom: 1rem;
    }
    
    /* Loading animation */
    .lab-loader {
        text-align: center;
        padding: 2rem;
    }
    
    .analyzer-beam {
        width: 100%;
        height: 4px;
        background: linear-gradient(90deg, transparent, #667eea, transparent);
        animation: scan 2s infinite;
        margin: 2rem 0;
    }
    
    @keyframes scan {
        0% { transform: translateX(-100%); }
        100% { transform: translateX(100%); }
    }
    
    .module-status {
        display: flex;
        flex-wrap: wrap;
        gap: 1rem;
        justify-content: center;
        margin: 2rem 0;
    }
    
    .module-pill {
        background: white;
        border: 1px solid #e0e0e0;
        border-radius: 30px;
        padding: 0.5rem 1.5rem;
        font-size: 0.9rem;
        color: #666;
        transition: all 0.3s ease;
    }
    
    .module-pill.active {
        background: #667eea;
        border-color: #667eea;
        color: white;
        box-shadow: 0 5px 15px rgba(102,126,234,0.3);
    }
    
    .module-pill.completed {
        background: #10b981;
        border-color: #10b981;
        color: white;
    }
    
    /* Results dashboard */
    .score-card {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        border-radius: 24px;
        padding: 2rem;
        color: white;
        box-shadow: 0 20px 40px rgba(102,126,234,0.3);
    }
    
    .score-number {
        font-size: 5rem;
        font-weight: 700;
        line-height: 1;
        margin-bottom: 0.5rem;
    }
    
    .score-label {
        font-size: 1.2rem;
        opacity: 0.9;
        text-transform: uppercase;
        letter-spacing: 2px;
    }
    
    .verdict-badge {
        display: inline-block;
        padding: 0.5rem 1.5rem;
        border-radius: 30px;
        font-weight: 600;
        font-size: 1.1rem;
        margin-top: 1rem;
    }
    
    .verdict-high {
        background: #fee;
        color: #c00;
    }
    
    .verdict-medium {
        background: #fff3e0;
        color: #f90;
    }
    
    .verdict-low {
        background: #e8f5e9;
        color: #2e7d32;
    }
    
    /* Module spectrum */
    .spectrum-container {
        background: white;
        border-radius: 16px;
        padding: 1.5rem;
        margin: 1rem 0;
        box-shadow: 0 5px 15px rgba(0,0,0,0.05);
    }
    
    .spectrum-bar {
        height: 40px;
        border-radius: 20px;
        background: linear-gradient(90deg, #10b981, #fbbf24, #ef4444);
        margin: 1rem 0;
        position: relative;
    }
    
    .spectrum-marker {
        width: 4px;
        height: 50px;
        background: #333;
        position: absolute;
        top: -5px;
        border-radius: 2px;
    }
    
    .module-contribution {
        display: flex;
        align-items: center;
        gap: 1rem;
        margin: 0.5rem 0;
        padding: 0.5rem;
        border-radius: 8px;
        transition: all 0.2s ease;
    }
    
    .module-contribution:hover {
        background: #f5f5f5;
    }
    
    .contribution-bar {
        height: 8px;
        border-radius: 4px;
        background: #667eea;
        transition: width 0.3s ease;
    }
    
    /* Metric cards */
    .metric-grid {
        display: grid;
        grid-template-columns: repeat(auto-fit, minmax(250px, 1fr));
        gap: 1rem;
        margin: 1rem 0;
    }
    
    .metric-card {
        background: white;
        border-radius: 16px;
        padding: 1.25rem;
        box-shadow: 0 5px 15px rgba(0,0,0,0.05);
        transition: all 0.3s ease;
        border: 1px solid rgba(0,0,0,0.05);
    }
    
    .metric-card:hover {
        transform: translateY(-2px);
        box-shadow: 0 10px 25px rgba(0,0,0,0.1);
    }
    
    .metric-title {
        font-size: 0.9rem;
        color: #666;
        margin-bottom: 0.5rem;
        text-transform: uppercase;
        letter-spacing: 0.5px;
    }
    
    .metric-value {
        font-size: 2rem;
        font-weight: 600;
        color: #333;
    }
    
    .metric-unit {
        font-size: 0.9rem;
        color: #999;
        margin-left: 0.25rem;
    }
    
    .metric-trend {
        font-size: 0.8rem;
        margin-top: 0.25rem;
    }
    
    .trend-up { color: #10b981; }
    .trend-down { color: #ef4444; }
    
    /* Example cards */
    .example-card {
        background: #f8f9fa;
        border-left: 4px solid #667eea;
        border-radius: 8px;
        padding: 1rem;
        margin: 0.5rem 0;
        font-size: 0.9rem;
    }
    
    .example-phrase {
        font-weight: 600;
        color: #667eea;
    }
    
    /* Tabs styling */
    .stTabs [data-baseweb="tab-list"] {
        gap: 0.5rem;
        background: white;
        padding: 0.5rem;
        border-radius: 12px;
        box-shadow: 0 2px 10px rgba(0,0,0,0.05);
    }
    
    .stTabs [data-baseweb="tab"] {
        border-radius: 8px;
        padding: 0.5rem 1rem;
        font-weight: 500;
    }
    
    .stTabs [aria-selected="true"] {
        background: #667eea !important;
        color: white !important;
    }
    
    /* Expander styling */
    .streamlit-expanderHeader {
        background: white;
        border-radius: 8px;
        border: 1px solid #eaeaea;
        font-weight: 500;
    }
    
    /* Hide default streamlit elements */
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    .stDeployButton {display:none;}
    
    /* Custom scrollbar */
    ::-webkit-scrollbar {
        width: 8px;
        height: 8px;
    }
    
    ::-webkit-scrollbar-track {
        background: #f1f1f1;
        border-radius: 10px;
    }
    
    ::-webkit-scrollbar-thumb {
        background: #c1c1c1;
        border-radius: 10px;
    }
    
    ::-webkit-scrollbar-thumb:hover {
        background: #a8a8a8;
    }
</style>
""", unsafe_allow_html=True)

# ============================================================================
# Classes and functions for analysis
# ============================================================================

class ReferenceCutoff:
    """Cut text before References section"""
    
    @staticmethod
    def cut_at_references(text: str) -> str:
        """Cut text at the beginning of references section"""
        if not text:
            return text
        
        # Reference markers
        reference_markers = [
            r'\nreferences?\s*\n',
            r'\nbibliography\s*\n',
            r'\nliterature cited\s*\n',
            r'\nworks cited\s*\n',
            r'\ncited literature\s*\n',
            r'\nreference list\s*\n',
            r'\nлитература\s*\n',
            r'\nсписок литературы\s*\n',
            r'\nбиблиография\s*\n',
            r'\nreferences and notes\s*\n',
            r'\nreferences & notes\s*\n'
        ]
        
        # Find first marker
        text_lower = text.lower()
        cutoff_pos = len(text)
        
        for marker in reference_markers:
            # Search case insensitive
            match = re.search(marker, text_lower)
            if match:
                # Find position in original text
                pos = text_lower.find(marker.strip())
                if 0 < pos < cutoff_pos:
                    cutoff_pos = pos
        
        # If marker found, cut
        if cutoff_pos < len(text):
            return text[:cutoff_pos].strip()
        
        return text

class UnicodeArtifactDetector:
    """Unicode artifact detector (level 1)"""
    
    def __init__(self):
        # Superscript and subscript characters
        self.sup_sub_chars = set([
            '⁰', '¹', '²', '³', '⁴', '⁵', '⁶', '⁷', '⁸', '⁹',  # superscript digits
            '₀', '₁', '₂', '₃', '₄', '₅', '₆', '₇', '₈', '₉',  # subscript digits
            '⁺', '⁻', '⁼', '⁽', '⁾',  # superscript signs
            '₊', '₋', '₌', '₍', '₎',  # subscript signs
            'ª', 'º',  # feminine/masculine ordinal
        ])
        
        # Fullwidth digits
        self.fullwidth_digits = set([chr(i) for i in range(0xFF10, 0xFF1A)])
        
        # Homoglyphs - characters that look like Latin/Cyrillic
        self.homoglyphs = {
            'а': 'a', 'е': 'e', 'о': 'o', 'р': 'p', 'с': 'c', 'у': 'y', 'х': 'x',
            'А': 'A', 'В': 'B', 'Е': 'E', 'К': 'K', 'М': 'M', 'Н': 'H',
            'О': 'O', 'Р': 'P', 'С': 'C', 'Т': 'T', 'Х': 'X',
        }
        
        # Greek characters (common in scientific papers)
        self.greek_chars = set([
            'α', 'β', 'γ', 'δ', 'ε', 'ζ', 'η', 'θ', 'ι', 'κ', 'λ', 'μ',
            'ν', 'ξ', 'ο', 'π', 'ρ', 'σ', 'τ', 'υ', 'φ', 'χ', 'ψ', 'ω',
            'Α', 'Β', 'Γ', 'Δ', 'Ε', 'Ζ', 'Η', 'Θ', 'Ι', 'Κ', 'Λ', 'Μ',
            'Ν', 'Ξ', 'Ο', 'Π', 'Ρ', 'Σ', 'Τ', 'Υ', 'Φ', 'Χ', 'Ψ', 'Ω'
        ])
        
        # Cyrillic
        self.cyrillic_chars = set([
            'а', 'б', 'в', 'г', 'д', 'е', 'ё', 'ж', 'з', 'и', 'й', 'к',
            'л', 'м', 'н', 'о', 'п', 'р', 'с', 'т', 'у', 'ф', 'х', 'ц',
            'ч', 'ш', 'щ', 'ъ', 'ы', 'ь', 'э', 'ю', 'я',
            'А', 'Б', 'В', 'Г', 'Д', 'Е', 'Ё', 'Ж', 'З', 'И', 'Й', 'К',
            'Л', 'М', 'Н', 'О', 'П', 'Р', 'С', 'Т', 'У', 'Ф', 'Х', 'Ц',
            'Ч', 'Ш', 'Щ', 'Ъ', 'Ы', 'Ь', 'Э', 'Ю', 'Я'
        ])
        
        # Arabic characters
        self.arabic_chars = set([
            'ا', 'ب', 'ت', 'ث', 'ج', 'ح', 'خ', 'د', 'ذ', 'ر', 'ز', 'س',
            'ش', 'ص', 'ض', 'ط', 'ظ', 'ع', 'غ', 'ف', 'ق', 'ك', 'ل', 'م',
            'ن', 'ه', 'و', 'ي', 'ء', 'آ', 'أ', 'ؤ', 'إ', 'ئ'
        ])
        
        # All suspicious characters
        self.all_suspicious = self.sup_sub_chars.union(self.fullwidth_digits)
    
    def analyze(self, text: str) -> Dict:
        """Analyze text for Unicode artifacts and non-Latin characters"""
        results = {
            'sup_sub_count': 0,
            'fullwidth_count': 0,
            'homoglyph_count': 0,
            'greek_count': 0,
            'cyrillic_count': 0,
            'arabic_count': 0,
            'non_latin_total': 0,
            'non_latin_per_1000': 0,
            'density_per_10k': 0,
            'suspicious_chunks': [],
            'all_suspicious_chunks': [],  # All found examples
            'risk_level': 'none',
            'risk_score': 0,
            'confidence': 0,
            # Additional statistics
            'statistics': {
                'mean_density': 0,
                'median_density': 0,
                'max_density': 0,
                'distribution': []
            }
        }
        
        total_chars = len(text)
        if total_chars == 0:
            return results
        
        # Count words for normalization
        words = text.split()
        total_words = len(words)
        
        # Go through text (full analysis, no limit)
        for i, char in enumerate(text):
            # Check superscript/subscript
            if char in self.sup_sub_chars:
                results['sup_sub_count'] += 1
                context = text[max(0, i-30):min(len(text), i+30)]
                chunk = {
                    'char': char,
                    'context': context,
                    'type': 'superscript/subscript'
                }
                results['suspicious_chunks'].append(chunk)
                results['all_suspicious_chunks'].append(chunk)
            
            # Check fullwidth digits
            elif char in self.fullwidth_digits:
                results['fullwidth_count'] += 1
                context = text[max(0, i-30):min(len(text), i+30)]
                chunk = {
                    'char': char,
                    'context': context,
                    'type': 'fullwidth digit'
                }
                results['suspicious_chunks'].append(chunk)
                results['all_suspicious_chunks'].append(chunk)
            
            # Check homoglyphs
            elif char in self.homoglyphs and i > 0 and i < len(text)-1:
                surrounding = text[max(0, i-3):i] + text[i+1:min(len(text), i+4)]
                if all(ord(c) < 128 or c.isspace() or c in '.,;:!?' for c in surrounding):
                    results['homoglyph_count'] += 1
            
            # Count non-Latin characters
            if char in self.greek_chars:
                results['greek_count'] += 1
                results['non_latin_total'] += 1
            elif char in self.cyrillic_chars:
                results['cyrillic_count'] += 1
                results['non_latin_total'] += 1
            elif char in self.arabic_chars:
                results['arabic_count'] += 1
                results['non_latin_total'] += 1
        
        # Density per 10,000 characters
        total_suspicious = results['sup_sub_count'] + results['fullwidth_count']
        results['density_per_10k'] = (total_suspicious * 10000) / total_chars if total_chars > 0 else 0
        
        # Non-Latin characters per 1000 words
        if total_words > 0:
            results['non_latin_per_1000'] = (results['non_latin_total'] * 1000) / total_words
        
        # Distribution statistics (split text into 1000 character segments)
        segment_size = 1000
        segments = [text[i:i+segment_size] for i in range(0, len(text), segment_size)]
        densities = []
        
        for segment in segments:
            seg_suspicious = sum(1 for c in segment if c in self.all_suspicious)
            seg_density = (seg_suspicious * 10000) / len(segment) if len(segment) > 0 else 0
            densities.append(seg_density)
        
        if densities:
            results['statistics']['distribution'] = densities
            results['statistics']['mean_density'] = float(np.mean(densities))
            results['statistics']['median_density'] = float(np.median(densities))
            results['statistics']['max_density'] = float(np.max(densities))
        
        # Risk assessment - УВЕЛИЧЕННЫЕ ЗНАЧЕНИЯ
        risk_score = 0
        confidence = 0.5  # base confidence
        
        # Consider artifact density - УВЕЛИЧИВАЕМ ВЕС
        if results['density_per_10k'] > 8:
            risk_score += 5                      # было 3, стало 5
            confidence = 0.95
        elif results['density_per_10k'] > 3:
            risk_score += 4                      # было 3, стало 4
            confidence = 0.85
        elif results['density_per_10k'] > 0:
            risk_score += 3                      # было 2, стало 3
            confidence = 0.75
        
        # Consider homoglyphs (more dangerous)
        if results['homoglyph_count'] > 5:
            risk_score += 4                      # было 3, стало 4
            confidence = min(confidence + 0.2, 1.0)
        elif results['homoglyph_count'] > 0:
            risk_score += 3                      # было 2, стало 3
            confidence = min(confidence + 0.15, 1.0)
        
        # Too many non-Latin characters (except Greek in scientific) - suspicious
        if results['non_latin_total'] > 100 and results['greek_count'] < results['non_latin_total'] * 0.8:
            risk_score += 3                      # было 2, стало 3
        
        # Если очень много артефактов - максимум риска
        if results['density_per_10k'] > 15 or results['sup_sub_count'] > 20:
            risk_score = 6                       # было 3, стало 6 (максимум)
            confidence = 1.0
        
        # Ограничиваем максимальный risk_score
        risk_score = min(risk_score, 6)           # добавляем ограничение сверху
        
        results['risk_score'] = risk_score
        results['confidence'] = confidence
        
        # Convert to levels for backward compatibility
        if risk_score >= 5:
            results['risk_level'] = 'critical'
        elif risk_score >= 4:
            results['risk_level'] = 'high'
        elif risk_score >= 2:
            results['risk_level'] = 'medium'
        elif risk_score > 0:
            results['risk_level'] = 'low'
        
        return results

class DashAnalyzer:
    """Multiple long dash analysis with paired dash detection"""
    
    def __init__(self):
        self.em_dash = '—'  # U+2014
        self.en_dash = '–'   # U+2013 (also common)
    
    def analyze(self, sentences: List[str]) -> Dict:
        """Analyze sentences for multiple dashes"""
        results = {
            'total_sentences': len(sentences),
            'sentences_with_multiple_dashes': [],
            'heavy_sentences': [],
            'all_dash_sentences': [],
            'double_dash_sentences': [],      # NEW: sentences with exactly two dashes
            'paired_dash_structures': [],     # NEW: detailed analysis of paired dashes
            'percentage_heavy': 0,
            'examples': [],
            'risk_level': 'none',
            'risk_score': 0,
            'confidence': 0,
            'statistics': {
                'mean_dashes_per_sentence': 0,
                'median_dashes_per_sentence': 0,
                'max_dashes_in_sentence': 0,
                'distribution': []
            }
        }
        
        if not sentences:
            return results
        
        dash_counts = []
        
        for sent_idx, sent in enumerate(sentences):
            if not sent or len(sent.strip()) == 0:
                continue
            
            # Count both em-dash and en-dash
            dash_count = sent.count(self.em_dash) + sent.count(self.en_dash)
            dash_counts.append(dash_count)
            word_count = len(sent.split())
            
            # Find positions of dashes
            dash_positions = []
            for i, char in enumerate(sent):
                if char in (self.em_dash, self.en_dash):
                    dash_positions.append(i)
            
            # Determine if heavy
            is_heavy = False
            if dash_count >= 3:
                is_heavy = True
            elif dash_count >= 2 and word_count > 90:
                is_heavy = True
            
            sent_data = {
                'sentence_idx': sent_idx,
                'sentence': sent,
                'dash_count': dash_count,
                'word_count': word_count,
                'is_heavy': is_heavy,
                'dash_positions': dash_positions
            }
            
            # Save all sentences with dashes
            if dash_count > 0:
                results['all_dash_sentences'].append(sent_data)
            
            # Special handling for sentences with exactly two dashes (paired structure)
            if dash_count == 2:
                results['double_dash_sentences'].append(sent_data)
                
                # Analyze paired dash structure
                pos1, pos2 = dash_positions[0], dash_positions[1]
                
                # Extract the text between dashes
                between_dashes = sent[pos1+1:pos2].strip()
                before_first = sent[:pos1].strip()
                after_second = sent[pos2+1:].strip()
                
                paired_data = {
                    'sentence_idx': sent_idx,
                    'full_sentence': sent,
                    'before': before_first[:100],
                    'between': between_dashes,
                    'between_word_count': len(between_dashes.split()),
                    'after': after_second[:100],
                    'structure_type': self._classify_paired_structure(before_first, between_dashes, after_second)
                }
                results['paired_dash_structures'].append(paired_data)
            
            if dash_count >= 2:
                results['sentences_with_multiple_dashes'].append(sent_data)
            
            if is_heavy:
                results['heavy_sentences'].append(sent_data)
                if len(results['examples']) < 20:
                    results['examples'].append(sent[:200])
        
        # Calculate statistics
        if sentences:
            results['percentage_heavy'] = (len(results['heavy_sentences']) / len(sentences)) * 100
            
            if dash_counts:
                results['statistics']['mean_dashes_per_sentence'] = float(np.mean(dash_counts))
                results['statistics']['median_dashes_per_sentence'] = float(np.median(dash_counts))
                results['statistics']['max_dashes_in_sentence'] = float(np.max(dash_counts))
                results['statistics']['distribution'] = dash_counts[:200]
        
        # Enhanced risk assessment considering paired dashes
        risk_score = 0
        confidence = 0.5
        
        # Paired dashes are a strong human indicator (complex syntax)
        paired_count = len(results['paired_dash_structures'])
        if paired_count > 5:
            risk_score = max(0, risk_score - 2)  # Reduce risk (human-like)
            confidence = min(confidence + 0.2, 1.0)
        elif paired_count > 2:
            risk_score = max(0, risk_score - 1)
            confidence = min(confidence + 0.1, 1.0)
        
        # Heavy dash usage is suspicious
        if results['percentage_heavy'] > 5:
            risk_score += 4
            confidence = 0.85
        elif results['percentage_heavy'] > 2:
            risk_score += 3
            confidence = 0.7
        elif results['percentage_heavy'] > 0:
            risk_score += 2
            confidence = 0.6
        
        # Many sentences with multiple dashes
        if len(results['sentences_with_multiple_dashes']) > len(sentences) * 0.1:
            risk_score = max(risk_score, 3)
            confidence = min(confidence + 0.1, 1.0)
        
        results['risk_score'] = min(risk_score, 6)  # Cap at 6
        results['confidence'] = confidence
        
        if risk_score >= 4:
            results['risk_level'] = 'critical'
        elif risk_score >= 3:
            results['risk_level'] = 'high'
        elif risk_score >= 2:
            results['risk_level'] = 'medium'
        elif risk_score >= 1:
            results['risk_level'] = 'low'
        
        return results
    
    def _classify_paired_structure(self, before: str, between: str, after: str) -> str:
        """
        Classify the structure of a paired dash construction.
        Common types in academic writing:
        - appositive: explanatory phrase between dashes
        - parenthetical: additional information
        - list_intro: before dash is list introduction
        - elaboration: after dash elaborates on between
        """
        between_words = len(between.split())
        
        # Appositive: between is a short explanation
        if between_words <= 5 and (',' in between or 'or' in between):
            return 'appositive'
        
        # Parenthetical: between contains additional but non-essential info
        if '(' in between or ')' in between or ';' in between:
            return 'parenthetical'
        
        # List introduction: before ends with colon or contains list indicator
        if before.strip().endswith(':') or 'following' in before.lower() or 'several' in before.lower():
            return 'list_introduction'
        
        # Elaboration: after is longer and explains between
        if len(after.split()) > between_words * 2:
            return 'elaboration'
        
        # Default
        return 'general'

class AIPhraseDetector:
    """Detector of characteristic AI phrases and clichés (level 1)"""
    
    def __init__(self):
        # Main AI phrases (updated 2025-2026)
        self.ai_phrases = [
            # "Fashionable" words 2024-2025
            'delve into', 'testament to', 'pivotal role', 'sheds light',
            'in the tapestry', 'in the realm', 'underscores', 'harnesses',
            'ever-evolving landscape', 'nuanced understanding', 'robust framework',
            'holistic approach', 'paradigm shift', 'cutting-edge',
            
            # New markers 2025-2026
            'crucial', 'pivotal', 'paramount', 'underscores', 'sheds light',
            'highlights the importance', 'testament to', 'integral', 'realm',
            'landscape', 'ever-evolving', 'tapestry', 'harness', 'delve',
            'delves', 'delving', 'intricate', 'meticulously', 'nuanced',
            'robust', 'unveiling', 'findings', 'revealed', 'demonstrated',
            'underscore', 'elucidate', 'illuminate', 'data-driven',
            'paves the way', 'leverage', 'leverages', 'leveraging',
            
            # Additional AI phrases
            'pathway', 'pathways',
            'signaling', 'signals',
            'Collectively',
            'manifest',
            'paradigm',
            
            # Stable connectives
            'it is worth noting', 'it is important to note',
            'it should be noted', 'as we delve deeper',
            'in the context of', 'with respect to', 'in terms of',
            'it is crucial to', 'it is paramount to',
            
            # Redundant transitions
            'moreover', 'furthermore', 'in addition', 'consequently',
            'therefore', 'thus', 'hence', 'nonetheless', 'nevertheless',
            'accordingly', 'as a result', 'for this reason',
            
            # Certainty amplifiers
            'significantly', 'substantially', 'dramatically',
            'remarkably', 'notably', 'strikingly', 'profoundly'
        ]
        
        # Metadiscourse markers (contrastive connectives)
        self.metadiscourse_markers = [
            'however', 'nevertheless', 'nonetheless', 'yet',
            'although', 'even though', 'despite', 'in spite of',
            'conversely', 'in contrast', 'on the contrary',
            'on the one hand', 'on the other hand'
        ]
        
        self.transition_threshold = 12
    
    def analyze(self, text: str, sentences: List[str]) -> Dict:
        """Analyze text for AI phrases and clichés"""
        results = {
            'ai_phrase_count': 0,
            'ai_phrases_found': [],
            'all_phrase_occurrences': [],  # Все найденные вхождения с контекстом
            'top_phrases': [],
            'metadiscourse_count': 0,
            'metadiscourse_markers': [],
            'transition_score': 0,
            'risk_level': 'none',
            'risk_score': 0,
            'confidence': 0,
            'statistics': {
                'mean_phrases_per_sentence': 0,
                'median_phrases_per_sentence': 0,
                'max_phrases_in_sentence': 0,
                'distribution': []
            }
        }
        
        if not text or not sentences:
            return results
        
        text_lower = text.lower()
        total_sentences = len(sentences)
        
        # Search for AI phrases
        phrase_counts = Counter()
        phrase_positions = []
        
        for phrase in self.ai_phrases:
            # Find all occurrences with context
            matches = list(re.finditer(re.escape(phrase), text_lower, re.IGNORECASE))
            if matches:
                phrase_counts[phrase] += len(matches)
                
                for match in matches[:20]:  # Limit per phrase for performance
                    start_pos = match.start()
                    # Get context (50 chars before and after)
                    context_start = max(0, start_pos - 100)
                    context_end = min(len(text), start_pos + len(phrase) + 100)
                    context = text[context_start:context_end].strip()
                    
                    occurrence = {
                        'phrase': phrase,
                        'context': context,
                        'position': start_pos
                    }
                    results['all_phrase_occurrences'].append(occurrence)
        
        # Search for metadiscourse markers
        for marker in self.metadiscourse_markers:
            count = text_lower.count(marker)
            if count > 0:
                results['metadiscourse_count'] += count
                results['metadiscourse_markers'].append({'marker': marker, 'count': count})
        
        # Calculate statistics
        results['ai_phrase_count'] = len(results['all_phrase_occurrences'])
        
        # Get top phrases
        if phrase_counts:
            results['top_phrases'] = phrase_counts.most_common(20)
        
        # Distribution by sentence
        phrases_per_sentence = []
        for sent in sentences[:200]:
            sent_lower = sent.lower()
            sent_count = sum(1 for phrase in self.ai_phrases if phrase in sent_lower)
            phrases_per_sentence.append(sent_count)
        
        if phrases_per_sentence:
            results['statistics']['mean_phrases_per_sentence'] = float(np.mean(phrases_per_sentence))
            results['statistics']['median_phrases_per_sentence'] = float(np.median(phrases_per_sentence))
            results['statistics']['max_phrases_in_sentence'] = float(np.max(phrases_per_sentence))
            results['statistics']['distribution'] = phrases_per_sentence[:100]
        
        # Transition score (ratio of metadiscourse to total phrases)
        if results['ai_phrase_count'] > 0:
            results['transition_score'] = results['metadiscourse_count'] / results['ai_phrase_count']
        
        # Risk assessment
        risk_score = 0
        confidence = 0.5
        
        # More AI phrases = higher risk
        if results['ai_phrase_count'] > 50:
            risk_score = 6
            confidence = 0.95
        elif results['ai_phrase_count'] > 25:
            risk_score = 5
            confidence = 0.9
        elif results['ai_phrase_count'] > 10:
            risk_score = 4
            confidence = 0.85
        elif results['ai_phrase_count'] > 5:
            risk_score = 3
            confidence = 0.7
        elif results['ai_phrase_count'] > 2:
            risk_score = 2
            confidence = 0.6
        elif results['ai_phrase_count'] > 0:
            risk_score = 1
            confidence = 0.5
        
        # Adjust for high transition score (more transitions = more human-like)
        if results['transition_score'] > 0.5:
            risk_score = max(0, risk_score - 1)
            confidence = min(confidence + 0.1, 1.0)
        
        results['risk_score'] = risk_score
        results['confidence'] = confidence
        
        if risk_score >= 5:
            results['risk_level'] = 'critical'
        elif risk_score >= 4:
            results['risk_level'] = 'high'
        elif risk_score >= 2:
            results['risk_level'] = 'medium'
        elif risk_score >= 1:
            results['risk_level'] = 'low'
        
        return results
    
    # Split into paragraphs (improved)
    paragraphs = self.split_paragraphs(text)
    
    # ===== IMPROVED GERUND DETECTION (pre-sentence analysis) =====
    gerund_pattern = r'\b([a-zA-Z]+ing)\b'
    gerund_positions = []
    
    for match in re.finditer(gerund_pattern, text, re.IGNORECASE):
        gerund_positions.append(match.end())
    
    # Analyze what follows each gerund
    for pos in gerund_positions:
        context = text[pos:pos+40]
        
        if re.match(r'\s+the\b', context, re.IGNORECASE):
            results['gerund_the_count'] += 1
            # Store example
            gerund_match = re.search(r'\b(\w+ing)\b', text[max(0, pos-20):pos])
            if gerund_match and len(results['gerund_contexts']) < 30:
                after_text = context[context.find('the')+4:context.find('the')+40].strip()
                results['gerund_contexts'].append({
                    'type': 'the',
                    'phrase': f"{gerund_match.group(1)} the {after_text}"[:100]
                })
        
        elif re.match(r'\s+of\b', context, re.IGNORECASE):
            results['gerund_of_count'] += 1
            gerund_match = re.search(r'\b(\w+ing)\b', text[max(0, pos-20):pos])
            if gerund_match and len(results['gerund_contexts']) < 30:
                after_text = context[context.find('of')+3:context.find('of')+40].strip()
                results['gerund_contexts'].append({
                    'type': 'of',
                    'phrase': f"{gerund_match.group(1)} of {after_text}"[:100]
                })
        
        elif re.match(r'\s+a\b', context, re.IGNORECASE):
            results['gerund_a_count'] += 1
        
        elif re.match(r'\s+an\b', context, re.IGNORECASE):
            results['gerund_an_count'] += 1
    
    # ===== SENTENCE-LEVEL STATISTICS =====
    for sentence in sentences:
        if not sentence or len(sentence.strip()) < 2:
            continue
        
        words = sentence.split()
        sent_length = len(words)
        results['sentence_lengths'].append(sent_length)
        
        # Commas
        comma_count = sentence.count(',')
        results['commas_per_sentence'].append(comma_count)
        
        # Apostrophes
        apostrophe_count = len(re.findall(r"'", sentence))
        results['apostrophes_per_sentence'].append(apostrophe_count)
        
        # -ly adverbs
        ly_adverbs = len(re.findall(r'\b\w+ly\b', sentence, re.IGNORECASE))
        results['ly_adverbs_per_sentence'].append(ly_adverbs)
        
        # Gerund patterns per sentence
        gerund_the_sent = len(re.findall(r'\b\w+ing\s+the\b', sentence, re.IGNORECASE))
        gerund_of_sent = len(re.findall(r'\b\w+ing\s+of\b', sentence, re.IGNORECASE))
        results['gerund_the_per_sentence'].append(gerund_the_sent)
        results['gerund_of_per_sentence'].append(gerund_of_sent)
        
        # Indefinite articles
        articles = len(re.findall(r'\b(a|an|A|An)\b', sentence))
        results['indefinite_articles_per_sentence'].append(articles)
        
        # Figure/Table/Supplementary mentions (with full sentence context)
        for pattern in self.figure_patterns:
            matches = re.findall(pattern, sentence, re.IGNORECASE)
            for match in matches:
                results['figure_mentions'].append({
                    'sentence': sentence.strip(),
                    'match': match
                })
        
        for pattern in self.table_patterns:
            matches = re.findall(pattern, sentence, re.IGNORECASE)
            for match in matches:
                results['table_mentions'].append({
                    'sentence': sentence.strip(),
                    'match': match
                })
        
        for pattern in self.supplementary_patterns:
            matches = re.findall(pattern, sentence, re.IGNORECASE)
            for match in matches:
                results['supplementary_mentions'].append({
                    'sentence': sentence.strip(),
                    'match': match
                })
    
    # ===== PARAGRAPH-LEVEL STATISTICS =====
    for paragraph in paragraphs:
        if not paragraph:
            continue
        
        para_words = paragraph.split()
        para_length = len(para_words)
        results['paragraph_lengths'].append(para_length)
        
        results['commas_per_paragraph'].append(paragraph.count(','))
        results['apostrophes_per_paragraph'].append(len(re.findall(r"'", paragraph)))
        results['ly_adverbs_per_paragraph'].append(len(re.findall(r'\b\w+ly\b', paragraph, re.IGNORECASE)))
        results['gerund_the_per_paragraph'].append(len(re.findall(r'\b\w+ing\s+the\b', paragraph, re.IGNORECASE)))
        results['gerund_of_per_paragraph'].append(len(re.findall(r'\b\w+ing\s+of\b', paragraph, re.IGNORECASE)))
        results['indefinite_articles_per_paragraph'].append(len(re.findall(r'\b(a|an|A|An)\b', paragraph)))
    
    # ===== CALCULATE STATISTICS =====
    def calculate_stats(values):
        if not values:
            return {'min': 0, 'max': 0, 'mean': 0, 'median': 0, 'total': 0}
        return {
            'min': float(np.min(values)),
            'max': float(np.max(values)),
            'mean': float(np.mean(values)),
            'median': float(np.median(values)),
            'total': sum(values),
            'count': len(values)
        }
    
    # Sentence-level stats
    results['statistics']['sentence_length'] = calculate_stats(results['sentence_lengths'])
    results['statistics']['commas_per_sentence'] = calculate_stats(results['commas_per_sentence'])
    results['statistics']['apostrophes_per_sentence'] = calculate_stats(results['apostrophes_per_sentence'])
    results['statistics']['ly_adverbs_per_sentence'] = calculate_stats(results['ly_adverbs_per_sentence'])
    results['statistics']['gerund_the_per_sentence'] = calculate_stats(results['gerund_the_per_sentence'])
    results['statistics']['gerund_of_per_sentence'] = calculate_stats(results['gerund_of_per_sentence'])
    results['statistics']['indefinite_articles_per_sentence'] = calculate_stats(results['indefinite_articles_per_sentence'])
    
    # Paragraph-level stats
    results['statistics']['paragraph_length'] = calculate_stats(results['paragraph_lengths'])
    results['statistics']['commas_per_paragraph'] = calculate_stats(results['commas_per_paragraph'])
    results['statistics']['apostrophes_per_paragraph'] = calculate_stats(results['apostrophes_per_paragraph'])
    results['statistics']['ly_adverbs_per_paragraph'] = calculate_stats(results['ly_adverbs_per_paragraph'])
    results['statistics']['gerund_the_per_paragraph'] = calculate_stats(results['gerund_the_per_paragraph'])
    results['statistics']['gerund_of_per_paragraph'] = calculate_stats(results['gerund_of_per_paragraph'])
    results['statistics']['indefinite_articles_per_paragraph'] = calculate_stats(results['indefinite_articles_per_paragraph'])
    
    # Additional counts
    results['figure_count'] = len(results['figure_mentions'])
    results['table_count'] = len(results['table_mentions'])
    results['supplementary_count'] = len(results['supplementary_mentions'])
    
    # Add warning flag if paragraph detection failed
    if len(paragraphs) <= 1 and len(text.split()) > 500:
        results['paragraph_warning'] = True
    
    return results

class TorturedPhraseDetector:
    """Detector for tortured phrases - distorted academic terminology"""
    
    def __init__(self):
        # Dictionary of correct terms -> list of tortured variations
        self.tortured_phrases = {
            # Chemistry
            'amino acid': ['amino corrosive'],
            'ascorbic acid': ['ascorbic corrosive'],
            'hydrochloric acid': ['hydrochloric corrosive'],
            'nitric acid': ['nitric corrosive'],
            'sulfuric acid': ['sulfuric corrosive'],
            'acetic acid': ['acetic corrosive'],
            'citric acid': ['citric corrosive'],
            'phosphoric acid': ['phosphoric corrosive'],
            'formic acid': ['formic corrosive'],
            'lactic acid': ['lactic corrosive'],
            'nucleic acid': ['nucleic corrosive'],
            'aqueous solution': ['watery arrangement', 'watery solution'],
            'ph neutral': ['impartial ph', 'neutral ph', 'ph impartial'],
            
            # Physics/Materials
            'surface area': ['surface region', 'surface territory'],
            'thermal conductivity': ['warm conductivity', 'full conductivity', 'common conductivity', 'general conductivity'],
            'heat transfer': ['warmth move', 'heat move'],
            'heat exchanger': ['warmth exchanger'],
            'thermal conductivity': ['warm conductivity'],
            'turbulent flow': ['turbulant flow', 'turbulent stream'],
            'boundary layer': ['limit layer'],
            'computational fluid dynamics': ['computational liquid elements'],
            'energy consumption': ['vitality utilization', 'energy utilization'],
            'energy efficiency': ['vitality effective', 'energy effective'],
            'energy saving': ['vitality sparing', 'energy sparing'],
            'residual energy': ['remaining vitality', 'leftover vitality'],
            'magnetic resonance': ['attractive reverberation'],
            
            # Medical/Imaging
            'computed tomography': ['processed tomography', 'figured tomography', 'ct'],
            'magnetic resonance imaging': ['attractive reverberation imaging', 'mri'],
            
            # Environmental
            'greenhouse gas': ['ozone harming substance', 'ozone depleting substance'],
            'greenhouse gas emissions': ['ozone harming substance discharge', 'ozone depleting substance emissions'],
            'global warming': ['an earth-wide temperature boost', 'earth wide temperature boost'],
            'solar energy': ['sun oriented energy', 'sunlight-based energy'],
            'alternative energy': ['elective energy'],
            'environmental degradation': ['ecological corruption', 'natural debasement'],
            'heavy metals': ['substantial metals'],
            
            # Nanotechnology
            'polymeric nanofiber': ['polymeric nanofiber'],
            'quantum dots': ['quantum dabs'],
            'drug delivery': ['medication conveyance'],
            'negatively charged': ['contrarily charged', 'negatively charged'],
            'transition metal': ['progress metal'],
            
            # AI/CS
            'artificial intelligence': ['counterfeit consciousness', 'artificial consciousness'],
            'deep neural network': ['profound neural organization', 'deep neural organization'],
            'workflow engine': ['work process motor'],
            'global parameters': ['worldwide parameters'],
            
            # Common tortured phrases
            'certification': ['sertification'],
            'methodology': ['methodologie', 'methodolgy'],
            'analysis': ['analysys', 'analisys'],
            'synthesis': ['synthethis', 'syntesis'],
            'characterization': ['characterisation', 'characterisation'],
            'significant': ['significent', 'signifcant'],
            'important': ['importent', 'importnant'],
            'demonstrate': ['demonstate', 'demostrate'],
        }
        
        # Flatten for easier searching
        self.all_tortured = []
        self.tortured_to_correct = {}
        for correct, tortured_list in self.tortured_phrases.items():
            for tortured in tortured_list:
                self.all_tortured.append(tortured)
                self.tortured_to_correct[tortured] = correct
    
    def analyze(self, text: str) -> Dict:
        """Analyze text for tortured phrases"""
        results = {
            'tortured_phrases_found': [],
            'all_occurrences': [],  # Все найденные вхождения с контекстом
            'unique_tortured_count': 0,
            'total_occurrences': 0,
            'tortured_per_1000_words': 0,
            'examples': [],
            'risk_score': 0,
            'confidence': 0,
            'risk_level': 'none',
            'statistics': {
                'distribution': {}
            }
        }
        
        if not text:
            return results
        
        text_lower = text.lower()
        words = text.split()
        total_words = len(words)
        
        # Search for each tortured phrase
        for tortured in self.all_tortured:
            # Use word boundary to avoid partial matches
            pattern = r'\b' + re.escape(tortured) + r'\b'
            matches = list(re.finditer(pattern, text_lower))
            
            if matches:
                correct_term = self.tortured_to_correct[tortured]
                
                # Сохраняем каждое вхождение
                for match in matches[:20]:  # Лимит на однотипные для производительности
                    start_pos = match.start()
                    # Получаем контекст (предложение или окружающий текст)
                    context_start = max(0, text_lower.rfind('.', 0, start_pos))
                    context_end = text_lower.find('.', start_pos + len(tortured))
                    if context_end == -1:
                        context_end = len(text)
                    context = text[context_start:context_end+1].strip()
                    
                    occurrence = {
                        'tortured': tortured,
                        'correct': correct_term,
                        'context': context[:300] if len(context) > 300 else context,
                        'position': start_pos
                    }
                    results['all_occurrences'].append(occurrence)
                    
                    if len(results['examples']) < 50:
                        results['examples'].append({
                            'tortured': tortured,
                            'correct': correct_term,
                            'context': context[:200]
                        })
                
                # Сохраняем статистику по фразе
                results['tortured_phrases_found'].append({
                    'tortured': tortured,
                    'correct': correct_term,
                    'count': len(matches)
                })
                
                results['total_occurrences'] += len(matches)
        
        results['unique_tortured_count'] = len(results['tortured_phrases_found'])
        
        # Normalize per 1000 words
        if total_words > 0:
            results['tortured_per_1000_words'] = (results['total_occurrences'] * 1000) / total_words
        
        # Risk assessment
        risk_score = 0
        confidence = 0.5
        
        if results['total_occurrences'] > 10:
            risk_score = 6  # Максимальный риск - явные признаки tortured phrases
            confidence = 1.0
        elif results['total_occurrences'] > 5:
            risk_score = 5
            confidence = 0.95
        elif results['total_occurrences'] > 2:
            risk_score = 4
            confidence = 0.85
        elif results['total_occurrences'] > 0:
            risk_score = 3
            confidence = 0.7
        
        results['risk_score'] = risk_score
        results['confidence'] = confidence
        
        if risk_score >= 5:
            results['risk_level'] = 'critical'
        elif risk_score >= 4:
            results['risk_level'] = 'high'
        elif risk_score >= 3:
            results['risk_level'] = 'medium'
        elif risk_score >= 1:
            results['risk_level'] = 'low'
        
        return results

class BurstinessAnalyzer:
    """Sentence length variability analysis (level 2) - updated with Yule's I"""
    
    def analyze(self, sentences: List[str]) -> Dict:
        """Analyze text burstiness using Yule's I"""
        results = {
            'sentence_lengths': [],
            'mean_length': 0,
            'std_length': 0,
            'cv': 0,
            'iqr': 0,
            'yules_i': 0,  # Yule's I characteristic
            'sic': 0,      # Syntactic Irregularity Coefficient
            'burstiness': 'unknown',
            'risk_level': 'none',
            'risk_score': 0,
            'confidence': 0,
            'statistics': {
                'median_length': 0,
                'max_length': 0,
                'min_length': 0,
                'percentile_25': 0,
                'percentile_75': 0,
                'distribution': []
            }
        }
        
        if len(sentences) < 5:
            results['error'] = 'Too few sentences for analysis'
            return results
        
        # Sentence length in words
        sent_lengths = [len(sent.split()) for sent in sentences if len(sent.strip()) > 0]
        if not sent_lengths:
            return results
            
        results['sentence_lengths'] = sent_lengths
        
        # Basic statistics
        results['mean_length'] = float(np.mean(sent_lengths))
        results['std_length'] = float(np.std(sent_lengths))
        
        # Coefficient of variation (CV)
        if results['mean_length'] > 0:
            results['cv'] = float(results['std_length'] / results['mean_length'])
        
        # Interquartile range (IQR)
        q75, q25 = np.percentile(sent_lengths, [75, 25])
        results['iqr'] = float(q75 - q25)
        
        # Yule's I (variance measure robust to text length)
        # I = (σ² / μ) * 10000
        variance = np.var(sent_lengths)
        if results['mean_length'] > 0:
            results['yules_i'] = float((variance / results['mean_length']) * 10000)
        
        # Additional statistics
        results['statistics']['median_length'] = float(np.median(sent_lengths))
        results['statistics']['max_length'] = float(np.max(sent_lengths))
        results['statistics']['min_length'] = float(np.min(sent_lengths))
        results['statistics']['percentile_25'] = float(q25)
        results['statistics']['percentile_75'] = float(q75)
        results['statistics']['distribution'] = sent_lengths[:200]
        
        # Syntactic Irregularity Coefficient (SIC) - simple metric
        # Ratio of max to min length with adjustment
        if results['statistics']['min_length'] > 0:
            results['sic'] = float(results['statistics']['max_length'] / results['statistics']['min_length'])
        
        # Probabilistic burstiness assessment based on Yule's I
        risk_score = 0
        confidence = 0.5
        
        # Normal Yule's I values for scientific texts: ~200-800
        if results['yules_i'] < 150:
            results['burstiness'] = 'very_low'
            risk_score = 3
            confidence = 0.9
        elif results['yules_i'] < 250:
            results['burstiness'] = 'low'
            risk_score = 2
            confidence = 0.7
        elif results['yules_i'] < 500:
            results['burstiness'] = 'normal'
            risk_score = 1
            confidence = 0.5
        else:
            results['burstiness'] = 'high'
            risk_score = 0
            confidence = 0.3
        
        results['risk_score'] = risk_score
        results['confidence'] = confidence
        
        if risk_score >= 3:
            results['risk_level'] = 'high'
        elif risk_score >= 2:
            results['risk_level'] = 'medium'
        elif risk_score >= 1:
            results['risk_level'] = 'low'
        
        return results


class RepetitivenessAnalyzer:
    """N-gram repetitiveness analysis (new module)"""
    
    def __init__(self):
        self.n_values = [3, 4]  # 3-grams and 4-grams
    
    def analyze(self, text: str, sentences: List[str]) -> Dict:
        """Analyze phrase repetitiveness in text"""
        results = {
            'ngram_repetition_scores': {},
            'unique_ngram_ratios': {},
            'repeated_phrases': [],
            'all_repetitions': [],  # All found repetitions
            'risk_score': 0,
            'confidence': 0,
            'risk_level': 'none',
            'statistics': {
                'mean_repetition_rate': 0,
                'max_repetition_rate': 0,
                'repetition_distribution': {}
            }
        }
        
        if not text or len(text.split()) < 50:
            return results
        
        words = text.lower().split()
        
        for n in self.n_values:
            if len(words) < n:
                continue
                
            # Collect all N-grams
            ngrams = []
            for i in range(len(words) - n + 1):
                ngram = ' '.join(words[i:i+n])
                ngrams.append(ngram)
            
            # Count frequencies
            ngram_counts = Counter(ngrams)
            total_ngrams = len(ngrams)
            unique_ngrams = len(ngram_counts)
            
            # Unique ratio (lower means more repetition)
            unique_ratio = unique_ngrams / total_ngrams if total_ngrams > 0 else 1.0
            
            # Collect repeated phrases (occur >= 3 times)
            repeated = [(ngram, count) for ngram, count in ngram_counts.items() if count >= 3]
            repeated.sort(key=lambda x: x[1], reverse=True)
            
            # Save results
            results['ngram_repetition_scores'][f'{n}gram'] = 1.0 - unique_ratio  # Closer to 1 means more repetition
            results['unique_ngram_ratios'][f'{n}gram'] = unique_ratio
            
            # Save top repetitions
            for ngram, count in repeated[:20]:
                results['all_repetitions'].append({
                    'ngram': ngram,
                    'count': count,
                    'type': f'{n}-gram'
                })
            
            results['statistics']['repetition_distribution'][f'{n}gram'] = [c for _, c in repeated[:50]]
        
        if results['ngram_repetition_scores']:
            results['statistics']['mean_repetition_rate'] = float(np.mean(list(results['ngram_repetition_scores'].values())))
            results['statistics']['max_repetition_rate'] = float(np.max(list(results['ngram_repetition_scores'].values())))
        
        # Risk assessment
        risk_score = 0
        confidence = 0.5
        
        # Check 3-grams (more sensitive to repetitions)
        if '3gram' in results['ngram_repetition_scores']:
            rep_rate = results['ngram_repetition_scores']['3gram']
            if rep_rate > 0.25:
                risk_score += 3
                confidence = min(confidence + 0.3, 1.0)
            elif rep_rate > 0.15:
                risk_score += 2
                confidence = min(confidence + 0.2, 1.0)
            elif rep_rate > 0.08:
                risk_score += 1
                confidence = min(confidence + 0.1, 1.0)
        
        # Check 4-grams
        if '4gram' in results['ngram_repetition_scores']:
            rep_rate = results['ngram_repetition_scores']['4gram']
            if rep_rate > 0.15:
                risk_score += 2
                confidence = min(confidence + 0.2, 1.0)
            elif rep_rate > 0.08:
                risk_score += 1
                confidence = min(confidence + 0.1, 1.0)
        
        results['risk_score'] = risk_score
        results['confidence'] = confidence
        
        if risk_score >= 4:
            results['risk_level'] = 'high'
        elif risk_score >= 2:
            results['risk_level'] = 'medium'
        elif risk_score >= 1:
            results['risk_level'] = 'low'
        
        return results


class LexicalDiversityAnalyzer:
    """Lexical diversity analysis (MTLD, MATTR, HD-D)"""
    
    def __init__(self):
        self.available = LEXICAL_DIVERSITY_AVAILABLE
    
    def analyze(self, text: str) -> Dict:
        """Analyze text lexical diversity"""
        results = {
            'ttr': 0,  # Type-Token Ratio (basic)
            'mtld': 0,  # Measure of Textual Lexical Diversity
            'mattr': 0,  # Moving Average Type-Token Ratio
            'hdd': 0,    # Hypergeometric Distribution Diversity
            'hapax_legomena': 0,  # Words that occur once
            'hapax_dislegomena': 0,  # Words that occur twice
            'hapax_ratio': 0,
            'risk_score': 0,
            'confidence': 0,
            'risk_level': 'none',
            'note': 'Lexical diversity analysis requires lexical-diversity library'
        }
        
        if not text or len(text.split()) < 50:
            return results
        
        words = text.lower().split()
        unique_words = set(words)
        
        # Basic TTR
        results['ttr'] = len(unique_words) / len(words) if words else 0
        
        # Count hapax legomena (words that occur once)
        word_counts = Counter(words)
        results['hapax_legomena'] = sum(1 for count in word_counts.values() if count == 1)
        results['hapax_dislegomena'] = sum(1 for count in word_counts.values() if count == 2)
        results['hapax_ratio'] = results['hapax_legomena'] / len(words) if words else 0
        
        # If lexical-diversity library is available, use advanced metrics
        if self.available:
            try:
                # Convert to lexical-diversity format
                text_for_lex = ' '.join(words)
                
                # MTLD
                results['mtld'] = lexdiv.mtld(text_for_lex)
                
                # MATTR (window size = 50)
                results['mattr'] = lexdiv.mattr(text_for_lex, window_size=50)
                
                # HD-D
                results['hdd'] = lexdiv.hdd(text_for_lex, sample_size=42)
                
                results['note'] = 'Full lexical diversity analysis'
            except Exception as e:
                results['note'] = f'Error in advanced lexical analysis: {str(e)}'
        
        # Risk assessment based on lexical diversity
        risk_score = 0
        confidence = 0.5
        
        # Low lexical diversity - AI indicator
        if results['mtld'] > 0:
            if results['mtld'] < 40:
                risk_score += 3
                confidence = min(confidence + 0.3, 1.0)
            elif results['mtld'] < 60:
                risk_score += 2
                confidence = min(confidence + 0.2, 1.0)
            elif results['mtld'] < 80:
                risk_score += 1
                confidence = min(confidence + 0.1, 1.0)
        else:
            # Fallback to TTR
            if results['ttr'] < 0.4:
                risk_score += 2
                confidence = min(confidence + 0.2, 1.0)
            elif results['ttr'] < 0.5:
                risk_score += 1
                confidence = min(confidence + 0.1, 1.0)
        
        # Low hapax legomena ratio - template indicator
        if results['hapax_ratio'] < 0.5:
            risk_score += 1
            confidence = min(confidence + 0.1, 1.0)
        
        results['risk_score'] = risk_score
        results['confidence'] = confidence
        
        if risk_score >= 4:
            results['risk_level'] = 'high'
        elif risk_score >= 2:
            results['risk_level'] = 'medium'
        elif risk_score >= 1:
            results['risk_level'] = 'low'
        
        return results


class LogProbAnalyzer:
    """Log-probability analysis through open models"""
    
    def __init__(self):
        self.available = TRANSFORMERS_AVAILABLE
        self.model = None
        self.tokenizer = None
        
        if self.available:
            try:
                # Use small Qwen2.5-7B model (reduced for memory)
                model_name = "Qwen/Qwen2.5-7B-Instruct"
                self.tokenizer = AutoTokenizer.from_pretrained(model_name, trust_remote_code=True)
                self.model = AutoModelForCausalLM.from_pretrained(
                    model_name,
                    torch_dtype=torch.float16,
                    device_map="auto",
                    trust_remote_code=True
                )
            except Exception as e:
                try:
                    # Fallback to lighter model
                    model_name = "microsoft/Phi-3.5-mini-instruct"
                    self.tokenizer = AutoTokenizer.from_pretrained(model_name, trust_remote_code=True)
                    self.model = AutoModelForCausalLM.from_pretrained(
                        model_name,
                        torch_dtype=torch.float16,
                        device_map="auto",
                        trust_remote_code=True
                    )
                except Exception as e2:
                    self.available = False
    
    def calculate_log_probs(self, text: str) -> Dict:
        """Calculate average log-probabilities for text"""
        if not self.available or len(text) < 100:
            return {}
        
        try:
            # Tokenize
            inputs = self.tokenizer(text, return_tensors="pt", truncation=True, max_length=512)
            
            # Calculate logits
            with torch.no_grad():
                outputs = self.model(**inputs)
                logits = outputs.logits
                
                # Calculate log probabilities
                log_probs = torch.log_softmax(logits, dim=-1)
                
                # Get probabilities for actual tokens
                token_log_probs = []
                for i in range(inputs['input_ids'].shape[1] - 1):
                    token_id = inputs['input_ids'][0, i+1]
                    token_log_prob = log_probs[0, i, token_id].item()
                    token_log_probs.append(token_log_prob)
                
                # Statistics
                results = {
                    'mean_log_prob': float(np.mean(token_log_probs)),
                    'median_log_prob': float(np.median(token_log_probs)),
                    'std_log_prob': float(np.std(token_log_probs)),
                    'min_log_prob': float(np.min(token_log_probs)),
                    'max_log_prob': float(np.max(token_log_probs))
                }
                
                return results
        except Exception as e:
            return {}
    
    def analyze(self, text: str) -> Dict:
        """Analyze text through log-probabilities"""
        results = {
            'log_prob_stats': {},
            'mean_log_prob': 0,
            'perplexity_estimate': 0,
            'risk_score': 0,
            'confidence': 0,
            'risk_level': 'none',
            'note': 'Log-probability analysis requires transformers'
        }
        
        if not self.available or len(text) < 100:
            results['note'] = 'Log-probability analysis unavailable or text too short'
            return results
        
        log_prob_data = self.calculate_log_probs(text)
        if log_prob_data:
            results['log_prob_stats'] = log_prob_data
            results['mean_log_prob'] = log_prob_data['mean_log_prob']
            
            # Approximate perplexity from log-prob
            results['perplexity_estimate'] = float(np.exp(-log_prob_data['mean_log_prob']))
            
            # Risk assessment
            risk_score = 0
            confidence = 0.5
            
            # Higher log-probabilities (less negative) - AI indicator
            if results['mean_log_prob'] > -2.0:
                risk_score = 3
                confidence = 0.9
            elif results['mean_log_prob'] > -3.0:
                risk_score = 2
                confidence = 0.7
            elif results['mean_log_prob'] > -4.0:
                risk_score = 1
                confidence = 0.6
            
            results['risk_score'] = risk_score
            results['confidence'] = confidence
            
            if risk_score >= 3:
                results['risk_level'] = 'high'
            elif risk_score >= 2:
                results['risk_level'] = 'medium'
            elif risk_score >= 1:
                results['risk_level'] = 'low'
        
        return results


class MLClassifier:
    """ML classifier based on BERT/RoBERTa (fine-tuned)"""
    
    def __init__(self):
        self.available = TRANSFORMERS_AVAILABLE
        self.model = None
        self.tokenizer = None
        
        if self.available:
            try:
                # Use pre-trained AI detector (if available)
                # In real project, this would be your fine-tuned model
                model_name = "roberta-base-openai-detector"  # Example
                self.tokenizer = AutoTokenizer.from_pretrained(model_name)
                self.model = AutoModelForSequenceClassification.from_pretrained(model_name)
            except:
                # If no ready model, use stub
                self.available = False
    
    def analyze(self, text: str) -> Dict:
        """Analyze text using ML model"""
        results = {
            'ml_score': 0,
            'ml_probability': 0,
            'ml_confidence': 0,
            'risk_level': 'none',
            'risk_score': 0,
            'confidence': 0,
            'note': 'ML classifier not available'
        }
        
        if not self.available or len(text) < 100:
            return results
        
        try:
            # Tokenize
            inputs = self.tokenizer(text, return_tensors="pt", truncation=True, max_length=512)
            
            # Prediction
            with torch.no_grad():
                outputs = self.model(**inputs)
                probs = torch.softmax(outputs.logits, dim=-1)
                
                # Assume binary classification: 0 - human, 1 - AI
                ai_prob = probs[0, 1].item()
                
                results['ml_score'] = ai_prob * 100
                results['ml_probability'] = ai_prob
                results['ml_confidence'] = max(probs[0]).item()
                
                # Risk assessment
                risk_score = 0
                confidence = results['ml_confidence']
                
                if ai_prob > 0.8:
                    risk_score = 3
                elif ai_prob > 0.6:
                    risk_score = 2
                elif ai_prob > 0.4:
                    risk_score = 1
                
                results['risk_score'] = risk_score
                results['confidence'] = confidence
                
                if risk_score >= 3:
                    results['risk_level'] = 'high'
                elif risk_score >= 2:
                    results['risk_level'] = 'medium'
                elif risk_score >= 1:
                    results['risk_level'] = 'low'
                
                results['note'] = 'ML classifier analysis'
                
        except Exception as e:
            results['note'] = f'Error in ML classification: {str(e)}'
        
        return results


class GrammarAnalyzer:
    """Grammar feature analysis (without spacy, simplified version)"""
    
    def __init__(self):
        self.nominalization_suffixes = ['tion', 'ment', 'ance', 'ence', 'ing', 'ity', 'ism', 'sis', 'ure', 'age']
        
        # Modal verbs for hedging
        self.modal_verbs = ['may', 'might', 'could', 'would', 'should', 'can']
        
        # Epistemic markers
        self.epistemic_markers = [
            'seem', 'appear', 'suggest', 'indicate', 'likely', 'unlikely',
            'possibly', 'probably', 'perhaps', 'maybe', 'potentially',
            'presumably', 'arguably', 'tentatively'
        ]
        
        # Certainty boosters
        self.certainty_boosters = [
            'crucial', 'pivotal', 'paramount', 'essential', 'vital',
            'undoubtedly', 'certainly', 'definitely', 'clearly', 'obviously',
            'demonstrates', 'proves', 'confirms', 'establishes'
        ]
    
    def analyze(self, text: str) -> Dict:
        """Simplified grammar analysis without spacy"""
        results = {
            'passive_indicators': 0,
            'nominalization_count': 0,
            'modal_count': 0,
            'epistemic_count': 0,
            'certainty_boosters_count': 0,
            'total_sentences': 0,
            'passive_percentage': 0,
            'nominalizations_per_1000': 0,
            'modals_per_1000': 0,
            'epistemic_per_1000': 0,
            'boosters_per_1000': 0,
            'examples': [],
            'risk_level': 'none',
            'risk_score': 0,
            'confidence': 0,
            'statistics': {
                'mean_passive_per_sentence': 0,
                'median_passive_per_sentence': 0,
                'max_passive_in_sentence': 0,
                'distribution': []
            }
        }
        
        if not text:
            return results
        
        # Simple sentence segmentation
        sentences = re.split(r'[.!?]+', text)
        valid_sentences = [s for s in sentences if len(s.strip()) > 10]
        results['total_sentences'] = len(valid_sentences)
        
        text_lower = text.lower()
        words = text_lower.split()
        total_words = len(words)
        
        # Search for passive indicators (was/were + ed/en)
        passive_pattern = r'\b(was|were|is|are|been|being)\s+(\w+ed|\w+en)\b'
        passive_matches = re.findall(passive_pattern, text_lower)
        results['passive_indicators'] = len(passive_matches)
        
        # Nominalization suffixes
        for word in words:
            for suffix in self.nominalization_suffixes:
                if word.endswith(suffix) and len(word) > len(suffix) + 2:
                    results['nominalization_count'] += 1
                    break
        
        # Modal verbs
        for modal in self.modal_verbs:
            results['modal_count'] += sum(1 for w in words if w == modal)
        
        # Epistemic markers
        for marker in self.epistemic_markers:
            if ' ' in marker:
                results['epistemic_count'] += text_lower.count(marker)
            else:
                results['epistemic_count'] += sum(1 for w in words if w == marker)
        
        # Certainty boosters
        for booster in self.certainty_boosters:
            results['certainty_boosters_count'] += sum(1 for w in words if w == booster)
        
        # Normalize per 1000 words
        if total_words > 0:
            results['nominalizations_per_1000'] = (results['nominalization_count'] * 1000) / total_words
            results['modals_per_1000'] = (results['modal_count'] * 1000) / total_words
            results['epistemic_per_1000'] = (results['epistemic_count'] * 1000) / total_words
            results['boosters_per_1000'] = (results['certainty_boosters_count'] * 1000) / total_words
        
        if results['total_sentences'] > 0:
            results['passive_percentage'] = (results['passive_indicators'] / results['total_sentences']) * 100
        
        # Statistics by sentence
        passive_per_sentence = []
        for sent in valid_sentences[:100]:
            sent_lower = sent.lower()
            sent_passive = len(re.findall(passive_pattern, sent_lower))
            passive_per_sentence.append(sent_passive)
        
        if passive_per_sentence:
            results['statistics']['mean_passive_per_sentence'] = float(np.mean(passive_per_sentence))
            results['statistics']['median_passive_per_sentence'] = float(np.median(passive_per_sentence))
            results['statistics']['max_passive_in_sentence'] = float(np.max(passive_per_sentence))
            results['statistics']['distribution'] = passive_per_sentence
        
        # Probabilistic risk assessment
        risk_score = 0
        confidence = 0.5
        
        # Passive (too many - suspicious for AI?)
        if results['passive_percentage'] > 40:
            risk_score += 2
            confidence = min(confidence + 0.2, 1.0)
        elif results['passive_percentage'] > 30:
            risk_score += 1
            confidence = min(confidence + 0.1, 1.0)
        
        # Nominalizations (too many - academic style, but could be AI)
        if results['nominalizations_per_1000'] > 25:
            risk_score += 2
            confidence = min(confidence + 0.2, 1.0)
        elif results['nominalizations_per_1000'] > 15:
            risk_score += 1
            confidence = min(confidence + 0.1, 1.0)
        
        # Modal verbs (too few - insufficient hedging)
        if results['modals_per_1000'] < 3 and total_words > 500:
            risk_score += 2
            confidence = min(confidence + 0.2, 1.0)
        elif results['modals_per_1000'] < 5 and total_words > 500:
            risk_score += 1
            confidence = min(confidence + 0.1, 1.0)
        
        # Epistemic markers (too few)
        if results['epistemic_per_1000'] < 2 and total_words > 500:
            risk_score += 1
            confidence = min(confidence + 0.1, 1.0)
        
        # Certainty boosters (too many)
        if results['boosters_per_1000'] > 5 and total_words > 500:
            risk_score += 1
            confidence = min(confidence + 0.1, 1.0)
        
        results['risk_score'] = risk_score
        results['confidence'] = confidence
        
        if risk_score >= 4:
            results['risk_level'] = 'high'
        elif risk_score >= 2:
            results['risk_level'] = 'medium'
        elif risk_score >= 1:
            results['risk_level'] = 'low'
        
        # Passive examples
        for match in passive_matches[:50]:
            results['examples'].append(f"...{' '.join(match)}...")
        
        return results


class HedgingAnalyzer:
    """Hedging analysis (words of uncertainty)"""
    
    def __init__(self):
        self.hedging_words = [
            'may', 'might', 'could', 'would', 'should',
            'possibly', 'probably', 'perhaps', 'maybe',
            'likely', 'unlikely', 'potentially',
            'appears', 'seems', 'suggests', 'indicates',
            'generally', 'typically', 'usually', 'often',
            'to some extent', 'in some cases', 'relatively',
            'somewhat', 'partially', 'tentatively', 'arguably'
        ]
        
        self.personal_pronouns = ['we', 'our', 'us', 'i', 'my', 'mine']
        
        # Categorical expressions (anti-hedging)
        self.certainty_phrases = [
            'clearly', 'obviously', 'undoubtedly', 'certainly',
            'definitely', 'absolutely', 'without doubt',
            'it is clear that', 'it is obvious that', 'there is no doubt'
        ]
        
    def analyze(self, text: str) -> Dict:
        """Analyze hedging usage"""
        results = {
            'hedging_count': 0,
            'hedging_per_1000': 0,
            'certainty_count': 0,
            'certainty_per_1000': 0,
            'personal_count': 0,
            'personal_per_1000': 0,
            'hedging_ratio': 0,  # hedging / certainty
            'total_words': 0,
            'risk_level': 'none',
            'risk_score': 0,
            'confidence': 0,
            'statistics': {
                'mean_hedging_per_sentence': 0,
                'median_hedging_per_sentence': 0,
                'max_hedging_in_sentence': 0,
                'distribution': []
            }
        }
        
        if not text:
            return results
        
        text_lower = text.lower()
        words = text_lower.split()
        results['total_words'] = len(words)
        
        # Count hedging
        for word in self.hedging_words:
            if ' ' in word:
                results['hedging_count'] += text_lower.count(word)
            else:
                results['hedging_count'] += sum(1 for w in words if w == word)
        
        # Count categorical expressions
        for phrase in self.certainty_phrases:
            if ' ' in phrase:
                results['certainty_count'] += text_lower.count(phrase)
            else:
                results['certainty_count'] += sum(1 for w in words if w == phrase)
        
        # Count personal pronouns
        for pronoun in self.personal_pronouns:
            results['personal_count'] += sum(1 for w in words if w == pronoun)
        
        # Normalize per 1000 words
        if results['total_words'] > 0:
            results['hedging_per_1000'] = (results['hedging_count'] * 1000) / results['total_words']
            results['certainty_per_1000'] = (results['certainty_count'] * 1000) / results['total_words']
            results['personal_per_1000'] = (results['personal_count'] * 1000) / results['total_words']
            
            if results['certainty_count'] > 0:
                results['hedging_ratio'] = results['hedging_count'] / results['certainty_count']
        
        # Statistics by sentence (simple)
        sentences = re.split(r'[.!?]+', text)
        hedging_per_sentence = []
        
        for sent in sentences[:100]:
            sent_lower = sent.lower()
            sent_hedging = sum(1 for word in self.hedging_words if word in sent_lower)
            hedging_per_sentence.append(sent_hedging)
        
        if hedging_per_sentence:
            results['statistics']['mean_hedging_per_sentence'] = float(np.mean(hedging_per_sentence))
            results['statistics']['median_hedging_per_sentence'] = float(np.median(hedging_per_sentence))
            results['statistics']['max_hedging_in_sentence'] = float(np.max(hedging_per_sentence))
            results['statistics']['distribution'] = hedging_per_sentence
        
        # Probabilistic risk assessment
        risk_score = 0
        confidence = 0.5
        
        # Too little hedging - main marker
        if results['hedging_per_1000'] < 3:
            risk_score += 3
            confidence = min(confidence + 0.3, 1.0)
        elif results['hedging_per_1000'] < 5:
            risk_score += 2
            confidence = min(confidence + 0.2, 1.0)
        elif results['hedging_per_1000'] < 7:
            risk_score += 1
            confidence = min(confidence + 0.1, 1.0)
        
        # Too much certainty
        if results['certainty_per_1000'] > 5:
            risk_score += 2
            confidence = min(confidence + 0.2, 1.0)
        elif results['certainty_per_1000'] > 3:
            risk_score += 1
            confidence = min(confidence + 0.1, 1.0)
        
        # Poor hedging/certainty ratio
        if results['hedging_ratio'] < 0.5 and results['certainty_count'] > 5:
            risk_score += 1
            confidence = min(confidence + 0.1, 1.0)
        
        # Personal pronouns (too few - impersonal style, could be AI)
        if results['personal_per_1000'] < 0.5 and results['total_words'] > 500:
            risk_score += 1
            confidence = min(confidence + 0.1, 1.0)
        
        results['risk_score'] = risk_score
        results['confidence'] = confidence
        
        if risk_score >= 5:
            results['risk_level'] = 'high'
        elif risk_score >= 3:
            results['risk_level'] = 'medium-high'
        elif risk_score >= 2:
            results['risk_level'] = 'medium'
        elif risk_score >= 1:
            results['risk_level'] = 'low'
        
        return results


class ParenthesisAnalyzer:
    """Long parenthesis content analysis (new module)"""
    
    def __init__(self):
        self.min_words_for_long = 5  # 5+ words in parentheses - human indicator
    
    def analyze(self, text: str) -> Dict:
        """Analyze content inside parentheses"""
        results = {
            'total_parentheses': 0,
            'short_parentheses': 0,  # <5 words
            'long_parentheses': 0,   # >=5 words
            'long_percentage': 0,
            'long_examples': [],
            'all_parentheses': [],  # All found parentheses
            'words_in_parentheses': [],
            'risk_level': 'none',
            'risk_score': 0,
            'confidence': 0,
            'statistics': {
                'mean_words_in_parentheses': 0,
                'median_words_in_parentheses': 0,
                'max_words_in_parentheses': 0,
                'distribution': []
            }
        }
        
        if not text:
            return results

        # Find all occurrences in parentheses
        pattern = r'\(([^)]+)\)'
        matches = re.findall(pattern, text)
        
        results['total_parentheses'] = len(matches)
        
        word_counts = []
        
        for match in matches:
            words = match.split()
            word_count = len(words)
            word_counts.append(word_count)
            
            # Добавляем в all_parentheses ТОЛЬКО если слов >= 4
            if word_count >= self.min_words_for_long:
                parenthesis_data = {
                    'text': match[:200] + '...' if len(match) > 200 else match,
                    'word_count': word_count
                }
                results['all_parentheses'].append(parenthesis_data)
                
                results['long_parentheses'] += 1
                if len(results['long_examples']) < 20:
                    results['long_examples'].append(parenthesis_data)
            else:
                results['short_parentheses'] += 1
            
            results['words_in_parentheses'].append(word_count)
        
        if results['total_parentheses'] > 0:
            results['long_percentage'] = (results['long_parentheses'] / results['total_parentheses']) * 100
        
        # Statistics
        if word_counts:
            results['statistics']['mean_words_in_parentheses'] = float(np.mean(word_counts))
            results['statistics']['median_words_in_parentheses'] = float(np.median(word_counts))
            results['statistics']['max_words_in_parentheses'] = float(np.max(word_counts))
            results['statistics']['distribution'] = word_counts
        
        # Risk assessment (more long parentheses = lower AI risk)
        risk_score = 0
        confidence = 0.5
        
        if results['long_percentage'] > 15:
            risk_score = 0
            confidence = 0.8
        elif results['long_percentage'] > 8:
            risk_score = 1
            confidence = 0.6
        elif results['long_percentage'] > 3:
            risk_score = 2
            confidence = 0.5
        else:
            risk_score = 3
            confidence = 0.7
        
        results['risk_score'] = risk_score
        results['confidence'] = confidence
        
        if risk_score >= 3:
            results['risk_level'] = 'high'
        elif risk_score >= 2:
            results['risk_level'] = 'medium'
        elif risk_score >= 1:
            results['risk_level'] = 'low'
        else:
            results['risk_level'] = 'very_low'
        
        return results


class PunctuationAnalyzer:
    """Punctuation analysis ! ? ; (new module)"""
    
    def __init__(self):
        self.punctuation_marks = {
            'exclamation': '!',
            'question': '?',
            'semicolon': ';'
        }
    
    def analyze(self, text: str, sentences: List[str]) -> Dict:
        """Analyze usage of rare punctuation marks"""
        results = {
            'exclamation_count': 0,
            'question_count': 0,
            'semicolon_count': 0,
            'exclamation_per_1000': 0,
            'question_per_1000': 0,
            'semicolon_per_1000': 0,
            'semicolon_contexts': [],
            'all_semicolon_contexts': [],  # All contexts with ;
            'risk_level': 'none',
            'risk_score': 0,
            'confidence': 0,
            'statistics': {
                'mean_punctuation_per_sentence': 0,
                'median_punctuation_per_sentence': 0,
                'max_punctuation_in_sentence': 0,
                'distribution': {}
            }
        }
        
        if not text or not sentences:
            return results
        
        words = text.split()
        total_words = len(words)
        
        # Count punctuation marks
        results['exclamation_count'] = text.count('!')
        results['question_count'] = text.count('?')
        results['semicolon_count'] = text.count(';')
        
        # Normalize per 1000 words
        if total_words > 0:
            results['exclamation_per_1000'] = (results['exclamation_count'] * 1000) / total_words
            results['question_per_1000'] = (results['question_count'] * 1000) / total_words
            results['semicolon_per_1000'] = (results['semicolon_count'] * 1000) / total_words
        
        # Context for semicolons (especially important)
        semicolon_positions = [m.start() for m in re.finditer(r';', text)]
        for pos in semicolon_positions[:50]:  # Limit for performance
            context = text[max(0, pos-60):min(len(text), pos+60)]
            # Check if it separates long and short sentences
            before = context[:60].split(';')[0] if ';' in context[:60] else ''
            after = context[60:].split(';')[0] if ';' in context[60:] else ''
            
            before_words = len(before.split())
            after_words = len(after.split())
            
            context_data = {
                'context': context.replace('\n', ' ').strip(),
                'before_words': before_words,
                'after_words': after_words,
                'is_contrast': after_words < before_words * 0.3
            }
            results['all_semicolon_contexts'].append(context_data)
            if len(results['semicolon_contexts']) < 10:
                results['semicolon_contexts'].append(context_data)
        
        # Statistics by sentence
        punct_per_sentence = []
        for sent in sentences[:100]:
            sent_punct = sent.count('!') + sent.count('?') + sent.count(';')
            punct_per_sentence.append(sent_punct)
        
        if punct_per_sentence:
            results['statistics']['mean_punctuation_per_sentence'] = float(np.mean(punct_per_sentence))
            results['statistics']['median_punctuation_per_sentence'] = float(np.median(punct_per_sentence))
            results['statistics']['max_punctuation_in_sentence'] = float(np.max(punct_per_sentence))
            results['statistics']['distribution'] = {
                'exclamation': results['exclamation_count'],
                'question': results['question_count'],
                'semicolon': results['semicolon_count']
            }
        
        # Risk assessment (rare punctuation = human indicator)
        risk_score = 0
        confidence = 0.5
        
        # If exclamation marks exist (rare in science) - human indicator
        if results['exclamation_per_1000'] > 0.1:
            risk_score -= 1
            confidence = min(confidence + 0.1, 1.0)
        
        # If question marks exist (rhetorical questions) - human indicator
        if results['question_per_1000'] > 0.2:
            risk_score -= 1
            confidence = min(confidence + 0.1, 1.0)
        
        # Semicolon - strong human indicator (AI avoids)
        if results['semicolon_per_1000'] > 0.3:
            risk_score -= 2
            confidence = min(confidence + 0.2, 1.0)
        elif results['semicolon_per_1000'] > 0.1:
            risk_score -= 1
            confidence = min(confidence + 0.1, 1.0)
        
        # Normalize risk (cannot be negative)
        risk_score = max(0, risk_score)
        
        results['risk_score'] = risk_score
        results['confidence'] = confidence
        
        if risk_score == 0:
            results['risk_level'] = 'very_low'
        elif risk_score == 1:
            results['risk_level'] = 'low'
        elif risk_score == 2:
            results['risk_level'] = 'medium'
        else:
            results['risk_level'] = 'high'
        
        return results

class ApostropheAnalyzer:
    """Apostrophe 's analysis with improved detection"""
    
    def __init__(self):
        self.apostrophe_pattern = r"\b\w+'\w+\b"  # Words with apostrophe
        self.contraction_pattern = r"\b\w+'(?:t|re|ve|ll|m|d)\b"  # Common contractions
        self.possessive_pattern = r"\b\w+'(?:s)?\b"  # Possessives
    
    def analyze(self, text: str) -> Dict:
        """Analyze apostrophe usage"""
        results = {
            'apostrophe_count': 0,
            'apostrophe_per_1000': 0,
            'possessive_examples': [],
            'contraction_examples': [],
            'all_apostrophes': [],
            'risk_level': 'none',
            'risk_score': 0,
            'confidence': 0,
            'statistics': {
                'mean_apostrophes_per_paragraph': 0,
                'median_apostrophes_per_paragraph': 0,
                'max_apostrophes_in_paragraph': 0,
                'distribution': []
            }
        }
        
        if not text:
            return results
        
        words = text.split()
        total_words = len(words)
        
        # Find all words with apostrophes (improved pattern)
        apostrophe_matches = re.findall(self.apostrophe_pattern, text)
        
        # Also catch standalone apostrophes (like in "O'Hayre")
        standalone_apostrophe = re.findall(r"\b[A-Z]'[A-Za-z]+\b", text)
        apostrophe_matches.extend(standalone_apostrophe)
        
        # Deduplicate while preserving order
        seen = set()
        unique_matches = []
        for m in apostrophe_matches:
            if m not in seen:
                seen.add(m)
                unique_matches.append(m)
        
        results['apostrophe_count'] = len(unique_matches)
        results['all_apostrophes'] = unique_matches
        
        # Normalize per 1000 words
        if total_words > 0:
            results['apostrophe_per_1000'] = (results['apostrophe_count'] * 1000) / total_words
        
        # Classify examples
        for match in unique_matches:
            if match.endswith("'s") and len(match) > 2:
                base_word = match[:-2]
                if base_word.isalpha() or (base_word[0].isupper() and len(base_word) > 1):
                    if len(results['possessive_examples']) < 50:
                        results['possessive_examples'].append(match)
            elif any(cont in match for cont in ["'t", "'re", "'ve", "'ll", "'m", "'d"]):
                if len(results['contraction_examples']) < 30:
                    results['contraction_examples'].append(match)
        
        # Statistics by paragraph (improved)
        paragraphs = re.split(r'\n\s*\n', text)
        if len(paragraphs) <= 1:
            # Try alternative paragraph splitting
            paragraphs = [p for p in text.split('\n') if len(p.strip()) > 50]
        
        apostrophes_per_paragraph = []
        for para in paragraphs[:50]:
            para_apostrophes = len(re.findall(self.apostrophe_pattern, para))
            apostrophes_per_paragraph.append(para_apostrophes)
        
        if apostrophes_per_paragraph:
            results['statistics']['mean_apostrophes_per_paragraph'] = float(np.mean(apostrophes_per_paragraph))
            results['statistics']['median_apostrophes_per_paragraph'] = float(np.median(apostrophes_per_paragraph))
            results['statistics']['max_apostrophes_in_paragraph'] = float(np.max(apostrophes_per_paragraph))
            results['statistics']['distribution'] = apostrophes_per_paragraph
        
        # Risk assessment (inverted: more apostrophes = more human-like)
        risk_score = 0
        confidence = 0.5
        
        if results['apostrophe_per_1000'] > 2.0:
            risk_score = 1  # Low risk (human-like)
            confidence = 0.7
        elif results['apostrophe_per_1000'] > 1.0:
            risk_score = 2  # Medium-low risk
            confidence = 0.6
        elif results['apostrophe_per_1000'] > 0.3:
            risk_score = 3  # Medium risk
            confidence = 0.5
        else:
            risk_score = 4  # Higher risk (AI avoids apostrophes)
            confidence = 0.65
        
        # Adjust: if we found many possessives but no contractions, that's normal for academic
        if len(results['possessive_examples']) > 5 and len(results['contraction_examples']) == 0:
            risk_score = max(1, risk_score - 1)  # Less suspicious
        
        results['risk_score'] = min(risk_score, 6)
        results['confidence'] = confidence
        
        if risk_score >= 5:
            results['risk_level'] = 'critical'
        elif risk_score >= 4:
            results['risk_level'] = 'high'
        elif risk_score >= 2:
            results['risk_level'] = 'medium'
        elif risk_score >= 1:
            results['risk_level'] = 'low'
        else:
            results['risk_level'] = 'very_low'
        
        return results

class EnumerationAnalyzer:
    """Strict enumeration analysis - improved version that catches all three-item enumerations"""
    
    def __init__(self):
        # Pattern for three-item enumerations: X, Y, and Z
        # Allows for multi-word items with spaces
        self.enumeration_pattern = r'([^,]+(?:,\s+[^,]+)+,\s+and\s+[^,]+)'
        # Pattern for enumerations with specifically
        self.specifically_pattern = r'\bspecifically\s+[^,]+,\s+[^,]+,\s+and\s+[^,]+'
    
    def analyze(self, text: str, sentences: List[str]) -> Dict:
        """Analyze strict three-or-more item enumerations"""
        results = {
            'three_item_count': 0,
            'three_item_per_1000_sentences': 0,
            'specifically_count': 0,
            'examples': [],
            'all_enumerations': [],  # Все найденные перечисления с полным контекстом
            'invalid_enumerations': [],  # Перечисления с 4+ элементами (для отладки)
            'risk_level': 'none',
            'risk_score': 0,
            'confidence': 0,
            'statistics': {
                'mean_enumerations_per_sentence': 0,
                'max_enumerations_in_sentence': 0,
                'distribution': []
            }
        }
        
        if not text or not sentences:
            return results
        
        total_sentences = len(sentences)
        
        # Анализируем каждое предложение отдельно для лучшего контекста
        for sent_idx, sentence in enumerate(sentences):
            if not sentence or len(sentence.strip()) < 20:
                continue
            
            # Ищем все возможные перечисления в предложении
            matches = re.findall(self.enumeration_pattern, sentence, re.IGNORECASE | re.DOTALL)
            
            for match in matches:
                # Очищаем от лишних пробелов и переносов
                clean_match = re.sub(r'\s+', ' ', match).strip()
                
                # Подсчитываем количество запятых в найденном фрагменте
                comma_count = clean_match.count(',')
                
                # Для трех элементов должно быть ровно 2 запятые (перед and)
                # Если запятых больше 2, значит это перечисление из 4+ элементов
                if comma_count == 2:
                    # Убеждаемся, что это действительно 3 элемента
                    parts = clean_match.split(',')
                    if len(parts) == 3:  # Должно быть "X, Y and Z"
                        # Проверяем наличие "and" во второй части
                        if ' and ' in parts[1]:
                            enumeration_data = {
                                'sentence_idx': sent_idx,
                                'full_sentence': sentence[:500] if len(sentence) > 500 else sentence,
                                'enumeration': clean_match,
                                'comma_count': comma_count,
                                'word_count': len(clean_match.split()),
                                'context': self._get_context(sentence, clean_match)
                            }
                            results['all_enumerations'].append(enumeration_data)
                            
                            if len(results['examples']) < 100:
                                results['examples'].append(clean_match)
                else:
                    # Сохраняем для отладки, если нужно
                    if comma_count > 2:
                        results['invalid_enumerations'].append({
                            'enumeration': clean_match,
                            'comma_count': comma_count
                        })
        
        # Ищем перечисления со словом "specifically"
        specifically_matches = re.findall(self.specifically_pattern, text, re.IGNORECASE)
        results['specifically_count'] = len(specifically_matches)
        
        results['three_item_count'] = len(results['all_enumerations'])
        
        # Нормализуем на 1000 предложений
        if total_sentences > 0:
            results['three_item_per_1000_sentences'] = (results['three_item_count'] * 1000) / total_sentences
        
        # Статистика по предложениям
        enumerations_per_sentence = []
        for sent in sentences[:200]:
            sent_enumerations = len(re.findall(self.enumeration_pattern, sent, re.IGNORECASE))
            if sent_enumerations > 0:
                enumerations_per_sentence.append(sent_enumerations)
        
        if enumerations_per_sentence:
            results['statistics']['mean_enumerations_per_sentence'] = float(np.mean(enumerations_per_sentence))
            results['statistics']['max_enumerations_in_sentence'] = float(np.max(enumerations_per_sentence))
            results['statistics']['distribution'] = enumerations_per_sentence
        
        # Risk assessment
        risk_score = 0
        confidence = 0.5
        
        if results['three_item_per_1000_sentences'] > 15:
            risk_score = 3
            confidence = 0.9
        elif results['three_item_per_1000_sentences'] > 8:
            risk_score = 2
            confidence = 0.7
        elif results['three_item_per_1000_sentences'] > 3:
            risk_score = 1
            confidence = 0.6
        
        # Enumerations with specifically - even stronger indicator
        if results['specifically_count'] > 3:
            risk_score = min(risk_score + 1, 3)
            confidence = min(confidence + 0.1, 1.0)
        
        results['risk_score'] = risk_score
        results['confidence'] = confidence
        
        if risk_score >= 3:
            results['risk_level'] = 'high'
        elif risk_score >= 2:
            results['risk_level'] = 'medium'
        elif risk_score >= 1:
            results['risk_level'] = 'low'
        else:
            results['risk_level'] = 'very_low'
        
        return results
    
    def _get_context(self, sentence: str, enumeration: str) -> str:
        """Get context around the enumeration within the sentence"""
        if enumeration in sentence:
            parts = sentence.split(enumeration)
            before = parts[0][-50:] if len(parts[0]) > 50 else parts[0]
            after = parts[1][:50] if len(parts) > 1 and len(parts[1]) > 50 else (parts[1] if len(parts) > 1 else '')
            return f"...{before}[{enumeration}]{after}..."
        return enumeration

class ParagraphAnalyzer:
    """Paragraph-level analysis (new module)"""
    
    def __init__(self):
        self.available = SENTENCE_TRANSFORMERS_AVAILABLE
        self.model = None
        if self.available:
            try:
                self.model = SentenceTransformer('all-MiniLM-L6-v2')
            except:
                self.available = False
    
    def split_paragraphs(self, text: str) -> List[str]:
        """
        Improved paragraph splitting that handles:
        - Double newlines
        - Single newlines with indentation
        - Numbered/bulleted lists
        - Section headers
        """
        if not text:
            return []
        
        # First, normalize line endings
        text = text.replace('\r\n', '\n')
        
        paragraphs = []
        
        # Method 1: Try double newline separation
        double_newline_paras = re.split(r'\n\s*\n', text)
        if len(double_newline_paras) > 1:
            # Filter out empty paragraphs and very short ones
            for p in double_newline_paras:
                p_clean = p.strip()
                if len(p_clean) > 50:  # Reasonable minimum paragraph length
                    paragraphs.append(p_clean)
            
            # If we found multiple paragraphs, return them
            if len(paragraphs) > 1:
                return paragraphs
        
        # Method 2: Split by lines and detect paragraph boundaries
        lines = text.split('\n')
        current_paragraph = []
        paragraphs = []
        
        # Patterns that indicate a new paragraph
        new_para_patterns = [
            r'^\d+\.\s+',           # Numbered section: "1.", "2.1."
            r'^[A-Z][a-z]+\.\s+',   # Lettered section: "A.", "B."
            r'^[IVXLCDM]+\.\s+',    # Roman numerals: "I.", "II."
            r'^Abstract\.?\s*$',    # Abstract section
            r'^Introduction\.?\s*$', # Introduction
            r'^[A-Z][a-z]+\s+[A-Z][a-z]+\.?\s*$', # Section header like "Results and Discussion"
            r'^Keywords?:',         # Keywords line
            r'^\*+\s*$',            # Separator line
            r'^---+$',              # Horizontal rule
        ]
        
        # Combined pattern
        new_para_pattern = re.compile('|'.join(new_para_patterns), re.IGNORECASE)
        
        for line in lines:
            line_stripped = line.strip()
            
            # Skip empty lines
            if not line_stripped:
                if current_paragraph:
                    # End of paragraph on empty line
                    paragraph_text = ' '.join(current_paragraph)
                    if len(paragraph_text.split()) > 10:  # Minimum words for a paragraph
                        paragraphs.append(paragraph_text)
                    current_paragraph = []
                continue
            
            # Check if this line starts a new paragraph
            is_new_paragraph = False
            
            # Pattern-based detection
            if new_para_pattern.match(line_stripped):
                is_new_paragraph = True
            # Indentation detection (line starts with spaces)
            elif line.startswith(' ') or line.startswith('\t'):
                is_new_paragraph = True
            # Very short line that might be a header
            elif len(line_stripped.split()) <= 5 and len(line_stripped) < 40:
                is_new_paragraph = True
            # Line ends with colon (section header)
            elif line_stripped.endswith(':'):
                is_new_paragraph = True
            
            if is_new_paragraph and current_paragraph:
                # Save previous paragraph
                paragraph_text = ' '.join(current_paragraph)
                if len(paragraph_text.split()) > 10:
                    paragraphs.append(paragraph_text)
                current_paragraph = [line_stripped]
            else:
                current_paragraph.append(line_stripped)
        
        # Add last paragraph
        if current_paragraph:
            paragraph_text = ' '.join(current_paragraph)
            if len(paragraph_text.split()) > 10:
                paragraphs.append(paragraph_text)
        
        # Method 3: If still only one paragraph, try splitting by section headers
        if len(paragraphs) <= 1:
            # Look for numbered sections like "1. Introduction", "3.2. Results"
            section_pattern = r'\n(\d+(?:\.\d+)*\.?\s+[A-Z][^\n]+)\n'
            sections = re.split(section_pattern, text)
            
            if len(sections) > 2:
                paragraphs = []
                for i in range(1, len(sections), 2):
                    header = sections[i].strip()
                    content = sections[i+1].strip() if i+1 < len(sections) else ''
                    full_section = f"{header}\n{content}".strip()
                    if len(full_section.split()) > 20:
                        paragraphs.append(full_section)
        
        # Final cleaning
        paragraphs = [p for p in paragraphs if len(p.split()) > 15]
        
        # If still no paragraphs, treat entire text as one paragraph with warning
        if not paragraphs and len(text.split()) > 50:
            paragraphs = [text.strip()]
            self._paragraph_warning = True  # Store warning flag
        
        return paragraphs
    
    def analyze(self, text: str, sentences: List[str]) -> Dict:
        """Analyze intra and inter paragraph homogeneity"""
        results = {
            'paragraphs': [],
            'paragraph_lengths': [],
            'intra_paragraph_similarity': 0,
            'inter_paragraph_similarity': 0,
            'transition_smoothness': 0,
            'structure_repetition': 0,
            'risk_score': 0,
            'confidence': 0,
            'risk_level': 'none',
            'note': 'Paragraph analysis requires sentence-transformers',
            'statistics': {
                'mean_paragraph_length': 0,
                'median_paragraph_length': 0,
                'max_paragraph_length': 0,
                'min_paragraph_length': 0,
                'std_paragraph_length': 0,
                'distribution': []
            }
        }
        
        if not self.available or len(text) < 500:
            return results
        
        paragraphs = self.split_paragraphs(text)
        if len(paragraphs) < 3:
            return results
        
        results['paragraphs'] = paragraphs[:15]
        
        # Paragraph length statistics
        para_lengths = [len(p.split()) for p in paragraphs]
        results['paragraph_lengths'] = para_lengths[:50]
        
        if para_lengths:
            results['statistics']['mean_paragraph_length'] = float(np.mean(para_lengths))
            results['statistics']['median_paragraph_length'] = float(np.median(para_lengths))
            results['statistics']['max_paragraph_length'] = float(np.max(para_lengths))
            results['statistics']['min_paragraph_length'] = float(np.min(para_lengths))
            results['statistics']['std_paragraph_length'] = float(np.std(para_lengths))
            results['statistics']['distribution'] = para_lengths[:100]
        
        # If no model, use simple metrics
        if not self.model:
            # Simple metric: paragraph length variability
            if len(para_lengths) > 1:
                cv_lengths = np.std(para_lengths) / np.mean(para_lengths) if np.mean(para_lengths) > 0 else 0
                if cv_lengths < 0.3:
                    results['risk_score'] = 2
                    results['confidence'] = 0.6
                    results['risk_level'] = 'medium'
                    results['note'] = 'Simple paragraph length analysis'
            return results
        
        # Full analysis with embeddings
        try:
            # Get embeddings for paragraphs
            para_embeddings = self.model.encode(paragraphs[:25])
            
            # Intra-paragraph similarity
            intra_similarities = []
            for para in paragraphs[:8]:
                para_sents = [s for s in re.split(r'[.!?]+', para) if len(s.split()) > 5]
                if len(para_sents) >= 3:
                    sent_embeddings = self.model.encode(para_sents[:10])
                    similarities = []
                    for i in range(len(sent_embeddings)):
                        for j in range(i+1, len(sent_embeddings)):
                            sim = np.dot(sent_embeddings[i], sent_embeddings[j]) / (
                                np.linalg.norm(sent_embeddings[i]) * np.linalg.norm(sent_embeddings[j]))
                            similarities.append(sim)
                    if similarities:
                        intra_similarities.append(np.mean(similarities))
            
            if intra_similarities:
                results['intra_paragraph_similarity'] = float(np.mean(intra_similarities))
            
            # Inter-paragraph similarity
            if len(para_embeddings) > 1:
                inter_similarities = []
                for i in range(len(para_embeddings) - 1):
                    sim = np.dot(para_embeddings[i], para_embeddings[i+1]) / (
                        np.linalg.norm(para_embeddings[i]) * np.linalg.norm(para_embeddings[i+1]))
                    inter_similarities.append(sim)
                results['inter_paragraph_similarity'] = float(np.mean(inter_similarities))
            
            # Transition smoothness
            transition_sims = []
            for i in range(len(paragraphs) - 1):
                curr_sents = [s for s in re.split(r'[.!?]+', paragraphs[i]) if len(s.split()) > 5]
                next_sents = [s for s in re.split(r'[.!?]+', paragraphs[i+1]) if len(s.split()) > 5]
                
                if curr_sents and next_sents:
                    last_curr = self.model.encode([curr_sents[-1]])[0]
                    first_next = self.model.encode([next_sents[0]])[0]
                    sim = np.dot(last_curr, first_next) / (np.linalg.norm(last_curr) * np.linalg.norm(first_next))
                    transition_sims.append(sim)
            
            if transition_sims:
                results['transition_smoothness'] = float(np.mean(transition_sims))
            
            # Risk assessment
            risk_score = 0
            confidence = 0.5
            
            if results['intra_paragraph_similarity'] > 0.72:
                risk_score += 3
                confidence = min(confidence + 0.3, 1.0)
            elif results['intra_paragraph_similarity'] > 0.65:
                risk_score += 2
                confidence = min(confidence + 0.2, 1.0)
            elif results['intra_paragraph_similarity'] > 0.6:
                risk_score += 1
                confidence = min(confidence + 0.1, 1.0)
            
            if results['transition_smoothness'] > 0.65:
                risk_score += 2
                confidence = min(confidence + 0.2, 1.0)
            elif results['transition_smoothness'] > 0.55:
                risk_score += 1
                confidence = min(confidence + 0.1, 1.0)
            
            if results['inter_paragraph_similarity'] > 0.5:
                risk_score += 1
                confidence = min(confidence + 0.1, 1.0)
            
            results['risk_score'] = risk_score
            results['confidence'] = confidence
            
            if risk_score >= 4:
                results['risk_level'] = 'high'
            elif risk_score >= 2:
                results['risk_level'] = 'medium'
            elif risk_score >= 1:
                results['risk_level'] = 'low'
            
            results['note'] = 'Full semantic paragraph analysis'
            
        except Exception as e:
            results['note'] = f'Error in paragraph analysis: {str(e)}'
        
        return results


class PerplexityAnalyzer:
    """Perplexity analysis (full version)"""
    
    def __init__(self):
        self.available = TRANSFORMERS_AVAILABLE
        self.model = None
        self.tokenizer = None
        self.perplexity_pipeline = None
        
        if self.available:
            try:
                model_name = "gpt2"
                self.tokenizer = AutoTokenizer.from_pretrained(model_name)
                self.model = AutoModelForCausalLM.from_pretrained(model_name)
                self.perplexity_pipeline = pipeline(
                    "text-generation",
                    model=self.model,
                    tokenizer=self.tokenizer,
                    max_length=512
                )
            except Exception as e:
                self.available = False
    
    def calculate_perplexity(self, text: str) -> float:
        """Calculate text perplexity"""
        if not self.available or not text:
            return 0
        
        try:
            inputs = self.tokenizer(text, return_tensors="pt", truncation=True, max_length=512)
            
            with torch.no_grad():
                outputs = self.model(**inputs, labels=inputs["input_ids"])
                loss = outputs.loss
                perplexity = torch.exp(loss).item()
            
            return perplexity
        except Exception as e:
            return 0
    
    def analyze(self, text: str) -> Dict:
        """Full perplexity analysis"""
        results = {
            'perplexities': [],
            'mean_perplexity': 0,
            'median_perplexity': 0,
            'perplexity_variance': 0,
            'min_perplexity': 0,
            'max_perplexity': 0,
            'perplexity_segments': [],
            'all_perplexities': [],  # All perplexity values
            'risk_level': 'none',
            'risk_score': 0,
            'confidence': 0,
            'note': 'Perplexity analysis requires transformers library',
            'statistics': {
                'distribution': []
            }
        }
        
        if not self.available or len(text) < 100:
            results['note'] = 'Perplexity analysis unavailable or text too short'
            return results
        
        try:
            words = text.split()
            segment_size = 500
            segments = [' '.join(words[i:i+segment_size]) for i in range(0, len(words), segment_size)]
            segments = segments[:20]
            
            perplexities = []
            
            for i, segment in enumerate(segments):
                if len(segment.split()) < 50:
                    continue
                
                ppl = self.calculate_perplexity(segment)
                if ppl > 0:
                    perplexities.append(ppl)
                    results['perplexities'].append(ppl)
                    results['all_perplexities'].append(ppl)
                    results['perplexity_segments'].append({
                        'segment': i,
                        'perplexity': ppl
                    })
            
            if perplexities:
                results['mean_perplexity'] = float(np.mean(perplexities))
                results['median_perplexity'] = float(np.median(perplexities))
                results['min_perplexity'] = float(np.min(perplexities))
                results['max_perplexity'] = float(np.max(perplexities))
                results['perplexity_variance'] = float(np.var(perplexities))
                results['statistics']['distribution'] = perplexities
            
            # Risk assessment
            risk_score = 0
            confidence = 0.5
            
            if results['mean_perplexity'] < 15:
                risk_score = 3
                confidence = 0.8
            elif results['mean_perplexity'] < 25:
                risk_score = 2
                confidence = 0.7
            elif results['mean_perplexity'] < 40:
                risk_score = 1
                confidence = 0.6
            else:
                risk_score = 0
                confidence = 0.5
            
            results['risk_score'] = risk_score
            results['confidence'] = confidence
            
            if risk_score >= 3:
                results['risk_level'] = 'high'
            elif risk_score >= 2:
                results['risk_level'] = 'medium'
            elif risk_score >= 1:
                results['risk_level'] = 'low'
            else:
                results['risk_level'] = 'very_low'
            
            results['note'] = 'Full perplexity analysis'
            
        except Exception as e:
            results['note'] = f'Error in perplexity analysis: {str(e)}'
        
        return results


class SemanticAnalyzer:
    """Semantic similarity analysis (full version)"""
    
    def __init__(self):
        self.available = SENTENCE_TRANSFORMERS_AVAILABLE
        self.model = None
        if self.available:
            try:
                self.model = SentenceTransformer('all-MiniLM-L6-v2')
            except:
                self.available = False
    
    def analyze(self, sentences: List[str]) -> Dict:
        """Full semantic analysis"""
        results = {
            'similarities': [],
            'mean_similarity': 0,
            'similarity_variance': 0,
            'min_similarity': 0,
            'max_similarity': 0,
            'similarity_matrix': [],
            'semantic_clusters': 0,
            'all_similarities': [],  # All similarity values
            'risk_level': 'none',
            'risk_score': 0,
            'confidence': 0,
            'note': 'Semantic analysis requires sentence-transformers',
            'statistics': {
                'distribution': []
            }
        }
        
        if not self.available or len(sentences) < 10:
            return results
        
        try:
            sample_sentences = [s for s in sentences if len(s.split()) > 5][:50]
            
            if len(sample_sentences) < 5:
                return results
            
            embeddings = self.model.encode(sample_sentences)
            
            similarities = []
            n = len(embeddings)
            
            for i in range(n):
                for j in range(i+1, n):
                    sim = np.dot(embeddings[i], embeddings[j]) / (
                        np.linalg.norm(embeddings[i]) * np.linalg.norm(embeddings[j]))
                    similarities.append(sim)
            
            if similarities:
                results['similarities'] = similarities[:1000]
                results['all_similarities'] = similarities
                results['mean_similarity'] = float(np.mean(similarities))
                results['similarity_variance'] = float(np.var(similarities))
                results['min_similarity'] = float(np.min(similarities))
                results['max_similarity'] = float(np.max(similarities))
                results['statistics']['distribution'] = similarities[:500]
                
                # Simple clustering
                threshold = 0.7
                clusters = 0
                used = set()
                
                for i in range(n):
                    if i in used:
                        continue
                    cluster = [i]
                    for j in range(i+1, n):
                        if j in used:
                            continue
                        sim = np.dot(embeddings[i], embeddings[j]) / (
                            np.linalg.norm(embeddings[i]) * np.linalg.norm(embeddings[j]))
                        if sim > threshold:
                            cluster.append(j)
                            used.add(j)
                    if len(cluster) > 1:
                        clusters += 1
                    used.add(i)
                
                results['semantic_clusters'] = clusters
            
            # Risk assessment
            risk_score = 0
            confidence = 0.5
            
            if results['mean_similarity'] > 0.65:
                risk_score = 3
                confidence = 0.8
            elif results['mean_similarity'] > 0.55:
                risk_score = 2
                confidence = 0.7
            elif results['mean_similarity'] > 0.45:
                risk_score = 1
                confidence = 0.6
            else:
                risk_score = 0
                confidence = 0.5
            
            results['risk_score'] = risk_score
            results['confidence'] = confidence
            
            if risk_score >= 3:
                results['risk_level'] = 'high'
            elif risk_score >= 2:
                results['risk_level'] = 'medium'
            elif risk_score >= 1:
                results['risk_level'] = 'low'
            else:
                results['risk_level'] = 'very_low'
            
            results['note'] = 'Full semantic analysis'
            
        except Exception as e:
            results['note'] = f'Error in semantic analysis: {str(e)}'
        
        return results

class TextStatisticsAnalyzer:
    """Comprehensive text statistics analyzer for detailed metrics"""
    
    def __init__(self):
        # Patterns for detection
        self.ly_adverb_pattern = r'\b\w+ly\b'
        self.gerund_the_pattern = r'\b\w+ing\s+the\b'
        self.gerund_of_pattern = r'\b\w+ing\s+of\b'
        self.indefinite_article_pattern = r'\b(a|an|A|An)\b'
        
        # Patterns for figures and tables
        self.figure_patterns = [
            r'Figure\s+\d+',
            r'Figures\s+\d+\s+and\s+\d+',
            r'Figures\s+\d+[-–]\d+',
            r'Figure\s+\d+[A-Z]?',
            r'Figs\.?\s+\d+'
        ]
        
        self.table_patterns = [
            r'Table\s+\d+',
            r'Tables\s+\d+\s+and\s+\d+',
            r'Tables\s+\d+[-–]\d+',
            r'Table\s+\d+[A-Z]?',
            r'Tables?\s+\d+'
        ]
        
        self.supplementary_patterns = [
            r'supplementary materials?',
            r'supplementary information',
            r'electronic supplementary materials?',
            r'electronic supplementary information',
            r'Figure\s+S\d+',
            r'Figure\s+\d+S',
            r'Table\s+S\d+',
            r'Table\s+\d+S',
            r'Figures?\s+S\d+[-–]S\d+',
            r'Tables?\s+S\d+[-–]S\d+',
            r'Figures?\s+\d+S\s+and\s+\d+S',
            r'Tables?\s+\d+S\s+and\s+\d+S'
        ]
    
    def split_paragraphs(self, text: str) -> List[str]:
        """Split text into paragraphs"""
        # Split by double newlines
        paragraphs = re.split(r'\n\s*\n', text)
        # Filter out empty paragraphs
        return [p.strip() for p in paragraphs if p.strip()]
    
    def analyze(self, text: str, sentences: List[str]) -> Dict:
        """Analyze comprehensive text statistics"""
        results = {
            'sentence_lengths': [],
            'paragraph_lengths': [],
            'commas_per_sentence': [],
            'commas_per_paragraph': [],
            'apostrophes_per_sentence': [],
            'apostrophes_per_paragraph': [],
            'ly_adverbs_per_sentence': [],
            'ly_adverbs_per_paragraph': [],
            'gerund_the_per_sentence': [],
            'gerund_the_per_paragraph': [],
            'gerund_of_per_sentence': [],
            'gerund_of_per_paragraph': [],
            'indefinite_articles_per_sentence': [],
            'indefinite_articles_per_paragraph': [],
            'figure_mentions': [],
            'table_mentions': [],
            'supplementary_mentions': [],
            'statistics': {
                'sentence_length': {},
                'paragraph_length': {},
                'commas_per_sentence': {},
                'commas_per_paragraph': {},
                'apostrophes_per_sentence': {},
                'apostrophes_per_paragraph': {},
                'ly_adverbs_per_sentence': {},
                'ly_adverbs_per_paragraph': {},
                'gerund_the_per_sentence': {},
                'gerund_the_per_paragraph': {},
                'gerund_of_per_sentence': {},
                'gerund_of_per_paragraph': {},
                'indefinite_articles_per_sentence': {},
                'indefinite_articles_per_paragraph': {}
            }
        }
        
        if not text or not sentences:
            return results
        
        # Split into paragraphs
        paragraphs = self.split_paragraphs(text)
        
        # ===== SENTENCE-LEVEL STATISTICS =====
        for sentence in sentences:
            if not sentence or len(sentence.strip()) < 2:
                continue
            
            # Sentence length in words
            words = sentence.split()
            sent_length = len(words)
            results['sentence_lengths'].append(sent_length)
            
            # Commas count
            comma_count = sentence.count(',')
            results['commas_per_sentence'].append(comma_count)
            
            # Apostrophes count
            apostrophe_count = len(re.findall(r"'", sentence))
            results['apostrophes_per_sentence'].append(apostrophe_count)
            
            # -ly adverbs
            ly_adverbs = len(re.findall(self.ly_adverb_pattern, sentence, re.IGNORECASE))
            results['ly_adverbs_per_sentence'].append(ly_adverbs)
            
            # Gerund + the
            gerund_the = len(re.findall(self.gerund_the_pattern, sentence, re.IGNORECASE))
            results['gerund_the_per_sentence'].append(gerund_the)
            
            # Gerund + of
            gerund_of = len(re.findall(self.gerund_of_pattern, sentence, re.IGNORECASE))
            results['gerund_of_per_sentence'].append(gerund_of)
            
            # Indefinite articles
            articles = len(re.findall(self.indefinite_article_pattern, sentence))
            results['indefinite_articles_per_sentence'].append(articles)
            
            # Figure mentions with full sentence
            for pattern in self.figure_patterns:
                matches = re.findall(pattern, sentence, re.IGNORECASE)
                for match in matches:
                    results['figure_mentions'].append({
                        'sentence': sentence.strip(),
                        'match': match
                    })
            
            # Table mentions with full sentence
            for pattern in self.table_patterns:
                matches = re.findall(pattern, sentence, re.IGNORECASE)
                for match in matches:
                    results['table_mentions'].append({
                        'sentence': sentence.strip(),
                        'match': match
                    })
            
            # Supplementary mentions with full sentence
            for pattern in self.supplementary_patterns:
                matches = re.findall(pattern, sentence, re.IGNORECASE)
                for match in matches:
                    results['supplementary_mentions'].append({
                        'sentence': sentence.strip(),
                        'match': match
                    })
        
        # ===== PARAGRAPH-LEVEL STATISTICS =====
        for paragraph in paragraphs:
            if not paragraph:
                continue
            
            # Paragraph length in words
            para_words = paragraph.split()
            para_length = len(para_words)
            results['paragraph_lengths'].append(para_length)
            
            # Commas per paragraph
            comma_count = paragraph.count(',')
            results['commas_per_paragraph'].append(comma_count)
            
            # Apostrophes per paragraph
            apostrophe_count = len(re.findall(r"'", paragraph))
            results['apostrophes_per_paragraph'].append(apostrophe_count)
            
            # -ly adverbs per paragraph
            ly_adverbs = len(re.findall(self.ly_adverb_pattern, paragraph, re.IGNORECASE))
            results['ly_adverbs_per_paragraph'].append(ly_adverbs)
            
            # Gerund + the per paragraph
            gerund_the = len(re.findall(self.gerund_the_pattern, paragraph, re.IGNORECASE))
            results['gerund_the_per_paragraph'].append(gerund_the)
            
            # Gerund + of per paragraph
            gerund_of = len(re.findall(self.gerund_of_pattern, paragraph, re.IGNORECASE))
            results['gerund_of_per_paragraph'].append(gerund_of)
            
            # Indefinite articles per paragraph
            articles = len(re.findall(self.indefinite_article_pattern, paragraph))
            results['indefinite_articles_per_paragraph'].append(articles)
        
        # ===== CALCULATE STATISTICS =====
        def calculate_stats(values):
            if not values:
                return {'min': 0, 'max': 0, 'mean': 0, 'median': 0}
            return {
                'min': float(np.min(values)),
                'max': float(np.max(values)),
                'mean': float(np.mean(values)),
                'median': float(np.median(values))
            }
        
        # Sentence-level stats
        results['statistics']['sentence_length'] = calculate_stats(results['sentence_lengths'])
        results['statistics']['commas_per_sentence'] = calculate_stats(results['commas_per_sentence'])
        results['statistics']['apostrophes_per_sentence'] = calculate_stats(results['apostrophes_per_sentence'])
        results['statistics']['ly_adverbs_per_sentence'] = calculate_stats(results['ly_adverbs_per_sentence'])
        results['statistics']['gerund_the_per_sentence'] = calculate_stats(results['gerund_the_per_sentence'])
        results['statistics']['gerund_of_per_sentence'] = calculate_stats(results['gerund_of_per_sentence'])
        results['statistics']['indefinite_articles_per_sentence'] = calculate_stats(results['indefinite_articles_per_sentence'])
        
        # Paragraph-level stats
        results['statistics']['paragraph_length'] = calculate_stats(results['paragraph_lengths'])
        results['statistics']['commas_per_paragraph'] = calculate_stats(results['commas_per_paragraph'])
        results['statistics']['apostrophes_per_paragraph'] = calculate_stats(results['apostrophes_per_paragraph'])
        results['statistics']['ly_adverbs_per_paragraph'] = calculate_stats(results['ly_adverbs_per_paragraph'])
        results['statistics']['gerund_the_per_paragraph'] = calculate_stats(results['gerund_the_per_paragraph'])
        results['statistics']['gerund_of_per_paragraph'] = calculate_stats(results['gerund_of_per_paragraph'])
        results['statistics']['indefinite_articles_per_paragraph'] = calculate_stats(results['indefinite_articles_per_paragraph'])
        
        # Additional counts for figures/tables/supplementary
        results['figure_count'] = len(results['figure_mentions'])
        results['table_count'] = len(results['table_mentions'])
        results['supplementary_count'] = len(results['supplementary_mentions'])
        
        return results


class IntegratedRiskScorer:
    """Integrated risk assessment based on all modules"""
    
    def __init__(self):
        # Module weights (updated with tortured_phrases - weight 0.10)
        self.weights = {
            'unicode': 0.25,
            'dashes': 0.10,          # немного уменьшили, чтобы освободить место
            'phrases': 0.06,
            'tortured_phrases': 0.10,  # НОВЫЙ модуль с весом 10%
            'burstiness': 0.04,
            'grammar': 0.05,
            'hedging': 0.06,
            'paragraph': 0.03,
            'perplexity': 0.02,
            'semantic': 0.02,
            'parenthesis': 0.03,
            'punctuation': 0.03,
            'apostrophe': 0.07,
            'enumeration': 0.06,
            'repetitiveness': 0.03,
            'lexical_diversity': 0.02,
            'log_prob': 0.01,
            'ml_classifier': 0.02
        }
               
        # Normalize weights
        total = sum(self.weights.values())
        if total > 0:
            for key in self.weights:
                self.weights[key] /= total
    
    def calculate(self, results: Dict) -> Dict:
        """Calculate integrated risk with improved weighting"""
        total_score = 0
        total_confidence = 0
        weighted_score = 0
        module_scores = []
        
        # Проверяем, какие модули есть в результатах
        available_modules = []
        for module, weight in self.weights.items():
            if module in results and results[module]:
                available_modules.append(module)
        
        # Перенормализуем веса только для доступных модулей
        if available_modules:
            # Сумма весов доступных модулей
            total_available_weight = sum(self.weights[m] for m in available_modules)
            
            for module in available_modules:
                data = results[module]
                if 'risk_score' in data and data['risk_score'] is not None:
                    # Нормализуем вес относительно доступных модулей
                    normalized_weight = self.weights[module] / total_available_weight
                    
                    # Нормализуем risk_score с учетом увеличенного максимума (теперь до 6)
                    max_score = 6  # БЫЛО 3, СТАЛО 6 (учет увеличенных значений)
                    norm_score = min(data['risk_score'] / max_score, 1.0)
                    
                    # Для модулей с очень высоким риском - усиливаем сигнал
                    if data['risk_score'] >= 5:
                        norm_score = min(norm_score * 1.2, 1.0)  # буст для очень сильных сигналов
                    
                    # Учитываем confidence модуля
                    confidence = data.get('confidence', 0.5)
                    
                    # Для модулей, где низкий риск = человеческий, инвертируем
                    invert_modules = ['parenthesis', 'punctuation']
                    if module in invert_modules and 'risk_score' in data:
                        norm_score = 1.0 - norm_score
                    
                    # Вклад модуля в общий score
                    contribution = norm_score * normalized_weight * 100
                    
                    module_score = {
                        'module': module,
                        'raw_score': data['risk_score'],
                        'norm_score': norm_score,
                        'weight': normalized_weight,
                        'original_weight': self.weights[module],
                        'confidence': confidence,
                        'contribution': contribution
                    }
                    
                    module_scores.append(module_score)
                    weighted_score += norm_score * normalized_weight
                    total_confidence += confidence * normalized_weight
        
        # Финальный score 0-100 - УВЕЛИЧИВАЕМ ЧУВСТВИТЕЛЬНОСТЬ
        final_score = weighted_score * 100
        
        # Усиливаем влияние высоких скоров (нелинейное усиление)
        if weighted_score > 0.5:
            final_score = final_score * 1.2  # буст для явных случаев
        elif weighted_score > 0.7:
            final_score = final_score * 1.3  # еще больше буста
        
        # Корректировка на уверенность
        if total_confidence > 0:
            final_score = final_score * (0.5 + 0.5 * total_confidence)
        
        # Ограничиваем максимум 100
        final_score = min(final_score, 100)
        
        # Определяем уровень риска
        risk_level = 'unknown'
        if final_score < 15:
            risk_level = 'very_low'
        elif final_score < 20:
            risk_level = 'low'
        elif final_score < 30:
            risk_level = 'medium-low'
        elif final_score < 40:
            risk_level = 'medium'
        elif final_score < 50:
            risk_level = 'medium-high'
        else:
            risk_level = 'high'
        
        return {
            'final_score': final_score,
            'risk_level': risk_level,
            'weighted_score': weighted_score,
            'total_confidence': total_confidence,
            'module_scores': module_scores,
            'available_modules': available_modules
        }
        
class DocumentProcessor:
    """Uploaded document processor"""
    
    @staticmethod
    def read_docx(file) -> Optional[str]:
        """Read .docx file"""
        try:
            from docx import Document
            doc = Document(file)
            full_text = []
            for para in doc.paragraphs:
                if para.text:
                    full_text.append(para.text)
            
            # Read tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if para.text:
                                full_text.append(para.text)
            
            return '\n'.join(full_text)
        except Exception as e:
            st.error(f"Error reading DOCX: {e}")
            return None
    
    @staticmethod
    def read_doc(file) -> Optional[str]:
        """Read .doc file (if antiword available)"""
        try:
            # Save to temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix='.doc', mode='wb') as tmp:
                tmp.write(file.getvalue())
                tmp_path = tmp.name
            
            # Try using antiword
            try:
                result = subprocess.run(['antiword', tmp_path], 
                                       capture_output=True, text=True, timeout=10)
                os.unlink(tmp_path)
                if result.returncode == 0 and result.stdout:
                    return result.stdout
            except:
                pass
            
            # If antiword fails, try text fallback
            os.unlink(tmp_path)
            
            # Try reading as text file
            file.seek(0)
            content = file.getvalue().decode('utf-8', errors='ignore')
            if len(content) > 100:
                return content
            
            return None
        except Exception as e:
            return None
    
    @staticmethod
    def split_sentences_simple(text: str) -> List[str]:
        """
        Improved sentence segmentation with handling of abbreviations,
        numbers, and scientific notation.
        """
        if not text:
            return []
        
        # Common abbreviations that should NOT trigger sentence breaks
        abbreviations = {
            # Academic titles
            'Dr', 'Mr', 'Mrs', 'Ms', 'Prof', 'Rev', 'Hon', 'St', 'Sr', 'Jr',
            # Latin abbreviations
            'i.e', 'e.g', 'vs', 'et al', 'etc', 'cf', 'op cit', 'loc cit',
            # Units and measures
            'Fig', 'Figs', 'Table', 'Tables', 'Eq', 'Eqs', 'Sec', 'Secs', 'Ch', 'Chs',
            'Vol', 'Vols', 'No', 'Nos', 'p', 'pp', 'ed', 'eds', 'trans',
            # Time
            'a.m', 'p.m', 'AM', 'PM',
            # Scientific abbreviations (often in caps)
            'PCC', 'PCFC', 'PCEC', 'SOFC', 'SOEC', 'ORR', 'OER', 'TCO', 'MIEC',
            'TPB', 'DPB', 'YSZ', 'GDC', 'LSCF', 'BSCF', 'BCZY', 'LSCM', 'SFM',
            # Common in your article
            'CN', 'US', 'AU', 'KR', 'JP', 'GT', 'DTU', 'ICL', 'NJTECH', 'CU'
        }
        
        # Build pattern for abbreviations that should NOT break sentences
        # Pattern matches abbreviation followed by dot and then optional space and capital letter
        # But we'll protect them from being split
        protected_pattern = r'\b(' + '|'.join(re.escape(abbr) for abbr in abbreviations) + r')\.'
        
        # Temporarily protect abbreviations by replacing them with placeholders
        protected = {}
        counter = 0
        
        def protect_abbrev(match):
            nonlocal counter
            placeholder = f"__ABBREV_{counter}__"
            protected[placeholder] = match.group(0)
            counter += 1
            return placeholder
        
        # First, protect known abbreviations
        text_protected = re.sub(protected_pattern, protect_abbrev, text)
        
        # Also protect numbers with dots (like "5.3%", "Fig. 3", "Section 3.2")
        number_pattern = r'\b\d+\.\d+\b'
        text_protected = re.sub(number_pattern, protect_abbrev, text_protected)
        
        # Now do sentence splitting
        # Pattern: look for .!? followed by space and capital letter
        # But don't split on .!? inside quotes or parentheses (handled by regex)
        sentences = re.split(r'(?<=[.!?])\s+(?=[A-Z])', text_protected)
        
        # Restore protected abbreviations
        restored_sentences = []
        for sent in sentences:
            for placeholder, original in protected.items():
                sent = sent.replace(placeholder, original)
            # Clean up
            sent = sent.strip()
            if len(sent) > 10:  # Minimum sentence length
                restored_sentences.append(sent)
        
        # Additional pass: merge sentences that were split incorrectly
        # (e.g., after "et al." which wasn't in abbreviations list)
        merged_sentences = []
        skip_next = False
        
        for i, sent in enumerate(restored_sentences):
            if skip_next:
                skip_next = False
                continue
                
            # Check if this sentence is too short and might be part of previous
            if len(sent.split()) < 5 and i > 0:
                # Merge with previous
                merged_sentences[-1] = merged_sentences[-1] + " " + sent
            else:
                merged_sentences.append(sent)
        
        return merged_sentences
    
    @staticmethod
    def preprocess(text: str) -> str:
        """Basic text preprocessing"""
        if not text:
            return ""
        # Normalize spaces
        text = re.sub(r'\s+', ' ', text)
        # Remove extra line breaks
        text = re.sub(r'\n\s*\n', '\n\n', text)
        return text.strip()

def format_authors(authors_list):
    """Форматирование списка авторов для PDF"""
    if not authors_list:
        return "N/A"
    
    # Если авторы уже в виде строки
    if isinstance(authors_list, str):
        return authors_list
    
    # Если авторы в виде списка словарей или строк
    formatted = []
    for author in authors_list[:5]:
        if isinstance(author, dict):
            # Если это словарь с ключом 'name'
            name = author.get('name', '')
            if name:
                # Форматируем имя: Фамилия И.О.
                parts = name.split()
                if len(parts) >= 2:
                    last_name = parts[-1]
                    initials = ' '.join([p[0] + '.' for p in parts[:-1]])
                    formatted.append(f"{last_name} {initials}")
                else:
                    formatted.append(name)
        elif isinstance(author, str):
            formatted.append(author)
    
    result = ', '.join(formatted)
    if len(authors_list) > 5:
        result += f' et al. ({len(authors_list)} authors)'
    
    return result

def clean_text_for_pdf(text):
    """Очистка текста для PDF от HTML тегов и спецсимволов"""
    if not text:
        return ""
    if not isinstance(text, str):
        text = str(text)
    # Удаляем HTML теги
    text = re.sub(r'<[^>]+>', '', text)
    # Заменяем специальные символы
    text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
    # Убираем множественные пробелы
    text = re.sub(r'\s+', ' ', text)
    return text.strip()

def generate_pdf_report(results_data, topic_name="CT(A)I-detector Analysis"):
    """Генерация PDF отчета с результатами анализа"""
    
    buffer = io.BytesIO()
    
    # Настройка документа
    doc = SimpleDocTemplate(
        buffer, 
        pagesize=A4,
        topMargin=1*cm,
        bottomMargin=1*cm,
        leftMargin=1.5*cm,
        rightMargin=1.5*cm
    )
    
    styles = getSampleStyleSheet()
    
    # ========== СОЗДАНИЕ КАСТОМНЫХ СТИЛЕЙ ==========
    
    # Стиль для заголовка
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=20,
        textColor=colors.HexColor('#2C3E50'),
        spaceAfter=12,
        alignment=TA_CENTER,
        fontName='Helvetica-Bold'
    )
    
    # Стиль для подзаголовка
    subtitle_style = ParagraphStyle(
        'CustomSubtitle',
        parent=styles['Heading2'],
        fontSize=16,
        textColor=colors.HexColor('#34495E'),
        spaceAfter=8,
        alignment=TA_CENTER,
        fontName='Helvetica'
    )
    
    # Стиль для мета-информации
    meta_style = ParagraphStyle(
        'CustomMeta',
        parent=styles['Normal'],
        fontSize=10,
        textColor=colors.HexColor('#7F8C8D'),
        spaceAfter=3,
        alignment=TA_CENTER,
        fontName='Helvetica-Oblique'
    )
    
    # Стиль для секций
    section_style = ParagraphStyle(
        'CustomSection',
        parent=styles['Heading3'],
        fontSize=14,
        textColor=colors.HexColor('#2980B9'),
        spaceAfter=10,
        spaceBefore=15,
        alignment=TA_LEFT,
        fontName='Helvetica-Bold'
    )
    
    # Стиль для подсекций
    subsection_style = ParagraphStyle(
        'CustomSubsection',
        parent=styles['Heading4'],
        fontSize=12,
        textColor=colors.HexColor('#16A085'),
        spaceAfter=6,
        alignment=TA_LEFT,
        fontName='Helvetica-Bold'
    )
    
    # Стиль для обычного текста
    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontSize=10,
        textColor=colors.HexColor('#2C3E50'),
        spaceAfter=4,
        alignment=TA_LEFT,
        fontName='Helvetica'
    )
    
    # Стиль для метрик
    metrics_style = ParagraphStyle(
        'CustomMetrics',
        parent=styles['Normal'],
        fontSize=11,
        textColor=colors.HexColor('#27AE60'),
        spaceAfter=2,
        alignment=TA_LEFT,
        fontName='Helvetica-Bold'
    )
    
    # Стиль для примеров
    example_style = ParagraphStyle(
        'CustomExample',
        parent=styles['Normal'],
        fontSize=9,
        textColor=colors.HexColor('#7F8C8D'),
        spaceAfter=2,
        leftIndent=20,
        alignment=TA_LEFT,
        fontName='Helvetica-Oblique'
    )
    
    # Стиль для нижнего колонтитула
    footer_style = ParagraphStyle(
        'CustomFooter',
        parent=styles['Normal'],
        fontSize=8,
        textColor=colors.HexColor('#95A5A6'),
        spaceBefore=15,
        alignment=TA_CENTER,
        fontName='Helvetica-Oblique'
    )
    
    story = []
    
    # ========== ТИТУЛЬНАЯ СТРАНИЦА ==========
    
    story.append(Spacer(1, 2*cm))
    
    # Добавляем логотип
    try:
        # Пробуем несколько возможных путей
        possible_paths = [
            "logo.png",  # Текущая директория
            "./logo.png",  # Относительный путь
            "app/logo.png",  # Если в поддиректории
            os.path.join(os.path.dirname(__file__), "logo.png"),  # Абсолютный путь
            os.path.join(os.getcwd(), "logo.png")  # Текущая рабочая директория
        ]
        
        logo_path = None
        for path in possible_paths:
            if os.path.exists(path):
                logo_path = path
                break
        
        if logo_path:
            # Используем Image из reportlab
            logo = Image(logo_path, width=270, height=135)
            logo.hAlign = 'CENTER'
            story.append(logo)
            story.append(Spacer(1, 1*cm))
    except Exception as e:
        # Если логотип не загрузился, показываем эмодзи
        story.append(Paragraph("🔬", ParagraphStyle(
            'LogoEmoji',
            parent=styles['Heading1'],
            fontSize=40,
            textColor=colors.HexColor('#667eea'),
            alignment=TA_CENTER
        )))
        story.append(Spacer(1, 0.5*cm))
    
    # Заголовок
    story.append(Paragraph("Advanced AI Text Analysis Report", subtitle_style))
    story.append(Spacer(1, 1*cm))
    
    # Мета-информация
    current_date = datetime.now().strftime('%B %d, %Y at %H:%M')
    story.append(Paragraph(f"Generated on {current_date}", meta_style))
    
    if results_data:
        text = results_data.get('text', '')
        sentences = results_data.get('sentences', [])
        integrated = results_data.get('integrated', {})
        
        story.append(Paragraph(f"Total words analyzed: {len(text.split())}", meta_style))
        story.append(Paragraph(f"Total sentences: {len(sentences)}", meta_style))
        story.append(Paragraph(f"Risk score: {integrated.get('final_score', 0):.1f}/100", meta_style))
        story.append(Paragraph(f"Risk level: {integrated.get('risk_level', 'unknown').replace('_', ' ').title()}", meta_style))
    
    story.append(Spacer(1, 3*cm))
    story.append(Paragraph("© CT(A)I-detector", footer_style))
    story.append(Paragraph("https://chimicatechnoacta.ru developed by @daM", footer_style))
    
    story.append(PageBreak())
    
    # ========== ОСНОВНОЙ ОТЧЕТ ==========
    
    if not results_data:
        story.append(Paragraph("No data available for analysis", normal_style))
    else:
        text = results_data.get('text', '')
        sentences = results_data.get('sentences', [])
        results = results_data.get('results', {})
        integrated = results_data.get('integrated', {})
        text_stats = results_data.get('text_statistics', {})
        
        # 1. ОБЩИЙ СЧЕТ
        story.append(Paragraph("1. OVERALL RISK ASSESSMENT", section_style))
        
        # Основная метрика
        final_score = integrated.get('final_score', 0)
        risk_level = integrated.get('risk_level', 'unknown').replace('_', ' ').title()
        
        # Определяем цвет на основе риска
        if final_score < 30:
            risk_color = '#27AE60'  # зеленый
        elif final_score < 50:
            risk_color = '#F39C12'  # оранжевый
        elif final_score < 70:
            risk_color = '#E67E22'  # темно-оранжевый
        else:
            risk_color = '#E74C3C'  # красный
        
        story.append(Paragraph(f"AI Risk Score: {final_score:.1f}/100", 
                              ParagraphStyle('Score', parent=metrics_style, fontSize=16, textColor=colors.HexColor(risk_color))))
        story.append(Paragraph(f"Risk Level: {risk_level}", metrics_style))
        story.append(Spacer(1, 0.5*cm))
        
        # Статистика текста
        story.append(Paragraph("Text Statistics:", subsection_style))
        
        stats_data = [
            ["Metric", "Value"],
            ["Characters", len(text)],
            ["Words", len(text.split())],
            ["Sentences", len(sentences)],
            ["Avg. sentence length", f"{len(text.split()) / max(len(sentences), 1):.1f} words"],
            ["Modules analyzed", len(integrated.get('module_scores', []))]
        ]
        
        stats_table = Table(stats_data, colWidths=[doc.width/2.5, doc.width/3])
        stats_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#3498DB')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 11),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#F8F9FA')),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#D5DBDB')),
            ('FONTSIZE', (0, 1), (-1, -1), 10),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ]))
        
        story.append(stats_table)
        story.append(Spacer(1, 0.5*cm))
        
        # 1.5 ДЕТАЛЬНАЯ СТАТИСТИКА (NEW!)
        if text_stats:
            story.append(Paragraph("1.5 Detailed Text Statistics", section_style))
            
            # Sentence length statistics
            story.append(Paragraph("Sentence Length (words):", subsection_style))
            sent_stats = text_stats.get('statistics', {}).get('sentence_length', {})
            sent_data = [
                ["Min", f"{sent_stats.get('min', 0):.1f}"],
                ["Max", f"{sent_stats.get('max', 0):.1f}"],
                ["Mean", f"{sent_stats.get('mean', 0):.1f}"],
                ["Median", f"{sent_stats.get('median', 0):.1f}"]
            ]
            sent_table = Table(sent_data, colWidths=[doc.width/4, doc.width/4])
            sent_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#9B59B6')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#D5DBDB')),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
            ]))
            story.append(sent_table)
            story.append(Spacer(1, 0.3*cm))
            
            # Paragraph length statistics
            story.append(Paragraph("Paragraph Length (words):", subsection_style))
            para_stats = text_stats.get('statistics', {}).get('paragraph_length', {})
            para_data = [
                ["Min", f"{para_stats.get('min', 0):.1f}"],
                ["Max", f"{para_stats.get('max', 0):.1f}"],
                ["Mean", f"{para_stats.get('mean', 0):.1f}"],
                ["Median", f"{para_stats.get('median', 0):.1f}"]
            ]
            para_table = Table(para_data, colWidths=[doc.width/4, doc.width/4])
            para_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#9B59B6')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#D5DBDB')),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
            ]))
            story.append(para_table)
            story.append(Spacer(1, 0.3*cm))
            
            # Commas per sentence
            story.append(Paragraph("Commas per Sentence:", subsection_style))
            comma_stats = text_stats.get('statistics', {}).get('commas_per_sentence', {})
            comma_data = [
                ["Min", f"{comma_stats.get('min', 0):.1f}"],
                ["Max", f"{comma_stats.get('max', 0):.1f}"],
                ["Mean", f"{comma_stats.get('mean', 0):.1f}"],
                ["Median", f"{comma_stats.get('median', 0):.1f}"]
            ]
            comma_table = Table(comma_data, colWidths=[doc.width/4, doc.width/4])
            comma_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#E67E22')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#D5DBDB')),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
            ]))
            story.append(comma_table)
            story.append(Spacer(1, 0.3*cm))
            
            # Apostrophes per sentence
            story.append(Paragraph("Apostrophes per Sentence:", subsection_style))
            apost_stats = text_stats.get('statistics', {}).get('apostrophes_per_sentence', {})
            apost_data = [
                ["Min", f"{apost_stats.get('min', 0):.1f}"],
                ["Max", f"{apost_stats.get('max', 0):.1f}"],
                ["Mean", f"{apost_stats.get('mean', 0):.1f}"],
                ["Median", f"{apost_stats.get('median', 0):.1f}"]
            ]
            apost_table = Table(apost_data, colWidths=[doc.width/4, doc.width/4])
            apost_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#E74C3C')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#D5DBDB')),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
            ]))
            story.append(apost_table)
            story.append(Spacer(1, 0.3*cm))
            
            # -ly adverbs per sentence
            story.append(Paragraph("-ly Adverbs per Sentence:", subsection_style))
            ly_stats = text_stats.get('statistics', {}).get('ly_adverbs_per_sentence', {})
            ly_data = [
                ["Min", f"{ly_stats.get('min', 0):.1f}"],
                ["Max", f"{ly_stats.get('max', 0):.1f}"],
                ["Mean", f"{ly_stats.get('mean', 0):.1f}"],
                ["Median", f"{ly_stats.get('median', 0):.1f}"]
            ]
            ly_table = Table(ly_data, colWidths=[doc.width/4, doc.width/4])
            ly_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#3498DB')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#D5DBDB')),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
            ]))
            story.append(ly_table)
            story.append(Spacer(1, 0.3*cm))
            
            # Gerund + the per sentence
            story.append(Paragraph("Gerund + 'the' per Sentence:", subsection_style))
            gerund_the_stats = text_stats.get('statistics', {}).get('gerund_the_per_sentence', {})
            gerund_the_data = [
                ["Min", f"{gerund_the_stats.get('min', 0):.1f}"],
                ["Max", f"{gerund_the_stats.get('max', 0):.1f}"],
                ["Mean", f"{gerund_the_stats.get('mean', 0):.1f}"],
                ["Median", f"{gerund_the_stats.get('median', 0):.1f}"]
            ]
            gerund_the_table = Table(gerund_the_data, colWidths=[doc.width/4, doc.width/4])
            gerund_the_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#16A085')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#D5DBDB')),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
            ]))
            story.append(gerund_the_table)
            story.append(Spacer(1, 0.3*cm))
            
            # Gerund + of per sentence
            story.append(Paragraph("Gerund + 'of' per Sentence:", subsection_style))
            gerund_of_stats = text_stats.get('statistics', {}).get('gerund_of_per_sentence', {})
            gerund_of_data = [
                ["Min", f"{gerund_of_stats.get('min', 0):.1f}"],
                ["Max", f"{gerund_of_stats.get('max', 0):.1f}"],
                ["Mean", f"{gerund_of_stats.get('mean', 0):.1f}"],
                ["Median", f"{gerund_of_stats.get('median', 0):.1f}"]
            ]
            gerund_of_table = Table(gerund_of_data, colWidths=[doc.width/4, doc.width/4])
            gerund_of_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#F39C12')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#D5DBDB')),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
            ]))
            story.append(gerund_of_table)
            story.append(Spacer(1, 0.3*cm))
            
            # Indefinite articles per sentence
            story.append(Paragraph("Indefinite Articles (a/an) per Sentence:", subsection_style))
            article_stats = text_stats.get('statistics', {}).get('indefinite_articles_per_sentence', {})
            article_data = [
                ["Min", f"{article_stats.get('min', 0):.1f}"],
                ["Max", f"{article_stats.get('max', 0):.1f}"],
                ["Mean", f"{article_stats.get('mean', 0):.1f}"],
                ["Median", f"{article_stats.get('median', 0):.1f}"]
            ]
            article_table = Table(article_data, colWidths=[doc.width/4, doc.width/4])
            article_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#8E44AD')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#D5DBDB')),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
            ]))
            story.append(article_table)
            story.append(Spacer(1, 0.5*cm))
            
            # Figure mentions
            if text_stats.get('figure_mentions'):
                story.append(Paragraph(f"Figure Mentions (Total: {text_stats.get('figure_count', 0)}):", subsection_style))
                for mention in text_stats['figure_mentions'][:10]:
                    clean_sent = clean_text_for_pdf(mention['sentence'])[:200]
                    story.append(Paragraph(f"• {clean_sent}...", example_style))
                if len(text_stats['figure_mentions']) > 10:
                    story.append(Paragraph(f"... and {len(text_stats['figure_mentions']) - 10} more mentions", example_style))
                story.append(Spacer(1, 0.3*cm))
            
            # Table mentions
            if text_stats.get('table_mentions'):
                story.append(Paragraph(f"Table Mentions (Total: {text_stats.get('table_count', 0)}):", subsection_style))
                for mention in text_stats['table_mentions'][:10]:
                    clean_sent = clean_text_for_pdf(mention['sentence'])[:200]
                    story.append(Paragraph(f"• {clean_sent}...", example_style))
                if len(text_stats['table_mentions']) > 10:
                    story.append(Paragraph(f"... and {len(text_stats['table_mentions']) - 10} more mentions", example_style))
                story.append(Spacer(1, 0.3*cm))
            
            # Supplementary mentions
            if text_stats.get('supplementary_mentions'):
                story.append(Paragraph(f"Supplementary Mentions (Total: {text_stats.get('supplementary_count', 0)}):", subsection_style))
                for mention in text_stats['supplementary_mentions'][:10]:
                    clean_sent = clean_text_for_pdf(mention['sentence'])[:200]
                    story.append(Paragraph(f"• {clean_sent}...", example_style))
                if len(text_stats['supplementary_mentions']) > 10:
                    story.append(Paragraph(f"... and {len(text_stats['supplementary_mentions']) - 10} more mentions", example_style))
                story.append(Spacer(1, 0.3*cm))
        
        story.append(PageBreak())
        
        # 2. МОДУЛИ АНАЛИЗА
        story.append(Paragraph("2. MODULE CONTRIBUTIONS", section_style))
        
        module_scores = integrated.get('module_scores', [])
        if module_scores:
            # Сортируем по вкладу
            sorted_modules = sorted(module_scores, key=lambda x: x.get('contribution', 0), reverse=True)
            
            module_data = [["Module", "Raw Score", "Contribution", "Confidence"]]
            for ms in sorted_modules[:10]:  # Топ-10 модулей
                module_data.append([
                    ms.get('module', 'unknown').replace('_', ' ').title(),
                    f"{ms.get('raw_score', 0)}/6",
                    f"{ms.get('contribution', 0):.1f}%",
                    f"{ms.get('confidence', 0):.2f}"
                ])
            
            module_table = Table(module_data, colWidths=[doc.width/3, doc.width/6, doc.width/6, doc.width/6])
            module_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#9B59B6')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 11),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#D5DBDB')),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F2F4F4')]),
            ]))
            
            story.append(module_table)
        else:
            story.append(Paragraph("No module data available", normal_style))
        
        story.append(Spacer(1, 0.5*cm))
        
        # 3. ДЕТАЛЬНЫЙ АНАЛИЗ ПО МОДУЛЯМ
        story.append(Paragraph("3. DETAILED MODULE ANALYSIS", section_style))
        
        # Unicode артефакты
        if 'unicode' in results:
            story.append(Paragraph("3.1 Unicode Artifacts", subsection_style))
            unicode_data = results['unicode']
            story.append(Paragraph(f"• Suspicious characters: {unicode_data.get('sup_sub_count', 0) + unicode_data.get('fullwidth_count', 0)}", normal_style))
            story.append(Paragraph(f"• Density per 10k chars: {unicode_data.get('density_per_10k', 0):.2f}", normal_style))
            story.append(Paragraph(f"• Risk score: {unicode_data.get('risk_score', 0)}/6", normal_style))
            
            # Примеры
            chunks = unicode_data.get('all_suspicious_chunks', [])
            if chunks:
                story.append(Paragraph("Examples:", example_style))
                for chunk in chunks[:50]:
                    context = clean_text_for_pdf(chunk.get('context', ''))[:200]
                    story.append(Paragraph(f"  • '{chunk.get('char', '')}' → ...{context}...", example_style))
            story.append(Spacer(1, 0.3*cm))
        
        # AI фразы
        if 'phrases' in results:
            story.append(Paragraph("3.2 AI Phrases", subsection_style))
            phrases_data = results['phrases']
            occurrences = phrases_data.get('all_phrase_occurrences', [])
            story.append(Paragraph(f"• Total occurrences: {len(occurrences)}", normal_style))
            story.append(Paragraph(f"• Top phrases: {', '.join([p[0] for p in phrases_data.get('top_phrases', [])[:50]])}", normal_style))
            
            if occurrences:
                story.append(Paragraph("Examples:", example_style))
                for occ in occurrences[:50]:
                    context = clean_text_for_pdf(occ.get('context', ''))[:200]
                    story.append(Paragraph(f"  • '{occ.get('phrase', '')}' → ...{context}...", example_style))
            story.append(Spacer(1, 0.3*cm))
        
        # Tortured phrases
        if 'tortured_phrases' in results:
            story.append(Paragraph("3.3 Tortured Phrases", subsection_style))
            tp_data = results['tortured_phrases']
            occurrences = tp_data.get('all_occurrences', [])
            
            story.append(Paragraph(f"• Total occurrences: {tp_data.get('total_occurrences', 0)}", normal_style))
            story.append(Paragraph(f"• Unique tortured phrases: {tp_data.get('unique_tortured_count', 0)}", normal_style))
            story.append(Paragraph(f"• Frequency per 1000 words: {tp_data.get('tortured_per_1000_words', 0):.2f}", normal_style))
            
            if tp_data.get('tortured_phrases_found'):
                story.append(Paragraph("Tortured phrases found:", example_style))
                for phrase in tp_data['tortured_phrases_found']:
                    story.append(Paragraph(f"  • '{phrase['tortured']}' → correct: {phrase['correct']} (found {phrase['count']} times)", example_style))
            
            if occurrences:
                story.append(Spacer(1, 0.2*cm))
                story.append(Paragraph("Examples with context:", example_style))
                for occ in occurrences[:20]:
                    context = clean_text_for_pdf(occ.get('context', ''))[:200]
                    story.append(Paragraph(f"  • '{occ['tortured']}' → ...{context}...", example_style))
            story.append(Spacer(1, 0.3*cm))
        
        # Перечисления
        if 'enumeration' in results:
            story.append(Paragraph("3.4 Strict Enumerations", subsection_style))
            enum_data = results['enumeration']
            enumerations = enum_data.get('all_enumerations', [])
            story.append(Paragraph(f"• Found: {len(enumerations)} enumerations", normal_style))
            
            if enumerations:
                story.append(Paragraph("Examples:", example_style))
                for enum in enumerations[:5]:
                    story.append(Paragraph(f"  • {clean_text_for_pdf(enum)[:250]}...", example_style))
            story.append(Spacer(1, 0.3*cm))
        
        # Апострофы
        if 'apostrophe' in results:
            story.append(Paragraph("3.5 Apostrophe Usage", subsection_style))
            apost_data = results['apostrophe']
            apostrophes = apost_data.get('all_apostrophes', [])
            story.append(Paragraph(f"• Found: {len(apostrophes)} apostrophes", normal_style))
            story.append(Paragraph(f"• Per 1000 words: {apost_data.get('apostrophe_per_1000', 0):.2f}", normal_style))
            
            if apostrophes:
                examples = ', '.join(apostrophes[:100])
                story.append(Paragraph(f"Examples: {examples}", example_style))
            story.append(Spacer(1, 0.3*cm))
        
        # Пунктуация
        if 'punctuation' in results:
            story.append(Paragraph("3.6 Punctuation Analysis", subsection_style))
            punct_data = results['punctuation']
            story.append(Paragraph(f"• Exclamation marks: {punct_data.get('exclamation_count', 0)}", normal_style))
            story.append(Paragraph(f"• Question marks: {punct_data.get('question_count', 0)}", normal_style))
            story.append(Paragraph(f"• Semicolons: {punct_data.get('semicolon_count', 0)}", normal_style))
            
            semicolons = punct_data.get('all_semicolon_contexts', [])
            if semicolons:
                story.append(Paragraph("Semicolon contexts:", example_style))
                for ctx in semicolons[:3]:
                    context = clean_text_for_pdf(ctx.get('context', ''))[:100]
                    story.append(Paragraph(f"  • ...{context}...", example_style))
            story.append(Spacer(1, 0.3*cm))
        
        # Скобки
        if 'parenthesis' in results:
            story.append(Paragraph("3.7 Parenthesis Analysis", subsection_style))
            paren_data = results['parenthesis']
            parentheses = paren_data.get('all_parentheses', [])
            story.append(Paragraph(f"• Long parentheses (≥5 words): {paren_data.get('long_parentheses', 0)}", normal_style))
            story.append(Paragraph(f"• Percentage: {paren_data.get('long_percentage', 0):.1f}%", normal_style))
            
            if parentheses:
                story.append(Paragraph("Examples:", example_style))
                for p in parentheses[:50]:
                    text_content = clean_text_for_pdf(p.get('text', ''))[:100]
                    story.append(Paragraph(f"  • ({text_content}) — {p.get('word_count', 0)} words", example_style))
            story.append(Spacer(1, 0.3*cm))
        
        # Повторы
        if 'repetitiveness' in results:
            story.append(Paragraph("3.8 Repetitiveness Analysis", subsection_style))
            rep_data = results['repetitiveness']
            repetitions = rep_data.get('all_repetitions', [])
            story.append(Paragraph(f"• Found: {len(repetitions)} repeated phrases", normal_style))
            
            if repetitions:
                story.append(Paragraph("Top repetitions:", example_style))
                for rep in repetitions[:100]:
                    story.append(Paragraph(f"  • '{rep.get('ngram', '')}' — {rep.get('count', 0)} times", example_style))
            story.append(Spacer(1, 0.3*cm))
        
        # Тире
        if 'dashes' in results:
            story.append(Paragraph("3.9 Dash Analysis", subsection_style))
            dash_data = results['dashes']
            dashes = dash_data.get('all_dash_sentences', [])
            story.append(Paragraph(f"• Sentences with dashes: {len(dashes)}", normal_style))
            story.append(Paragraph(f"• Heavy sentences: {len(dash_data.get('heavy_sentences', []))}", normal_style))
            
            if dashes:
                story.append(Paragraph("Examples:", example_style))
                for d in dashes[:5]:
                    sentence = clean_text_for_pdf(d.get('sentence', ''))[:100]
                    story.append(Paragraph(f"  • Dashes: {d.get('dash_count', 0)} → ...{sentence}...", example_style))
            story.append(Spacer(1, 0.3*cm))
        
        # Лексическое разнообразие
        if 'lexical_diversity' in results:
            story.append(Paragraph("3.10 Lexical Diversity", subsection_style))
            lex_data = results['lexical_diversity']
            story.append(Paragraph(f"• TTR: {lex_data.get('ttr', 0):.3f}", normal_style))
            if lex_data.get('mtld', 0) > 0:
                story.append(Paragraph(f"• MTLD: {lex_data.get('mtld', 0):.1f}", normal_style))
            if lex_data.get('mattr', 0) > 0:
                story.append(Paragraph(f"• MATTR: {lex_data.get('mattr', 0):.3f}", normal_style))
            story.append(Paragraph(f"• Hapax ratio: {lex_data.get('hapax_ratio', 0):.3f}", normal_style))
        
        story.append(PageBreak())
        
        # 4. ПРИМЕРЫ ТЕКСТА
        story.append(Paragraph("4. TEXT SAMPLES", section_style))
        
        # Первые 5 предложений
        story.append(Paragraph("4.1 First 5 sentences:", subsection_style))
        for i, sent in enumerate(sentences[:5]):
            clean_sent = clean_text_for_pdf(sent)[:200]
            story.append(Paragraph(f"{i+1}. {clean_sent}...", example_style))
        
        story.append(Spacer(1, 0.5*cm))
        
        # Предложения с высоким риском
        high_risk_examples = []
        
        # Из модуля dashes
        if 'dashes' in results:
            for item in results['dashes'].get('heavy_sentences', [])[:30]:
                high_risk_examples.append(("Heavy dash usage", item.get('sentence', '')))
        
        # Из модуля enumerations
        if 'enumeration' in results:
            for enum in results['enumeration'].get('all_enumerations', [])[:30]:
                high_risk_examples.append(("Strict enumeration", enum))
        
        # Из модуля phrases
        if 'phrases' in results:
            for occ in results['phrases'].get('all_phrase_occurrences', [])[:30]:
                high_risk_examples.append(("AI phrase", occ.get('context', '')))
        
        if high_risk_examples:
            story.append(Paragraph("4.2 High-risk examples:", subsection_style))
            for i, (label, example) in enumerate(high_risk_examples[:50]):
                clean_example = clean_text_for_pdf(example)[:200]
                story.append(Paragraph(f"{label}: {clean_example}...", example_style))
        
        story.append(PageBreak())
        
        # 5. ЗАКЛЮЧЕНИЕ
        story.append(Paragraph("5. CONCLUSION & RECOMMENDATIONS", section_style))
        
        # Интерпретация результата
        story.append(Paragraph("Interpretation:", subsection_style))
        
        if final_score < 30:
            interpretation = [
                "• LOW RISK: Text shows patterns consistent with human writing.",
                "• High lexical diversity and natural variation in sentence structure.",
                "• Presence of human markers (varied punctuation, hedging, personal pronouns)."
            ]
        elif final_score < 50:
            interpretation = [
                "• MEDIUM-LOW RISK: Some AI-like patterns detected.",
                "• Moderate use of AI phrases and repetitive structures.",
                "• Consider reviewing specific flagged sections."
            ]
        elif final_score < 70:
            interpretation = [
                "• MEDIUM-HIGH RISK: Significant AI-like patterns detected.",
                "• High density of AI phrases and low lexical diversity.",
                "• Multiple modules show elevated risk scores."
            ]
        else:
            interpretation = [
                "• HIGH RISK: Text shows strong AI generation patterns.",
                "• Very low lexical diversity with high repetitiveness.",
                "• Many characteristic AI phrases and structures present."
            ]
        
        for line in interpretation:
            story.append(Paragraph(line, normal_style))
        
        story.append(Spacer(1, 0.5*cm))
        
        # Рекомендации
        story.append(Paragraph("Recommendations:", subsection_style))
        
        recommendations = [
            "• Review the flagged examples in each module section",
            "• Pay special attention to modules with raw scores ≥ 4/6",
            "• Consider the context of the text (academic vs. casual writing)",
            "• For research purposes, combine with other detection methods"
        ]
        
        for rec in recommendations:
            story.append(Paragraph(rec, normal_style))
        
        story.append(Spacer(1, 0.5*cm))
        
        # Техническая информация
        story.append(Paragraph("Technical Details:", subsection_style))
        story.append(Paragraph(f"• Analysis timestamp: {current_date}", example_style))
        story.append(Paragraph(f"• Modules with data: {', '.join([m.get('module', '') for m in module_scores[:8]])}", example_style))
        story.append(Paragraph(f"• Confidence score: {integrated.get('total_confidence', 0):.2f}", example_style))
    
    # Нижний колонтитул
    story.append(Spacer(1, 1*cm))
    story.append(Paragraph("─" * 70, footer_style))
    story.append(Paragraph("© CT(A)I-detector - Advanced AI Text Analysis Tool", footer_style))
    story.append(Paragraph(f"Report ID: {hashlib.md5(str(datetime.now()).encode()).hexdigest()[:8].upper()}", footer_style))
    
    # Генерация PDF
    doc.build(story)
    
    return buffer.getvalue()

# ============================================================================
# PDF Report Generation - Enhanced Version with Charts and TOC
# ============================================================================

def create_module_pie_chart(module_scores, width=400, height=200):
    """Create a pie/donut chart showing module contributions"""
    from reportlab.graphics.shapes import Drawing, String
    from reportlab.graphics.charts.piecharts import Pie
    
    if not module_scores:
        return None
    
    # Take top 6 modules by contribution, group others as "Others"
    sorted_modules = sorted(module_scores, key=lambda x: x.get('contribution', 0), reverse=True)
    
    if len(sorted_modules) > 6:
        top_modules = sorted_modules[:5]
        other_contribution = sum(m.get('contribution', 0) for m in sorted_modules[5:])
        if other_contribution > 0:
            top_modules.append({
                'module': 'others',
                'contribution': other_contribution,
                'raw_score': 0
            })
    else:
        top_modules = sorted_modules
    
    # Prepare data
    data = [m.get('contribution', 0) for m in top_modules]
    labels = [m.get('module', 'unknown').replace('_', ' ').title() for m in top_modules]
    
    # Create drawing
    drawing = Drawing(width, height)
    
    # Create pie chart
    pie = Pie()
    pie.x = 50
    pie.y = 30
    pie.width = 150
    pie.height = 150
    pie.data = data
    pie.labels = labels
    pie.sideLabels = True
    pie.simpleLabels = False
    
    # Colors
    colors_list = [
        colors.HexColor('#FF6B6B'),  # red
        colors.HexColor('#4ECDC4'),  # teal
        colors.HexColor('#45B7D1'),  # blue
        colors.HexColor('#96CEB4'),  # green
        colors.HexColor('#FFE194'),  # yellow
        colors.HexColor('#D4A5A5'),  # pink
    ]
    
    for i, color in enumerate(colors_list):
        if i < len(pie.slices):
            pie.slices[i].fillColor = color
    
    drawing.add(pie)
    
    # Add title
    title = String(200, 180, 'Module Contribution Distribution')
    title.fontName = 'Helvetica-Bold'
    title.fontSize = 10
    title.textAnchor = 'middle'
    drawing.add(title)
    
    return drawing

def create_module_bar_chart(module_scores, width=450, height=200):
    """Create a bar chart showing top modules by contribution"""
    from reportlab.graphics.shapes import Drawing
    from reportlab.graphics.charts.barcharts import VerticalBarChart
    
    if not module_scores:
        return None
    
    # Take top 8 modules
    top_modules = sorted(module_scores, key=lambda x: x.get('contribution', 0), reverse=True)[:8]
    
    # Prepare data
    contributions = [[m.get('contribution', 0) for m in top_modules]]
    raw_scores = [m.get('raw_score', 0) for m in top_modules]
    categories = [m.get('module', 'unknown').replace('_', ' ')[:12] for m in top_modules]
    
    # Create drawing
    drawing = Drawing(width, height)
    
    # Create bar chart
    bc = VerticalBarChart()
    bc.x = 50
    bc.y = 40
    bc.width = 350
    bc.height = 140
    bc.data = contributions
    bc.strokeColor = colors.black
    bc.strokeWidth = 0.5
    bc.categoryAxis.categoryNames = categories 
    bc.valueAxis.valueMin = 0
    bc.valueAxis.valueMax = max(max(contributions[0]) * 1.1, 30)
    bc.valueAxis.valueStep = 10
    bc.categoryAxis.labels.boxAnchor = 'ne'
    bc.categoryAxis.labels.dx = 8
    bc.categoryAxis.labels.dy = -2
    bc.categoryAxis.labels.angle = 45
    bc.bars[0].fillColor = colors.HexColor('#667eea')
    
    drawing.add(bc)
    
    # Add title
    from reportlab.graphics.shapes import String
    title = String(225, 180, 'Top Modules by Contribution')
    title.fontName = 'Helvetica-Bold'
    title.fontSize = 10
    title.textAnchor = 'middle'
    drawing.add(title)
    
    return drawing

def add_horizontal_rule(story, width=None, thickness=1, color=None):
    """Add a horizontal line separator with default color if None provided"""
    from reportlab.platypus import Flowable
    
    # Устанавливаем цвет по умолчанию, если передан None
    if color is None:
        color = colors.HexColor('#BDC3C7')
    
    class HRFlowable(Flowable):
        def __init__(self, width, thickness=1, color=colors.black):
            Flowable.__init__(self)
            self.width = width
            self.thickness = thickness
            self.color = color
        
        def draw(self):
            self.canv.setStrokeColor(self.color)
            self.canv.setLineWidth(self.thickness)
            self.canv.line(0, 0, self.width, 0)
        
        def wrap(self, availWidth, availHeight):
            if self.width is None:
                self.width = availWidth
            return (self.width, self.thickness)
    
    story.append(HRFlowable(width, thickness, color))
    story.append(Spacer(1, 0.3*cm))

def add_watermark(canvas, doc, text="CONFIDENTIAL", opacity=0.1):
    """Add watermark to every page"""
    canvas.saveState()
    canvas.setFillColor(colors.grey)
    canvas.setFillAlpha(opacity)
    canvas.setFont('Helvetica-Bold', 60)
    canvas.rotate(45)
    canvas.drawString(200, 100, text)
    canvas.restoreState()

def create_table_of_contents(story, doc):
    """Create clickable table of contents"""
    from reportlab.platypus import SimpleDocTemplate, Paragraph, PageBreak, Spacer
    from reportlab.lib import pagesizes
    
    story.append(PageBreak())
    story.append(Paragraph("TABLE OF CONTENTS", ParagraphStyle(
        'TOCHeading',
        parent=getSampleStyleSheet()['Heading1'],
        fontSize=16,
        textColor=colors.HexColor('#2C3E50'),
        spaceAfter=20,
        alignment=TA_CENTER,
        fontName='Helvetica-Bold'
    )))
    
    toc_entries = [
        ("1. OVERALL RISK ASSESSMENT", 1),
        ("2. MODULE CONTRIBUTIONS", 2),
        ("3. DETAILED MODULE ANALYSIS", 3),
        ("   3.1 Unicode Artifacts", 3),
        ("   3.2 AI Phrases", 3),
        ("   3.3 Tortured Phrases", 3),
        ("   3.4 Strict Enumerations", 3),
        ("   3.5 Apostrophe Usage", 3),
        ("   3.6 Punctuation Analysis", 3),
        ("   3.7 Parenthesis Analysis", 3),
        ("   3.8 Repetitiveness Analysis", 3),
        ("   3.9 Dash Analysis", 3),
        ("   3.10 Lexical Diversity", 3),
        ("4. TEXT SAMPLES", 4),
        ("5. CONCLUSION & RECOMMENDATIONS", 5),
    ]
    
    for title, page in toc_entries:
        # Create clickable link
        link_style = ParagraphStyle(
            'TOCEntry',
            parent=getSampleStyleSheet()['Normal'],
            fontSize=11,
            textColor=colors.HexColor('#3498DB'),
            leftIndent=20 if title.startswith('   ') else 0,
            fontName='Helvetica'
        )
        
        # Add dots between title and page number
        dots = '.' * (60 - len(title))
        p = Paragraph(f'<link href="#section{page}">{title}{dots}{page}</link>', link_style)
        story.append(p)
        story.append(Spacer(1, 0.1*cm))
    
    story.append(PageBreak())

def add_section_header(story, text, level=1, anchor=None):
    """Add section header with optional anchor for TOC"""
    styles = getSampleStyleSheet()
    
    if level == 1:
        style = ParagraphStyle(
            'Section1',
            parent=styles['Heading1'],
            fontSize=14,
            textColor=colors.HexColor('#2980B9'),
            spaceAfter=10,
            spaceBefore=15,
            fontName='Helvetica-Bold'
        )
    elif level == 2:
        style = ParagraphStyle(
            'Section2',
            parent=styles['Heading2'],
            fontSize=12,
            textColor=colors.HexColor('#16A085'),
            spaceAfter=8,
            spaceBefore=10,
            fontName='Helvetica-Bold'
        )
    else:
        style = ParagraphStyle(
            'Section3',
            parent=styles['Heading3'],
            fontSize=11,
            textColor=colors.HexColor('#34495E'),
            spaceAfter=6,
            spaceBefore=8,
            fontName='Helvetica-Bold'
        )
    
    if anchor:
        story.append(Paragraph(f'<a name="{anchor}"/>{text}', style))
    else:
        story.append(Paragraph(text, style))
    
    add_horizontal_rule(story, color=colors.HexColor('#BDC3C7') if level == 1 else None)

def generate_enhanced_pdf_report(results_data, topic_name="CT(A)I-detector Analysis", report_type="full"):
    """
    Generate PDF report with charts, TOC, and watermark
    
    Args:
        results_data: Analysis results
        topic_name: Report title
        report_type: "full" or "concise"
    """
    
    buffer = io.BytesIO()
    
    # Document setup with page number and watermark
    doc = SimpleDocTemplate(
        buffer, 
        pagesize=A4,
        topMargin=1.5*cm,
        bottomMargin=1.5*cm,
        leftMargin=1.5*cm,
        rightMargin=1.5*cm,
        title=topic_name,
        author="CT(A)I-detector"
    )
    
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=20,
        textColor=colors.HexColor('#2C3E50'),
        spaceAfter=12,
        alignment=TA_CENTER,
        fontName='Helvetica-Bold'
    )
    
    subtitle_style = ParagraphStyle(
        'CustomSubtitle',
        parent=styles['Heading2'],
        fontSize=16,
        textColor=colors.HexColor('#34495E'),
        spaceAfter=8,
        alignment=TA_CENTER,
        fontName='Helvetica'
    )
    
    meta_style = ParagraphStyle(
        'CustomMeta',
        parent=styles['Normal'],
        fontSize=10,
        textColor=colors.HexColor('#7F8C8D'),
        spaceAfter=3,
        alignment=TA_CENTER,
        fontName='Helvetica-Oblique'
    )
    
    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontSize=10,
        textColor=colors.HexColor('#2C3E50'),
        spaceAfter=4,
        alignment=TA_LEFT,
        fontName='Helvetica'
    )
    
    metrics_style = ParagraphStyle(
        'CustomMetrics',
        parent=styles['Normal'],
        fontSize=11,
        textColor=colors.HexColor('#27AE60'),
        spaceAfter=2,
        alignment=TA_LEFT,
        fontName='Helvetica-Bold'
    )
    
    example_style = ParagraphStyle(
        'CustomExample',
        parent=styles['Normal'],
        fontSize=9,
        textColor=colors.HexColor('#7F8C8D'),
        spaceAfter=2,
        leftIndent=20,
        alignment=TA_LEFT,
        fontName='Helvetica-Oblique'
    )
    
    footer_style = ParagraphStyle(
        'CustomFooter',
        parent=styles['Normal'],
        fontSize=8,
        textColor=colors.HexColor('#95A5A6'),
        spaceBefore=15,
        alignment=TA_CENTER,
        fontName='Helvetica-Oblique'
    )
    
    story = []
    
    # ========== TITLE PAGE ==========
    
    story.append(Spacer(1, 2*cm))
    
    # Logo or emoji
    try:
        possible_paths = ["logo.png", "./logo.png", "app/logo.png"]
        logo_path = next((p for p in possible_paths if os.path.exists(p)), None)
        if logo_path:
            logo = Image(logo_path, width=270, height=135)
            logo.hAlign = 'CENTER'
            story.append(logo)
        else:
            story.append(Paragraph("🔬", ParagraphStyle(
                'LogoEmoji',
                parent=styles['Heading1'],
                fontSize=40,
                textColor=colors.HexColor('#667eea'),
                alignment=TA_CENTER
            )))
    except:
        story.append(Paragraph("🔬", ParagraphStyle(
            'LogoEmoji',
            parent=styles['Heading1'],
            fontSize=40,
            textColor=colors.HexColor('#667eea'),
            alignment=TA_CENTER
        )))
    
    story.append(Spacer(1, 0.5*cm))
    story.append(Paragraph("Advanced AI Text Analysis Report", subtitle_style))
    story.append(Spacer(1, 1*cm))
    
    # Meta info
    current_date = datetime.now().strftime('%B %d, %Y at %H:%M')
    story.append(Paragraph(f"Generated on {current_date}", meta_style))
    
    if results_data:
        text = results_data.get('text', '')
        sentences = results_data.get('sentences', [])
        integrated = results_data.get('integrated', {})
        
        story.append(Paragraph(f"Total words analyzed: {len(text.split())}", meta_style))
        story.append(Paragraph(f"Total sentences: {len(sentences)}", meta_style))
        story.append(Paragraph(f"Risk score: {integrated.get('final_score', 0):.1f}/100", meta_style))
        story.append(Paragraph(f"Risk level: {integrated.get('risk_level', 'unknown').replace('_', ' ').title()}", meta_style))
    
    story.append(Spacer(1, 3*cm))
    story.append(Paragraph("© CT(A)I-detector", footer_style))
    story.append(Paragraph("https://chimicatechnoacta.ru developed by @daM", footer_style))
    
    # ========== TABLE OF CONTENTS (only for full report) ==========
    if report_type == "full":
        create_table_of_contents(story, doc)
    
    # ========== MAIN REPORT ==========
    
    if not results_data:
        story.append(Paragraph("No data available for analysis", normal_style))
    else:
        text = results_data.get('text', '')
        sentences = results_data.get('sentences', [])
        results = results_data.get('results', {})
        integrated = results_data.get('integrated', {})
        text_stats = results_data.get('text_statistics', {})
        module_scores = integrated.get('module_scores', [])
        
        # 1. OVERALL RISK ASSESSMENT
        add_section_header(story, "1. OVERALL RISK ASSESSMENT", level=1, anchor="section1")
        
        final_score = integrated.get('final_score', 0)
        risk_level = integrated.get('risk_level', 'unknown').replace('_', ' ').title()
        
        # Determine color
        if final_score < 30:
            risk_color = '#27AE60'
        elif final_score < 50:
            risk_color = '#F39C12'
        elif final_score < 70:
            risk_color = '#E67E22'
        else:
            risk_color = '#E74C3C'
        
        story.append(Paragraph(f"AI Risk Score: {final_score:.1f}/100", 
                              ParagraphStyle('Score', parent=metrics_style, fontSize=16, 
                                           textColor=colors.HexColor(risk_color))))
        story.append(Paragraph(f"Risk Level: {risk_level}", metrics_style))
        story.append(Spacer(1, 0.5*cm))
        
        # Charts (only for full report)
        if report_type == "full" and module_scores:
            pie_chart = create_module_pie_chart(module_scores)
            bar_chart = create_module_bar_chart(module_scores)
            
            if pie_chart and bar_chart:
                # Place charts one below another
                story.append(pie_chart)
                story.append(Spacer(1, 0.3*cm))
                story.append(bar_chart)
                story.append(Spacer(1, 0.5*cm))
        
        # Text Statistics
        add_section_header(story, "Text Statistics", level=2)
        
        stats_data = [
            ["Metric", "Value"],
            ["Characters", f"{len(text):,}"],
            ["Words", f"{len(text.split()):,}"],
            ["Sentences", f"{len(sentences):,}"],
            ["Avg. sentence length", f"{len(text.split()) / max(len(sentences), 1):.1f} words"],
            ["Modules analyzed", len(module_scores)]
        ]
        
        stats_table = Table(stats_data, colWidths=[doc.width/2.5, doc.width/3])
        stats_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#3498DB')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 11),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#F8F9FA')),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#D5DBDB')),
            ('FONTSIZE', (0, 1), (-1, -1), 10),
        ]))
        
        story.append(stats_table)
        story.append(Spacer(1, 0.5*cm))
        
        # 1.5 DETAILED TEXT STATISTICS (NEW!)
        if text_stats and report_type == "full":
            add_section_header(story, "Detailed Text Statistics", level=2)
            
            # Create a two-column layout for statistics
            stats_data_rows = []
            
            # Sentence length
            sent_stats = text_stats.get('statistics', {}).get('sentence_length', {})
            stats_data_rows.append(["Sentence Length (words)", 
                                   f"Min: {sent_stats.get('min', 0):.1f}, Max: {sent_stats.get('max', 0):.1f}, "
                                   f"Mean: {sent_stats.get('mean', 0):.1f}, Median: {sent_stats.get('median', 0):.1f}"])
            
            # Paragraph length
            para_stats = text_stats.get('statistics', {}).get('paragraph_length', {})
            stats_data_rows.append(["Paragraph Length (words)", 
                                   f"Min: {para_stats.get('min', 0):.1f}, Max: {para_stats.get('max', 0):.1f}, "
                                   f"Mean: {para_stats.get('mean', 0):.1f}, Median: {para_stats.get('median', 0):.1f}"])
            
            # Commas per sentence
            comma_stats = text_stats.get('statistics', {}).get('commas_per_sentence', {})
            stats_data_rows.append(["Commas per Sentence", 
                                   f"Min: {comma_stats.get('min', 0):.1f}, Max: {comma_stats.get('max', 0):.1f}, "
                                   f"Mean: {comma_stats.get('mean', 0):.1f}, Median: {comma_stats.get('median', 0):.1f}"])
            
            # Apostrophes per sentence
            apost_stats = text_stats.get('statistics', {}).get('apostrophes_per_sentence', {})
            stats_data_rows.append(["Apostrophes per Sentence", 
                                   f"Min: {apost_stats.get('min', 0):.1f}, Max: {apost_stats.get('max', 0):.1f}, "
                                   f"Mean: {apost_stats.get('mean', 0):.1f}, Median: {apost_stats.get('median', 0):.1f}"])
            
            # -ly adverbs per sentence
            ly_stats = text_stats.get('statistics', {}).get('ly_adverbs_per_sentence', {})
            stats_data_rows.append(["-ly Adverbs per Sentence", 
                                   f"Min: {ly_stats.get('min', 0):.1f}, Max: {ly_stats.get('max', 0):.1f}, "
                                   f"Mean: {ly_stats.get('mean', 0):.1f}, Median: {ly_stats.get('median', 0):.1f}"])
            
            # Gerund + the per sentence
            gerund_the_stats = text_stats.get('statistics', {}).get('gerund_the_per_sentence', {})
            stats_data_rows.append(["Gerund + 'the' per Sentence", 
                                   f"Min: {gerund_the_stats.get('min', 0):.1f}, Max: {gerund_the_stats.get('max', 0):.1f}, "
                                   f"Mean: {gerund_the_stats.get('mean', 0):.1f}, Median: {gerund_the_stats.get('median', 0):.1f}"])
            
            # Gerund + of per sentence
            gerund_of_stats = text_stats.get('statistics', {}).get('gerund_of_per_sentence', {})
            stats_data_rows.append(["Gerund + 'of' per Sentence", 
                                   f"Min: {gerund_of_stats.get('min', 0):.1f}, Max: {gerund_of_stats.get('max', 0):.1f}, "
                                   f"Mean: {gerund_of_stats.get('mean', 0):.1f}, Median: {gerund_of_stats.get('median', 0):.1f}"])
            
            # Indefinite articles per sentence
            article_stats = text_stats.get('statistics', {}).get('indefinite_articles_per_sentence', {})
            stats_data_rows.append(["Indefinite Articles per Sentence", 
                                   f"Min: {article_stats.get('min', 0):.1f}, Max: {article_stats.get('max', 0):.1f}, "
                                   f"Mean: {article_stats.get('mean', 0):.1f}, Median: {article_stats.get('median', 0):.1f}"])
            
            # Create table with statistics
            stats_table_detailed = Table(stats_data_rows, colWidths=[doc.width/3, doc.width/1.5])
            stats_table_detailed.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#16A085')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#F8F9FA')),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#D5DBDB')),
                ('FONTSIZE', (0, 1), (-1, -1), 9),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ]))
            
            story.append(stats_table_detailed)
            story.append(Spacer(1, 0.5*cm))
            
            # Figure mentions
            if text_stats.get('figure_mentions'):
                story.append(Paragraph(f"Figure Mentions (Total: {text_stats.get('figure_count', 0)}):", 
                                      ParagraphStyle('FigureHeading', parent=styles['Heading4'], fontSize=10)))
                for mention in text_stats['figure_mentions'][:5]:
                    clean_sent = clean_text_for_pdf(mention['sentence'])[:150]
                    story.append(Paragraph(f"• {clean_sent}...", example_style))
                if len(text_stats['figure_mentions']) > 5:
                    story.append(Paragraph(f"... and {len(text_stats['figure_mentions']) - 5} more mentions", example_style))
                story.append(Spacer(1, 0.3*cm))
            
            # Table mentions
            if text_stats.get('table_mentions'):
                story.append(Paragraph(f"Table Mentions (Total: {text_stats.get('table_count', 0)}):", 
                                      ParagraphStyle('TableHeading', parent=styles['Heading4'], fontSize=10)))
                for mention in text_stats['table_mentions'][:5]:
                    clean_sent = clean_text_for_pdf(mention['sentence'])[:150]
                    story.append(Paragraph(f"• {clean_sent}...", example_style))
                if len(text_stats['table_mentions']) > 5:
                    story.append(Paragraph(f"... and {len(text_stats['table_mentions']) - 5} more mentions", example_style))
                story.append(Spacer(1, 0.3*cm))
            
            # Supplementary mentions
            if text_stats.get('supplementary_mentions'):
                story.append(Paragraph(f"Supplementary Mentions (Total: {text_stats.get('supplementary_count', 0)}):", 
                                      ParagraphStyle('SuppHeading', parent=styles['Heading4'], fontSize=10)))
                for mention in text_stats['supplementary_mentions'][:5]:
                    clean_sent = clean_text_for_pdf(mention['sentence'])[:150]
                    story.append(Paragraph(f"• {clean_sent}...", example_style))
                if len(text_stats['supplementary_mentions']) > 5:
                    story.append(Paragraph(f"... and {len(text_stats['supplementary_mentions']) - 5} more mentions", example_style))
                story.append(Spacer(1, 0.3*cm))
        
        # 2. MODULE CONTRIBUTIONS
        add_section_header(story, "2. MODULE CONTRIBUTIONS", level=1, anchor="section2")
        
        if module_scores:
            sorted_modules = sorted(module_scores, key=lambda x: x.get('contribution', 0), reverse=True)
            
            module_data = [["Module", "Raw Score", "Contribution", "Confidence"]]
            for ms in sorted_modules[:10 if report_type == "full" else 5]:  # Show fewer in concise
                module_data.append([
                    ms.get('module', 'unknown').replace('_', ' ').title(),
                    f"{ms.get('raw_score', 0)}/6",
                    f"{ms.get('contribution', 0):.1f}%",
                    f"{ms.get('confidence', 0):.2f}"
                ])
            
            module_table = Table(module_data, colWidths=[doc.width/3, doc.width/6, doc.width/6, doc.width/6])
            module_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#9B59B6')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 11),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#D5DBDB')),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F2F4F4')]),
            ]))
            
            story.append(module_table)
        else:
            story.append(Paragraph("No module data available", normal_style))
        
        story.append(Spacer(1, 0.5*cm))
        
        # ========== FOR CONCISE REPORT: Skip to Conclusion ==========
        if report_type == "concise":
            # Add a few key examples
            add_section_header(story, "Key Examples", level=2)
            
            examples_added = 0
            
            # Top AI phrases
            if 'phrases' in results and results['phrases'].get('all_phrase_occurrences'):
                occurrences = results['phrases']['all_phrase_occurrences']
                if occurrences:
                    story.append(Paragraph("AI Phrases:", metrics_style))
                    for occ in occurrences[:3]:
                        context = clean_text_for_pdf(occ.get('context', ''))[:150]
                        story.append(Paragraph(f"• '{occ.get('phrase', '')}' → ...{context}...", example_style))
                    examples_added += 1
                    story.append(Spacer(1, 0.2*cm))

            # Tortured Phrases
            if 'tortured_phrases' in results:
                add_section_header(story, "3.3 Tortured Phrases", level=2, anchor="section3.3")
                tp_data = results['tortured_phrases']
                occurrences = tp_data.get('all_occurrences', [])
                
                story.append(Paragraph(f"• Total occurrences: {tp_data.get('total_occurrences', 0)}", normal_style))
                story.append(Paragraph(f"• Unique tortured phrases: {tp_data.get('unique_tortured_count', 0)}", normal_style))
                story.append(Paragraph(f"• Frequency per 1000 words: {tp_data.get('tortured_per_1000_words', 0):.2f}", normal_style))
                
                if tp_data.get('tortured_phrases_found'):
                    story.append(Spacer(1, 0.2*cm))
                    story.append(Paragraph("Tortured phrases found:", 
                                          ParagraphStyle('Bold', parent=normal_style, fontName='Helvetica-Bold')))
                    for phrase in tp_data['tortured_phrases_found']:
                        story.append(Paragraph(f"• '{phrase['tortured']}' → correct: {phrase['correct']} (found {phrase['count']} times)", 
                                              example_style))
                
                if occurrences:
                    story.append(Spacer(1, 0.2*cm))
                    story.append(Paragraph("Examples with context:", 
                                          ParagraphStyle('Bold', parent=normal_style, fontName='Helvetica-Bold')))
                    
                    by_phrase = {}
                    for occ in occurrences:
                        key = f"{occ['tortured']} → {occ['correct']}"
                        if key not in by_phrase:
                            by_phrase[key] = []
                        by_phrase[key].append(occ)
                    
                    for phrase_group, group_occurrences in list(by_phrase.items())[:5]:  # Топ-5 фраз
                        story.append(Paragraph(f"  {phrase_group}:", example_style))
                        for i, occ in enumerate(group_occurrences[:3]):  # По 3 примера на фразу
                            context = clean_text_for_pdf(occ['context'])[:250]
                            story.append(Paragraph(f"    {i+1}. ...{context}...", example_style))
                        story.append(Spacer(1, 0.1*cm))
                
                story.append(Spacer(1, 0.2*cm))
            
            # Enumerations
            if 'enumeration' in results and results['enumeration'].get('all_enumerations'):
                enumerations = results['enumeration']['all_enumerations']
                if enumerations:
                    story.append(Paragraph("Strict Enumerations:", metrics_style))
                    for enum in enumerations[:3]:
                        story.append(Paragraph(f"• {clean_text_for_pdf(enum)[:150]}...", example_style))
                    examples_added += 1
                    story.append(Spacer(1, 0.2*cm))
            
            if examples_added == 0:
                story.append(Paragraph("No notable examples found", normal_style))
            
            story.append(Spacer(1, 0.5*cm))
            
            # Skip to conclusion
            story.append(PageBreak())
            add_section_header(story, "5. CONCLUSION & RECOMMENDATIONS", level=1, anchor="section5")
        
        # ========== FOR FULL REPORT: Detailed Analysis ==========
        else:
            # 3. DETAILED MODULE ANALYSIS
            add_section_header(story, "3. DETAILED MODULE ANALYSIS", level=1, anchor="section3")
            
            # Unicode artifacts
            if 'unicode' in results:
                add_section_header(story, "3.1 Unicode Artifacts", level=2, anchor="section3.1")
                unicode_data = results['unicode']
                story.append(Paragraph(f"• Suspicious characters: {unicode_data.get('sup_sub_count', 0) + unicode_data.get('fullwidth_count', 0)}", normal_style))
                story.append(Paragraph(f"• Density per 10k chars: {unicode_data.get('density_per_10k', 0):.2f}", normal_style))
                
                chunks = unicode_data.get('all_suspicious_chunks', [])
                if chunks:
                    story.append(Paragraph("Examples:", example_style))
                    for chunk in chunks[:50]:
                        context = clean_text_for_pdf(chunk.get('context', ''))[:150]
                        story.append(Paragraph(f"  • '{chunk.get('char', '')}' → ...{context}...", example_style))
                story.append(Spacer(1, 0.2*cm))
            
            # AI phrases
            if 'phrases' in results:
                add_section_header(story, "3.2 AI Phrases", level=2, anchor="section3.2")
                phrases_data = results['phrases']
                occurrences = phrases_data.get('all_phrase_occurrences', [])
                story.append(Paragraph(f"• Total occurrences: {len(occurrences)}", normal_style))
                
                if occurrences:
                    story.append(Paragraph("Examples:", example_style))
                    for occ in occurrences[:50]:
                        context = clean_text_for_pdf(occ.get('context', ''))[:150]
                        story.append(Paragraph(f"  • '{occ.get('phrase', '')}' → ...{context}...", example_style))
                story.append(Spacer(1, 0.2*cm))
            
            # Tortured Phrases
            if 'tortured_phrases' in results:
                add_section_header(story, "3.3 Tortured Phrases", level=2, anchor="section3.3")
                tp_data = results['tortured_phrases']
                occurrences = tp_data.get('all_occurrences', [])
                
                story.append(Paragraph(f"• Total occurrences: {tp_data.get('total_occurrences', 0)}", normal_style))
                story.append(Paragraph(f"• Unique tortured phrases: {tp_data.get('unique_tortured_count', 0)}", normal_style))
                story.append(Paragraph(f"• Frequency per 1000 words: {tp_data.get('tortured_per_1000_words', 0):.2f}", normal_style))
                
                if tp_data.get('tortured_phrases_found'):
                    story.append(Spacer(1, 0.2*cm))
                    story.append(Paragraph("Tortured phrases found:", 
                                          ParagraphStyle('Bold', parent=normal_style, fontName='Helvetica-Bold')))
                    for phrase in tp_data['tortured_phrases_found']:
                        story.append(Paragraph(f"• '{phrase['tortured']}' → correct: {phrase['correct']} (found {phrase['count']} times)", 
                                              example_style))
                
                if occurrences:
                    story.append(Spacer(1, 0.2*cm))
                    story.append(Paragraph("Examples with context:", 
                                          ParagraphStyle('Bold', parent=normal_style, fontName='Helvetica-Bold')))
                    
                    by_phrase = {}
                    for occ in occurrences:
                        key = f"{occ['tortured']} → {occ['correct']}"
                        if key not in by_phrase:
                            by_phrase[key] = []
                        by_phrase[key].append(occ)
                    
                    for phrase_group, group_occurrences in list(by_phrase.items())[:5]:  # Топ-5 фраз
                        story.append(Paragraph(f"  {phrase_group}:", example_style))
                        for i, occ in enumerate(group_occurrences[:3]):  # По 3 примера на фразу
                            context = clean_text_for_pdf(occ['context'])[:250]
                            story.append(Paragraph(f"    {i+1}. ...{context}...", example_style))
                        story.append(Spacer(1, 0.1*cm))
                
                story.append(Spacer(1, 0.2*cm))
            
            # Enumerations
            if 'enumeration' in results:
                add_section_header(story, "3.4 Strict Enumerations", level=2, anchor="section3.4")
                enum_data = results['enumeration']
                enumerations = enum_data.get('all_enumerations', [])
                story.append(Paragraph(f"• Found: {len(enumerations)} enumerations", normal_style))
                
                if enumerations:
                    story.append(Paragraph("Examples:", example_style))
                    for enum in enumerations[:50]:
                        story.append(Paragraph(f"  • {clean_text_for_pdf(enum)[:200]}...", example_style))
                story.append(Spacer(1, 0.2*cm))
            
            # Apostrophes
            if 'apostrophe' in results:
                add_section_header(story, "3.5 Apostrophe Usage", level=2, anchor="section3.5")
                apost_data = results['apostrophe']
                apostrophes = apost_data.get('all_apostrophes', [])
                story.append(Paragraph(f"• Found: {len(apostrophes)} apostrophes", normal_style))
                
                if apostrophes:
                    examples = ', '.join(apostrophes[:50])
                    story.append(Paragraph(f"Examples: {examples}", example_style))
                story.append(Spacer(1, 0.2*cm))
            
            # Punctuation
            if 'punctuation' in results:
                add_section_header(story, "3.6 Punctuation Analysis", level=2, anchor="section3.6")
                punct_data = results['punctuation']
                story.append(Paragraph(f"• Exclamation marks: {punct_data.get('exclamation_count', 0)}", normal_style))
                story.append(Paragraph(f"• Question marks: {punct_data.get('question_count', 0)}", normal_style))
                story.append(Paragraph(f"• Semicolons: {punct_data.get('semicolon_count', 0)}", normal_style))
                
                semicolons = punct_data.get('all_semicolon_contexts', [])
                if semicolons:
                    story.append(Paragraph("Semicolon contexts:", example_style))
                    for ctx in semicolons[:3]:
                        context = clean_text_for_pdf(ctx.get('context', ''))[:100]
                        story.append(Paragraph(f"  • ...{context}...", example_style))
                story.append(Spacer(1, 0.2*cm))
            
            # Parentheses
            if 'parenthesis' in results:
                add_section_header(story, "3.7 Parenthesis Analysis", level=2, anchor="section3.7")
                paren_data = results['parenthesis']
                parentheses = paren_data.get('all_parentheses', [])
                story.append(Paragraph(f"• Long parentheses (≥5 words): {paren_data.get('long_parentheses', 0)}", normal_style))
                
                if parentheses:
                    story.append(Paragraph("Examples:", example_style))
                    for p in parentheses[:3]:
                        text_content = clean_text_for_pdf(p.get('text', ''))[:100]
                        story.append(Paragraph(f"  • ({text_content}) — {p.get('word_count', 0)} words", example_style))
                story.append(Spacer(1, 0.2*cm))
            
            # Repetitions
            if 'repetitiveness' in results:
                add_section_header(story, "3.8 Repetitiveness Analysis", level=2, anchor="section3.8")
                rep_data = results['repetitiveness']
                repetitions = rep_data.get('all_repetitions', [])
                story.append(Paragraph(f"• Found: {len(repetitions)} repeated phrases", normal_style))
                
                if repetitions:
                    story.append(Paragraph("Top repetitions:", example_style))
                    for rep in repetitions[:50]:
                        story.append(Paragraph(f"  • '{rep.get('ngram', '')}' — {rep.get('count', 0)} times", example_style))
                story.append(Spacer(1, 0.2*cm))
            
            # Dashes
            if 'dashes' in results:
                add_section_header(story, "3.9 Dash Analysis", level=2, anchor="section3.9")
                dash_data = results['dashes']
                
                # Общая статистика
                all_dashes = dash_data.get('all_dash_sentences', [])
                double_dashes = dash_data.get('double_dash_sentences', [])
                
                story.append(Paragraph(f"• Total sentences with dashes: {len(all_dashes)}", normal_style))
                story.append(Paragraph(f"• Sentences with TWO dashes: {len(double_dashes)}", normal_style))
                story.append(Paragraph(f"• Heavy sentences (3+ dashes): {len(dash_data.get('heavy_sentences', []))}", normal_style))
                
                # Отдельная секция для предложений с двумя тире
                if double_dashes:
                    story.append(Spacer(1, 0.2*cm))
                    story.append(Paragraph("Sentences with TWO dashes (critical pattern):", 
                                          ParagraphStyle('Bold', parent=normal_style, fontName='Helvetica-Bold')))
                    for i, item in enumerate(double_dashes[:50]):  # Лимит 50 для PDF, но можно убрать
                        story.append(Paragraph(f"{i+1}. {clean_text_for_pdf(item['sentence'])}", example_style))
                
                # Все предложения с тире (если нужно показать все, уберите [:50])
                if all_dashes and len(double_dashes) < len(all_dashes):
                    story.append(Spacer(1, 0.2*cm))
                    story.append(Paragraph("Other sentences with dashes:", 
                                          ParagraphStyle('Bold', parent=normal_style, fontName='Helvetica-Bold')))
                    for i, item in enumerate(all_dashes[:30]):  # Можно увеличить лимит до 100
                        if item not in double_dashes:  # Не повторяем уже показанные
                            story.append(Paragraph(f"{i+1}. {clean_text_for_pdf(item['sentence'][:300])}...", example_style))
                
                story.append(Spacer(1, 0.2*cm))
            
            # Lexical diversity
            if 'lexical_diversity' in results:
                add_section_header(story, "3.10 Lexical Diversity", level=2, anchor="section3.10")
                lex_data = results['lexical_diversity']
                story.append(Paragraph(f"• TTR: {lex_data.get('ttr', 0):.3f}", normal_style))
                if lex_data.get('mtld', 0) > 0:
                    story.append(Paragraph(f"• MTLD: {lex_data.get('mtld', 0):.1f}", normal_style))
                story.append(Paragraph(f"• Hapax ratio: {lex_data.get('hapax_ratio', 0):.3f}", normal_style))
            
            story.append(PageBreak())
            
            # 4. TEXT SAMPLES
            add_section_header(story, "4. TEXT SAMPLES", level=1, anchor="section4")
            
            # First 5 sentences
            add_section_header(story, "4.1 First 5 sentences", level=2)
            for i, sent in enumerate(sentences[:50]):
                clean_sent = clean_text_for_pdf(sent)[:200]
                story.append(Paragraph(f"{i+1}. {clean_sent}...", example_style))
            
            story.append(Spacer(1, 0.5*cm))
            
            # High-risk examples
            high_risk_examples = []
            
            if 'dashes' in results:
                for item in results['dashes'].get('heavy_sentences', [])[:3]:
                    high_risk_examples.append(("Heavy dash usage", item.get('sentence', '')))
            
            if 'enumeration' in results:
                for enum in results['enumeration'].get('all_enumerations', [])[:3]:
                    high_risk_examples.append(("Strict enumeration", enum))
            
            if 'phrases' in results:
                for occ in results['phrases'].get('all_phrase_occurrences', [])[:3]:
                    high_risk_examples.append(("AI phrase", occ.get('context', '')))
            
            if high_risk_examples:
                add_section_header(story, "4.2 High-risk examples", level=2)
                for i, (label, example) in enumerate(high_risk_examples[:50]):
                    clean_example = clean_text_for_pdf(example)[:150]
                    story.append(Paragraph(f"{label}: {clean_example}...", example_style))
            
            story.append(PageBreak())
            
            # 5. CONCLUSION & RECOMMENDATIONS
            add_section_header(story, "5. CONCLUSION & RECOMMENDATIONS", level=1, anchor="section5")
        
        # ========== CONCLUSION (common for both report types) ==========
        
        # Interpretation
        add_section_header(story, "Interpretation", level=2)
        
        if final_score < 30:
            interpretation = [
                "• LOW RISK: Text shows patterns consistent with human writing.",
                "• High lexical diversity and natural variation in sentence structure.",
                "• Presence of human markers (varied punctuation, hedging, personal pronouns)."
            ]
        elif final_score < 50:
            interpretation = [
                "• MEDIUM-LOW RISK: Some AI-like patterns detected.",
                "• Moderate use of AI phrases and repetitive structures.",
                "• Consider reviewing specific flagged sections."
            ]
        elif final_score < 70:
            interpretation = [
                "• MEDIUM-HIGH RISK: Significant AI-like patterns detected.",
                "• High density of AI phrases and low lexical diversity.",
                "• Multiple modules show elevated risk scores."
            ]
        else:
            interpretation = [
                "• HIGH RISK: Text shows strong AI generation patterns.",
                "• Very low lexical diversity with high repetitiveness.",
                "• Many characteristic AI phrases and structures present."
            ]
        
        for line in interpretation:
            story.append(Paragraph(line, normal_style))
        
        story.append(Spacer(1, 0.5*cm))
        
        # Recommendations
        add_section_header(story, "Recommendations", level=2)
        
        recommendations = [
            "• Review the flagged examples in each module section",
            "• Pay special attention to modules with raw scores ≥ 4/6",
            "• Consider the context of the text (academic vs. casual writing)",
            "• For research purposes, combine with other detection methods"
        ]
        
        for rec in recommendations:
            story.append(Paragraph(rec, normal_style))
        
        story.append(Spacer(1, 0.5*cm))
        
        # Technical info
        add_section_header(story, "Technical Details", level=2)
        story.append(Paragraph(f"• Analysis timestamp: {current_date}", example_style))
        story.append(Paragraph(f"• Modules with data: {', '.join([m.get('module', '') for m in module_scores[:50]])}", example_style))
        story.append(Paragraph(f"• Confidence score: {integrated.get('total_confidence', 0):.2f}", example_style))
    
    # Footer
    story.append(Spacer(1, 1*cm))
    story.append(Paragraph("─" * 70, footer_style))
    story.append(Paragraph("© CT(A)I-detector - Advanced AI Text Analysis Tool", footer_style))
    story.append(Paragraph(f"Report ID: {hashlib.md5(str(datetime.now()).encode()).hexdigest()[:8].upper()}", footer_style))
    
    # Build PDF with watermark
    if report_type == "full":
        doc.build(story, onFirstPage=add_watermark, onLaterPages=add_watermark)
    else:
        doc.build(story)
    
    return buffer.getvalue()

# ============================================================================
# Main application
# ============================================================================

def main():
    # Initialize session state for steps
    if 'step' not in st.session_state:
        st.session_state.step = 1
    if 'analyze_clicked' not in st.session_state:
        st.session_state.analyze_clicked = False
    if 'file_uploaded' not in st.session_state:
        st.session_state.file_uploaded = False
    if 'results' not in st.session_state:
        st.session_state.results = None
    
    # Header with logo
    col1, col2, col3 = st.columns([2, 2, 2])
    with col2:
        try:
            st.image('logo.png', use_column_width=True)
        except:
            st.markdown('<h1 style="text-align: center; color: #667eea;">CT(A)I-detector</h1>', unsafe_allow_html=True)
    
    # Step indicator
    st.markdown(f"""
    <div class="step-indicator">
        <div class="step-item">
            <div class="step-circle {'active' if st.session_state.step >= 1 else ''} {'completed' if st.session_state.step > 1 else ''}">1</div>
            <div class="step-label {'active' if st.session_state.step >= 1 else ''}">Upload Sample</div>
        </div>
        <div class="step-item">
            <div class="step-circle {'active' if st.session_state.step >= 2 else ''} {'completed' if st.session_state.step > 2 else ''}">2</div>
            <div class="step-label {'active' if st.session_state.step >= 2 else ''}">Analysis</div>
        </div>
        <div class="step-item">
            <div class="step-circle {'active' if st.session_state.step >= 3 else ''}">3</div>
            <div class="step-label {'active' if st.session_state.step >= 3 else ''}">Results</div>
        </div>
    </div>
    """, unsafe_allow_html=True)
    
    # Step 1: Upload
    if st.session_state.step == 1:
        with st.container():
            st.markdown('<div class="step-container">', unsafe_allow_html=True)
            
            st.markdown("""
            <div class="upload-area">
                <div class="upload-icon">🔬</div>
                <h3>Drop your sample here</h3>
                <p style="color: #666;">Supports .docx and .doc files</p>
            </div>
            """, unsafe_allow_html=True)
            
            uploaded_file = st.file_uploader("", type=['docx', 'doc'], label_visibility="collapsed")
            
            if uploaded_file:
                st.session_state.file_uploaded = True
                st.session_state.uploaded_file = uploaded_file
                
                # File info card
                st.markdown(f"""
                <div style="background: #f8f9fa; border-radius: 12px; padding: 1rem; margin-top: 1rem;">
                    <div style="display: flex; align-items: center; gap: 1rem;">
                        <div style="background: #667eea; color: white; width: 40px; height: 40px; border-radius: 8px; display: flex; align-items: center; justify-content: center;">📄</div>
                        <div>
                            <div style="font-weight: 600;">{uploaded_file.name}</div>
                            <div style="font-size: 0.9rem; color: #666;">{uploaded_file.size / 1024:.1f} KB</div>
                        </div>
                    </div>
                </div>
                """, unsafe_allow_html=True)
                
                col1, col2, col3 = st.columns([1, 1, 1])
                with col2:
                    if st.button("Start Analysis", type="primary", use_container_width=True):
                        st.session_state.step = 2
                        st.session_state.analyze_clicked = True
                        st.rerun()
            
            st.markdown('</div>', unsafe_allow_html=True)
    
    # Step 2: Analysis Process
    elif st.session_state.step == 2:
        with st.container():
            st.markdown('<div class="step-container">', unsafe_allow_html=True)
            
            st.markdown("""
            <div class="lab-loader">
                <h3>Analyzing Sample</h3>
                <div class="analyzer-beam"></div>
            </div>
            """, unsafe_allow_html=True)
            
            # Module status simulation
            modules = [
                "Unicode Artifacts", "Dashes", "AI Phrases", "Burstiness", "Grammar",
                "Hedging", "Parentheses", "Punctuation", "Apostrophes", "Enumerations",
                "Paragraphs", "Repetitiveness", "Lexical Diversity", "Semantic Analysis",
                "Text Statistics"
            ]
            
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            # Process analysis
            uploaded_file = st.session_state.uploaded_file
            
            try:
                # Read document
                status_text.text("Reading document...")
                progress_bar.progress(5)
                time.sleep(0.5)
                
                if uploaded_file.name.endswith('.docx'):
                    text = DocumentProcessor.read_docx(uploaded_file)
                else:
                    text = DocumentProcessor.read_doc(uploaded_file)
                
                if not text or len(text.strip()) < 100:
                    st.error("File is too small or empty")
                    st.session_state.step = 1
                    st.rerun()
                
                # Cut references
                status_text.text("Preprocessing text...")
                progress_bar.progress(10)
                time.sleep(0.5)
                
                text = ReferenceCutoff.cut_at_references(text)
                text = DocumentProcessor.preprocess(text)
                sentences = DocumentProcessor.split_sentences_simple(text)
                
                # Module status display
                module_container = st.empty()
                
                # Results storage
                results = {}
                
                # Run all analyzers with progress updates
                analyzers = [
                    ('unicode', UnicodeArtifactDetector(), 13),
                    ('dashes', DashAnalyzer(), 17),
                    ('phrases', AIPhraseDetector(), 21),
                    ('tortured_phrases', TorturedPhraseDetector(), 25),
                    ('burstiness', BurstinessAnalyzer(), 29),
                    ('grammar', GrammarAnalyzer(), 33),
                    ('hedging', HedgingAnalyzer(), 37),
                    ('parenthesis', ParenthesisAnalyzer(), 41),
                    ('punctuation', PunctuationAnalyzer(), 45),
                    ('apostrophe', ApostropheAnalyzer(), 49),
                    ('enumeration', EnumerationAnalyzer(), 53),
                    ('paragraph', ParagraphAnalyzer(), 57),
                    ('repetitiveness', RepetitivenessAnalyzer(), 61),
                    ('lexical_diversity', LexicalDiversityAnalyzer(), 65)
                ]
                
                # Show modules being analyzed
                for i, (name, analyzer, progress) in enumerate(analyzers):
                    status_text.text(f"Analyzing: {name.replace('_', ' ').title()}...")
                    
                    # Update module pills
                    module_html = '<div class="module-status">'
                    for j, mod in enumerate(modules):
                        if j < i:
                            module_html += f'<span class="module-pill completed">✓ {mod}</span>'
                        elif j == i:
                            module_html += f'<span class="module-pill active">⚡ {mod}</span>'
                        else:
                            module_html += f'<span class="module-pill">{mod}</span>'
                    module_html += '</div>'
                    module_container.markdown(module_html, unsafe_allow_html=True)
                    
                    # Run analyzer
                    if name in ['phrases', 'punctuation', 'enumeration', 'paragraph', 'repetitiveness']:
                        results[name] = analyzer.analyze(text, sentences)
                    else:
                        results[name] = analyzer.analyze(text)
                    
                    progress_bar.progress(progress)
                    time.sleep(0.3)
                
                # Text Statistics Analyzer (NEW!)
                status_text.text("Calculating detailed text statistics...")
                progress_bar.progress(68)
                text_stats_analyzer = TextStatisticsAnalyzer()
                text_stats = text_stats_analyzer.analyze(text, sentences)
                results['text_statistics'] = text_stats
                
                # Update module pills for text statistics
                module_html = '<div class="module-status">'
                for j, mod in enumerate(modules):
                    if j < len(analyzers):
                        module_html += f'<span class="module-pill completed">✓ {mod}</span>'
                    elif j == len(analyzers):
                        module_html += f'<span class="module-pill active">⚡ {mod}</span>'
                    else:
                        module_html += f'<span class="module-pill">{mod}</span>'
                module_html += '</div>'
                module_container.markdown(module_html, unsafe_allow_html=True)
                
                progress_bar.progress(72)
                time.sleep(0.3)
                
                # Deep analysis if available
                deep_analysis = st.checkbox("Enable Deep Analysis (slower)", value=False)
                if deep_analysis:
                    deep_analyzers = [
                        ('log_prob', LogProbAnalyzer(), 78),
                        ('perplexity', PerplexityAnalyzer(), 84),
                        ('semantic', SemanticAnalyzer(), 90),
                        ('ml_classifier', MLClassifier(), 95)
                    ]
                    
                    for name, analyzer, progress in deep_analyzers:
                        status_text.text(f"Deep analysis: {name.replace('_', ' ').title()}...")
                        if name == 'semantic':
                            results[name] = analyzer.analyze(sentences)
                        else:
                            results[name] = analyzer.analyze(text)
                        progress_bar.progress(progress)
                        time.sleep(0.5)
                
                # Calculate integrated risk
                status_text.text("Calculating integrated risk score...")
                scorer = IntegratedRiskScorer()
                integrated = scorer.calculate(results)
                results['integrated'] = integrated
                
                progress_bar.progress(100)
                status_text.text("Analysis complete!")
                time.sleep(0.5)
                
                # Store results
                st.session_state.results = {
                    'text': text,
                    'sentences': sentences,
                    'results': results,
                    'integrated': integrated,
                    'text_statistics': text_stats
                }
                
                # Move to step 3
                st.session_state.step = 3
                st.rerun()
                
            except Exception as e:
                st.error(f"Error during analysis: {str(e)}")
                if st.button("Back to Upload"):
                    st.session_state.step = 1
                    st.rerun()
            
            st.markdown('</div>', unsafe_allow_html=True)
    
    # Step 3: Results
    elif st.session_state.step == 3:
        data = st.session_state.results
        text = data['text']
        sentences = data['sentences']
        results = data['results']
        integrated = data['integrated']
        text_stats = data.get('text_statistics', {})
        
        with st.container():
            st.markdown('<div class="step-container">', unsafe_allow_html=True)
            
            # Main score card
            col1, col2 = st.columns([1, 1])
            
            with col1:
                risk_class = f"verdict-{integrated['risk_level'].split('_')[0]}"
                st.markdown(f"""
                <div class="score-card">
                    <div class="score-label">AI Risk Score</div>
                    <div class="score-number">{integrated['final_score']:.1f}</div>
                    <div class="score-label">out of 100</div>
                    <div class="verdict-badge {risk_class}">
                        {integrated['risk_level'].replace('_', ' ').title()}
                    </div>
                </div>
                """, unsafe_allow_html=True)
            
            with col2:
                st.markdown("""
                <div style="background: white; border-radius: 16px; padding: 1.5rem; height: 100%;">
                    <h4 style="margin-top: 0;">Sample Statistics</h4>
                """, unsafe_allow_html=True)
                
                # Создаем колонки Streamlit вместо ручного HTML
                col_metrics = st.columns(4)
                
                with col_metrics[0]:
                    st.markdown(f"""
                    <div class="metric-card">
                        <div class="metric-title">Sentences</div>
                        <div class="metric-value">{len(sentences)}<span class="metric-unit"></span></div>
                    </div>
                    """, unsafe_allow_html=True)
                
                with col_metrics[1]:
                    st.markdown(f"""
                    <div class="metric-card">
                        <div class="metric-title">Words</div>
                        <div class="metric-value">{len(text.split())}<span class="metric-unit"></span></div>
                    </div>
                    """, unsafe_allow_html=True)
                
                with col_metrics[2]:
                    st.markdown(f"""
                    <div class="metric-card">
                        <div class="metric-title">Characters</div>
                        <div class="metric-value">{len(text)}<span class="metric-unit"></span></div>
                    </div>
                    """, unsafe_allow_html=True)
                
                with col_metrics[3]:
                    st.markdown(f"""
                    <div class="metric-card">
                        <div class="metric-title">Modules</div>
                        <div class="metric-value">{len(integrated['module_scores'])}<span class="metric-unit"></span></div>
                    </div>
                    """, unsafe_allow_html=True)
            
            # Module spectrum
            st.markdown("""
            <div class="spectrum-container">
                <h4>Module Contribution Spectrum</h4>
                <div style="position: relative; height: 40px; margin: 2rem 0;">
                    <div class="spectrum-bar"></div>
                    <div class="spectrum-marker" style="left: 50%;"></div>
                </div>
            </div>
            """, unsafe_allow_html=True)
            
            # Module contributions
            st.markdown('<div class="metric-grid">', unsafe_allow_html=True)
            
            for module_score in integrated['module_scores']:
                # Определяем цвет и стиль на основе raw_score
                if module_score['raw_score'] >= 5:
                    color = "#9b1f1f"  # темно-красный для очень сильных сигналов
                    border = "2px solid #9b1f1f"
                    bg_color = "#ffeeee"
                    emoji = "🔴"
                elif module_score['raw_score'] >= 4:
                    color = "#ef4444"  # красный
                    border = "1px solid #ef4444"
                    bg_color = "#fff5f5"
                    emoji = "⚠️"
                elif module_score['raw_score'] >= 3:
                    color = "#f97316"  # оранжевый
                    border = "1px solid #f97316"
                    bg_color = "#fff7ed"
                    emoji = "⚡"
                elif module_score['raw_score'] >= 2:
                    color = "#fbbf24"  # желтый
                    border = "none"
                    bg_color = "white"
                    emoji = "📊"
                else:
                    color = "#10b981"  # зеленый
                    border = "none"
                    bg_color = "white"
                    emoji = "✅"
                
                st.markdown(f"""
                <div class="metric-card" style="border: {border}; background-color: {bg_color};">
                    <div class="metric-title">
                        {emoji} {module_score['module'].replace('_', ' ').title()}
                        <span style="float: right; font-size: 0.8rem; background: {color}; color: white; padding: 2px 8px; border-radius: 12px;">
                            raw: {module_score['raw_score']}/6
                        </span>
                    </div>
                    <div style="display: flex; align-items: center; gap: 1rem; margin-top: 0.5rem;">
                        <div style="flex: 1;">
                            <div style="height: 8px; background: #e0e0e0; border-radius: 4px;">
                                <div class="contribution-bar" style="width: {module_score['contribution']}%; background: {color};"></div>
                            </div>
                        </div>
                        <div style="font-weight: 600; color: {color};">{module_score['contribution']:.1f}%</div>
                    </div>
                </div>
                """, unsafe_allow_html=True)
            
            st.markdown('</div>', unsafe_allow_html=True)
            
            # NEW: Detailed Text Statistics Section
            if text_stats:
                st.markdown("---")
                st.markdown("## 📊 Detailed Text Statistics")
                
                # Create columns for statistics display
                col1, col2 = st.columns(2)
                
                with col1:
                    st.markdown("### Sentence-Level Statistics")
                    
                    # Sentence length
                    sent_stats = text_stats.get('statistics', {}).get('sentence_length', {})
                    st.metric("Sentence Length (words)", 
                             f"μ={sent_stats.get('mean', 0):.1f} | med={sent_stats.get('median', 0):.1f}",
                             f"min={sent_stats.get('min', 0):.0f} max={sent_stats.get('max', 0):.0f}")
                    
                    # Commas per sentence
                    comma_stats = text_stats.get('statistics', {}).get('commas_per_sentence', {})
                    st.metric("Commas per Sentence", 
                             f"μ={comma_stats.get('mean', 0):.1f} | med={comma_stats.get('median', 0):.1f}",
                             f"min={comma_stats.get('min', 0):.0f} max={comma_stats.get('max', 0):.0f}")
                    
                    # Apostrophes per sentence
                    apost_stats = text_stats.get('statistics', {}).get('apostrophes_per_sentence', {})
                    st.metric("Apostrophes per Sentence", 
                             f"μ={apost_stats.get('mean', 0):.1f} | med={apost_stats.get('median', 0):.1f}",
                             f"min={apost_stats.get('min', 0):.0f} max={apost_stats.get('max', 0):.0f}")
                    
                    # -ly adverbs per sentence
                    ly_stats = text_stats.get('statistics', {}).get('ly_adverbs_per_sentence', {})
                    st.metric("-ly Adverbs per Sentence", 
                             f"μ={ly_stats.get('mean', 0):.1f} | med={ly_stats.get('median', 0):.1f}",
                             f"min={ly_stats.get('min', 0):.0f} max={ly_stats.get('max', 0):.0f}")
                
                with col2:
                    st.markdown("### Sentence-Level Statistics (continued)")
                    
                    # Gerund + the per sentence
                    gerund_the_stats = text_stats.get('statistics', {}).get('gerund_the_per_sentence', {})
                    st.metric("Gerund + 'the' per Sentence", 
                             f"μ={gerund_the_stats.get('mean', 0):.1f} | med={gerund_the_stats.get('median', 0):.1f}",
                             f"min={gerund_the_stats.get('min', 0):.0f} max={gerund_the_stats.get('max', 0):.0f}")
                    
                    # Gerund + of per sentence
                    gerund_of_stats = text_stats.get('statistics', {}).get('gerund_of_per_sentence', {})
                    st.metric("Gerund + 'of' per Sentence", 
                             f"μ={gerund_of_stats.get('mean', 0):.1f} | med={gerund_of_stats.get('median', 0):.1f}",
                             f"min={gerund_of_stats.get('min', 0):.0f} max={gerund_of_stats.get('max', 0):.0f}")
                    
                    # Indefinite articles per sentence
                    article_stats = text_stats.get('statistics', {}).get('indefinite_articles_per_sentence', {})
                    st.metric("Indefinite Articles (a/an) per Sentence", 
                             f"μ={article_stats.get('mean', 0):.1f} | med={article_stats.get('median', 0):.1f}",
                             f"min={article_stats.get('min', 0):.0f} max={article_stats.get('max', 0):.0f}")
                    
                    # Paragraph length (if available)
                    para_stats = text_stats.get('statistics', {}).get('paragraph_length', {})
                    if para_stats:
                        st.metric("Paragraph Length (words)", 
                                 f"μ={para_stats.get('mean', 0):.1f} | med={para_stats.get('median', 0):.1f}",
                                 f"min={para_stats.get('min', 0):.0f} max={para_stats.get('max', 0):.0f}")
                
                # Figure, Table, Supplementary mentions
                st.markdown("### References to Figures, Tables, and Supplementary Materials")
                
                col_fig1, col_fig2, col_fig3 = st.columns(3)
                
                with col_fig1:
                    st.metric("Figure Mentions", text_stats.get('figure_count', 0))
                    if text_stats.get('figure_mentions'):
                        with st.expander("View Figure mentions"):
                            for mention in text_stats['figure_mentions'][:20]:
                                st.write(f"**{mention['match']}**: {mention['sentence'][:200]}...")
                
                with col_fig2:
                    st.metric("Table Mentions", text_stats.get('table_count', 0))
                    if text_stats.get('table_mentions'):
                        with st.expander("View Table mentions"):
                            for mention in text_stats['table_mentions'][:20]:
                                st.write(f"**{mention['match']}**: {mention['sentence'][:200]}...")
                
                with col_fig3:
                    st.metric("Supplementary Mentions", text_stats.get('supplementary_count', 0))
                    if text_stats.get('supplementary_mentions'):
                        with st.expander("View Supplementary mentions"):
                            for mention in text_stats['supplementary_mentions'][:20]:
                                st.write(f"**{mention['match']}**: {mention['sentence'][:200]}...")
            
            # Detailed results tabs
            tabs = st.tabs(["Artifacts", "Enumerations", "Repetitions", "Lexical", "All Examples"])
            
            with tabs[0]:
                col1, col2 = st.columns(2)
                
                with col1:
                    if 'unicode' in results:
                        with st.expander(f"Unicode Artifacts ({len(results['unicode']['all_suspicious_chunks'])} found)"):
                            for chunk in results['unicode']['all_suspicious_chunks'][:50]:
                                st.code(f"'{chunk['char']}' → ...{chunk['context'][:100]}...")
                    
                    if 'dashes' in results:
                        dash_data = results['dashes']
                        total_dashes = len(dash_data.get('all_dash_sentences', []))
                        double_dashes = len(dash_data.get('double_dash_sentences', []))
                        
                        with st.expander(f"— Dashes ({total_dashes} sentences with dashes, {double_dashes} with TWO dashes)"):
                            
                            # Сначала показываем предложения с двумя тире (как вы просили)
                            if dash_data.get('double_dash_sentences'):
                                st.markdown("### 🔴 Sentences with TWO dashes (— —)")
                                for i, item in enumerate(dash_data['double_dash_sentences']):
                                    # ПОЛНОЕ предложение, без обрезки
                                    st.markdown(f"**{i+1}.** {item['sentence']}")
                                    st.caption(f"*Sentence {item.get('sentence_idx', i)}, words: {item['word_count']}*")
                                    if i < len(dash_data['double_dash_sentences']) - 1:
                                        st.divider()
                            
                            # Затем показываем все остальные предложения с тире
                            other_dashes = [d for d in dash_data.get('all_dash_sentences', []) 
                                           if d.get('dash_count', 0) != 2]
                            if other_dashes:
                                if dash_data.get('double_dash_sentences'):
                                    st.markdown("### Other sentences with dashes")
                                for i, item in enumerate(other_dashes):
                                    # ПОЛНОЕ предложение, без обрезки
                                    st.markdown(f"**{i+1}.** {item['sentence']}")
                                    st.caption(f"*Dashes: {item['dash_count']}, words: {item['word_count']}*")
                                    if i < len(other_dashes) - 1:
                                        st.divider()
                                        
                    if 'tortured_phrases' in results:
                        tp_data = results['tortured_phrases']
                        if tp_data['total_occurrences'] > 0:
                            st.warning(f"⚠️ **Tortured Phrases Detected**: {tp_data['total_occurrences']} occurrences of {tp_data['unique_tortured_count']} unique tortured phrases")
                            
                            # Показываем топ-3 самые частые
                            if tp_data['tortured_phrases_found']:
                                top_phrases = sorted(tp_data['tortured_phrases_found'], 
                                                    key=lambda x: x['count'], reverse=True)[:3]
                                phrases_text = ", ".join([f"'{p['tortured']}' ({p['count']})" for p in top_phrases])
                                st.caption(f"Most frequent: {phrases_text}")
                
                with col2:
                    if 'punctuation' in results:
                        with st.expander(f"Semicolons ({len(results['punctuation']['all_semicolon_contexts'])} found)"):
                            for ctx in results['punctuation']['all_semicolon_contexts'][:30]:
                                st.code(f"{'📏 Contrast' if ctx['is_contrast'] else '➡️'} ...{ctx['context'][:150]}...")
                    
                    if 'apostrophe' in results:
                        with st.expander(f"Apostrophes ({len(results['apostrophe']['all_apostrophes'])} found)"):
                            st.write(", ".join(results['apostrophe']['all_apostrophes'][:100]))
            
            with tabs[1]:
                if 'enumeration' in results:
                    st.subheader(f"Enumerations found: {len(results['enumeration']['all_enumerations'])}")
                    for i, enum in enumerate(results['enumeration']['all_enumerations'][:100]):
                        st.markdown(f"{i+1}. `{enum}`")
                
                if 'phrases' in results:
                    with st.expander(f"AI Phrases ({len(results['phrases']['all_phrase_occurrences'])} occurrences)"):
                        for occ in results['phrases']['all_phrase_occurrences'][:50]:
                            st.markdown(f"**{occ['phrase']}** → {occ['context'][:150]}")
            
            with tabs[2]:
                if 'repetitiveness' in results:
                    st.subheader("Repeated Phrases")
                    for rep in results['repetitiveness']['all_repetitions'][:50]:
                        st.markdown(f"**{rep['ngram']}** — {rep['count']} times")
            
            with tabs[3]:
                if 'lexical_diversity' in results:
                    l = results['lexical_diversity']
                    col1, col2 = st.columns(2)
                    with col1:
                        st.metric("TTR", f"{l['ttr']:.3f}")
                        st.metric("MTLD", f"{l.get('mtld', 0):.1f}")
                        st.metric("MATTR", f"{l.get('mattr', 0):.3f}")
                    with col2:
                        st.metric("HD-D", f"{l.get('hdd', 0):.3f}")
                        st.metric("Hapax Ratio", f"{l['hapax_ratio']:.3f}")
                        st.metric("Unique Words", f"{len(set(l.get('words', [])))}")
            
            with tabs[4]:
                st.subheader("All Found Examples by Category")
                
                # Parentheses
                if 'parenthesis' in results and results['parenthesis']['all_parentheses']:
                    with st.expander(f"📝 Parentheses ({len(results['parenthesis']['all_parentheses'])} found)"):
                        for p in results['parenthesis']['all_parentheses'][:100]:
                            st.markdown(f"({p['text']}) — *{p['word_count']} words*")
                
                # Enumerations
                if 'enumeration' in results and results['enumeration']['all_enumerations']:
                    with st.expander(f"🔢 Three-item enumerations ({len(results['enumeration']['all_enumerations'])} found)"):
                        for i, enum in enumerate(results['enumeration']['all_enumerations']):
                            st.markdown(f"**{i+1}.** In sentence: *{enum['full_sentence']}*")
                            st.markdown(f"   → Enumeration: `{enum['enumeration']}`")
                            if i < len(results['enumeration']['all_enumerations']) - 1:
                                st.divider()
                
                # Apostrophes
                if 'apostrophe' in results and results['apostrophe']['all_apostrophes']:
                    with st.expander(f"🔤 Apostrophes ({len(results['apostrophe']['all_apostrophes'])} found)"):
                        st.write(", ".join(results['apostrophe']['all_apostrophes'][:200]))
                
                # AI Phrases
                if 'phrases' in results and results['phrases']['all_phrase_occurrences']:
                    with st.expander(f"🤖 AI Phrases ({len(results['phrases']['all_phrase_occurrences'])} occurrences)"):
                        for occ in results['phrases']['all_phrase_occurrences'][:100]:
                            st.markdown(f"**{occ['phrase']}** → {occ['context'][:150]}")
                
                # Semicolons
                if 'punctuation' in results and results['punctuation']['all_semicolon_contexts']:
                    with st.expander(f"‼️ Semicolons ({len(results['punctuation']['all_semicolon_contexts'])} found)"):
                        for ctx in results['punctuation']['all_semicolon_contexts'][:50]:
                            st.code(f"...{ctx['context'][:200]}...")
                
                # Repetitions
                if 'repetitiveness' in results and results['repetitiveness']['all_repetitions']:
                    with st.expander(f"🔄 Repetitions ({len(results['repetitiveness']['all_repetitions'])} found)"):
                        for rep in results['repetitiveness']['all_repetitions'][:100]:
                            st.markdown(f"**{rep['ngram']}** — {rep['count']} times")
                
                # Dashes
                if 'dashes' in results and (results['dashes']['all_dash_sentences'] or results['dashes']['double_dash_sentences']):
                    
                    if results['dashes']['double_dash_sentences']:
                        with st.expander(f"— Sentences with TWO dashes ({len(results['dashes']['double_dash_sentences'])} found) - CRITICAL FOR DETECTION"):
                            st.markdown("These sentences contain the **— —** pattern (two em-dashes):")
                            for i, item in enumerate(results['dashes']['double_dash_sentences']):
                                st.markdown(f"**{i+1}.** {item['sentence']}")
                                st.caption(f"*Sentence {item['sentence_idx']}, words: {item['word_count']}*")
                                st.divider()
                    
                    with st.expander(f"— All sentences with dashes ({len(results['dashes']['all_dash_sentences'])} total)"):
                        for i, item in enumerate(results['dashes']['all_dash_sentences']):
                            emoji = "🔴" if item['dash_count'] >= 3 else "🟡" if item['dash_count'] == 2 else "⚪"
                            st.markdown(f"{emoji} **{i+1}.** {item['sentence']}")
                            st.caption(f"*Dashes: {item['dash_count']}, words: {item['word_count']}*")
                            if i < len(results['dashes']['all_dash_sentences']) - 1:
                                st.divider()
                
                # Unicode artifacts
                if 'unicode' in results and results['unicode']['all_suspicious_chunks']:
                    with st.expander(f"🔣 Unicode Artifacts ({len(results['unicode']['all_suspicious_chunks'])} found)"):
                        for chunk in results['unicode']['all_suspicious_chunks'][:100]:
                            st.code(f"Character '{chunk['char']}' ({chunk['type']}) → ...{chunk['context'][:150]}...")

                # Tortured Phrases
                if 'tortured_phrases' in results and results['tortured_phrases']['all_occurrences']:
                    with st.expander(f"🔴 Tortured Phrases ({len(results['tortured_phrases']['all_occurrences'])} occurrences)"):
                        # Группируем по фразам
                        tortured_dict = {}
                        for occ in results['tortured_phrases']['all_occurrences']:
                            key = f"{occ['tortured']} → {occ['correct']}"
                            if key not in tortured_dict:
                                tortured_dict[key] = []
                            tortured_dict[key].append(occ)
                        
                        # Показываем статистику по группам
                        for phrase_group, occurrences in tortured_dict.items():
                            st.markdown(f"### {phrase_group} ({len(occurrences)} times)")
                            for i, occ in enumerate(occurrences[:10]):  # Показываем до 10 примеров на группу
                                st.markdown(f"**{i+1}.** Context: *{occ['context']}*")
                                if i < min(len(occurrences), 10) - 1:
                                    st.divider()
                            if len(occurrences) > 10:
                                st.caption(f"... and {len(occurrences) - 10} more occurrences")
            
            # Action buttons
            col1, col2, col3, col4 = st.columns([1, 1, 1, 1])
            with col1:
                if st.button("← New Analysis", use_container_width=True):
                    st.session_state.step = 1
                    st.session_state.analyze_clicked = False
                    st.session_state.file_uploaded = False
                    st.rerun()

            with col2:
                if st.button("📥 PDF Full Report", use_container_width=True):
                    with st.spinner("Generating full PDF report..."):
                        try:
                            pdf_data = {
                                'text': text,
                                'sentences': sentences,
                                'results': results,
                                'integrated': integrated,
                                'text_statistics': text_stats
                            }
                            
                            # Generate full report
                            pdf_bytes = generate_enhanced_pdf_report(pdf_data, report_type="full")
                            
                            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                            filename = f"ctai_detector_full_report_{timestamp}.pdf"
                            
                            st.download_button(
                                label="💾 Save Full Report",
                                data=pdf_bytes,
                                file_name=filename,
                                mime="application/pdf",
                                key="pdf_full_download"
                            )
                        except Exception as e:
                            st.error(f"Error generating PDF: {str(e)}")
                            import traceback
                            st.code(traceback.format_exc())

            with col3:
                if st.button("📄 PDF Concise Report", use_container_width=True):
                    with st.spinner("Generating concise PDF report..."):
                        try:
                            pdf_data = {
                                'text': text,
                                'sentences': sentences,
                                'results': results,
                                'integrated': integrated,
                                'text_statistics': text_stats
                            }
                            
                            # Generate concise report
                            pdf_bytes = generate_enhanced_pdf_report(pdf_data, report_type="concise")
                            
                            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                            filename = f"ctai_detector_concise_report_{timestamp}.pdf"
                            
                            st.download_button(
                                label="💾 Save Concise Report",
                                data=pdf_bytes,
                                file_name=filename,
                                mime="application/pdf",
                                key="pdf_concise_download"
                            )
                        except Exception as e:
                            st.error(f"Error generating PDF: {str(e)}")
                            import traceback
                            st.code(traceback.format_exc())
                    
            with col4:
                st.empty()

if __name__ == "__main__":
    main()
